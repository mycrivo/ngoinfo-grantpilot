"""Shared python-docx presentation primitives for client-facing document exports.

Generic only — no M&E or proposal domain logic. Used by core proposal export and
M&E report export renderers.
"""

from __future__ import annotations

import re
from datetime import date, datetime
from typing import Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# Print adaptation of NGOINFO_BRAND_GUIDELINES / BRAND_AND_FRONTEND_SPEC tokens.
NAVY = RGBColor(0x1A, 0x1F, 0x71)
TEXT_PRIMARY = RGBColor(0x1F, 0x29, 0x37)
TEXT_MUTED = RGBColor(0x64, 0x74, 0x8B)
BODY_FONT = "Calibri"
BODY_SIZE_PT = 11
HEADING_SIZES_PT = {1: 16, 2: 14, 3: 12}

ASSUMPTIONS_APPENDIX_TITLE = "Assumptions & Caveats"

SECTION_NOT_GENERATED = "[Section not generated]"
NOT_GENERATED_PREFIX = "[Not generated:"
NOT_GENERATED_RE = re.compile(r"^\[Not generated:[^\]]*\]\s*$", re.IGNORECASE)
MARKDOWN_HEADING_PREFIX_RE = re.compile(r"^#+\s*")


def format_client_date(when: datetime | date) -> str:
    """Human-readable date for client-facing cover blocks."""
    if isinstance(when, datetime):
        when = when.date()
    return f"{when.day} {when.strftime('%B %Y')}"


def strip_markdown_heading_prefix(text: str) -> str:
    """Remove accidental leading markdown '#' markers from heading text."""
    return MARKDOWN_HEADING_PREFIX_RE.sub("", text).strip()


def is_suppressed_placeholder(text: str) -> bool:
    """True when paragraph text is an internal artifact that must not appear in client docs."""
    stripped = text.strip()
    if not stripped:
        return False
    if stripped == SECTION_NOT_GENERATED:
        return True
    if stripped.startswith(NOT_GENERATED_PREFIX):
        return True
    return bool(NOT_GENERATED_RE.match(stripped))


def collect_deduped_assumptions(items: Sequence[str]) -> list[str]:
    deduped: list[str] = []
    seen: set[str] = set()
    for item in items:
        normalized = str(item).strip()
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        deduped.append(normalized)
    return deduped


def apply_house_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE_PT)
    normal.font.color.rgb = TEXT_PRIMARY

    for level, size_pt in HEADING_SIZES_PT.items():
        style = document.styles[f"Heading {level}"]
        style.font.name = BODY_FONT
        style.font.size = Pt(size_pt)
        style.font.color.rgb = NAVY
        style.font.bold = True


def add_branded_title_block(
    document: Document,
    *,
    org_name: str,
    document_title: str,
    subtitle_lines: Sequence[str] = (),
    document_date: datetime | date | None = None,
) -> None:
    """Client-facing cover: org name, document title, optional subtitles, human date."""
    org_para = document.add_paragraph()
    org_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    org_run = org_para.add_run(org_name.strip() or "Organisation")
    org_run.bold = True
    org_run.font.name = BODY_FONT
    org_run.font.size = Pt(20)
    org_run.font.color.rgb = NAVY

    title_para = document.add_paragraph()
    title_run = title_para.add_run(document_title.strip())
    title_run.bold = True
    title_run.font.name = BODY_FONT
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = NAVY

    for line in subtitle_lines:
        line_text = line.strip()
        if not line_text:
            continue
        sub_para = document.add_paragraph()
        sub_run = sub_para.add_run(line_text)
        sub_run.font.name = BODY_FONT
        sub_run.font.size = Pt(BODY_SIZE_PT)
        sub_run.font.color.rgb = TEXT_MUTED

    if document_date is not None:
        date_para = document.add_paragraph()
        date_run = date_para.add_run(format_client_date(document_date))
        date_run.font.name = BODY_FONT
        date_run.font.size = Pt(BODY_SIZE_PT)
        date_run.font.color.rgb = TEXT_MUTED

    document.add_paragraph()


def _append_field(run, field_code: str) -> None:
    """Insert a Word field (e.g. PAGE, NUMPAGES) using python-docx-supported OOXML."""
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(end)


def add_document_footer(
    document: Document,
    *,
    org_name: str,
    document_label: str | None = None,
    include_page_numbers: bool = True,
) -> None:
    """Footer with org name and optional page numbering on all document sections."""
    for section in document.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        if footer.paragraphs:
            para = footer.paragraphs[0]
            element = para._element
            for child in list(element):
                element.remove(child)
        else:
            para = footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        label_parts = [part.strip() for part in (org_name, document_label) if part and part.strip()]
        if label_parts:
            text_run = para.add_run("  |  ".join(label_parts))
            text_run.font.name = BODY_FONT
            text_run.font.size = Pt(9)
            text_run.font.color.rgb = TEXT_MUTED

        if include_page_numbers:
            if label_parts:
                sep_run = para.add_run("  |  Page ")
            else:
                sep_run = para.add_run("Page ")
            sep_run.font.name = BODY_FONT
            sep_run.font.size = Pt(9)
            sep_run.font.color.rgb = TEXT_MUTED

            page_run = para.add_run()
            page_run.font.name = BODY_FONT
            page_run.font.size = Pt(9)
            page_run.font.color.rgb = TEXT_MUTED
            _append_field(page_run, "PAGE")

            of_run = para.add_run(" of ")
            of_run.font.name = BODY_FONT
            of_run.font.size = Pt(9)
            of_run.font.color.rgb = TEXT_MUTED

            total_run = para.add_run()
            total_run.font.name = BODY_FONT
            total_run.font.size = Pt(9)
            total_run.font.color.rgb = TEXT_MUTED
            _append_field(total_run, "NUMPAGES")


def add_assumptions_appendix(document: Document, assumptions: Sequence[str]) -> None:
    """Single end-of-document assumptions block; preserves all assumption text."""
    deduped = collect_deduped_assumptions(assumptions)
    if not deduped:
        return
    document.add_page_break()
    document.add_heading(ASSUMPTIONS_APPENDIX_TITLE, level=1)
    for item in deduped:
        document.add_paragraph(item, style="List Bullet")
