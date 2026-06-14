"""Render donor_reports.content_json to a funder-structured .docx."""

from __future__ import annotations

import re
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document

from app.core.docx_presentation import (
    add_assumptions_appendix,
    add_branded_title_block,
    add_document_footer,
    apply_house_styles,
    strip_markdown_heading_prefix,
)
from app.reports.export.kb_table_renderer import (
    is_honest_empty_rows,
    table_headers_for_definition,
    table_rows_for_definition,
)
from app.reports.services.ngo_text_redaction import redact_internal_identifiers

_REPO_ROOT = Path(__file__).resolve().parents[3]

_TABLE_ROW_RE = re.compile(r"^\|.+\|$")
_TABLE_SEP_RE = re.compile(r"^\|[\s:\-|]+\|$")


def resolve_docx_template_path(docx_template_ref: str | None) -> Path | None:
    """Return a readable base .docx path, or None if ref is missing or not on disk."""
    ref = (docx_template_ref or "").strip()
    if not ref:
        return None
    direct = Path(ref)
    if direct.is_file():
        return direct
    from_repo = _REPO_ROOT / ref
    if from_repo.is_file():
        return from_repo
    return None


def _apply_basic_styles(document: Document) -> None:
    apply_house_styles(document)


def _strip_internal_tokens(text: str) -> str:
    """Route NGO-facing prose through the single identifier-redaction chokepoint."""
    return redact_internal_identifiers(text)


# D2: a model-emitted caveat claiming a table was unfillable because the schema
# lacked a table field is false (the template declares the table; the engine now
# renders it). Suppress only that misattribution; the engine emits the TRUE
# reason for any genuinely honest-empty table.
_FALSE_TABLE_SCHEMA_RE = re.compile(
    r"(?:output\s+)?schema\s+(?:did\s+not|does\s+not|did\s*n[o']t|does\s*n[o']t)\s+"
    r"(?:include|provide|define|contain|have)\s+a?\s*table",
    re.IGNORECASE,
)


def _is_false_table_schema_attribution(text: str) -> bool:
    return bool(_FALSE_TABLE_SCHEMA_RE.search(text or ""))


def _honest_empty_table_caveat(label: str) -> str:
    name = label.strip() or "required"
    return (
        f"The \"{name}\" table is included, but no verified figures were available "
        f"in the submitted records to populate it."
    )


def _parse_markdown_table_block(lines: list[str]) -> tuple[list[str], list[list[str]]] | None:
    if len(lines) < 2:
        return None
    if not all(_TABLE_ROW_RE.match(line.strip()) for line in lines[:2]):
        return None
    if not _TABLE_SEP_RE.match(lines[1].strip()):
        return None
    header = [cell.strip() for cell in lines[0].strip().strip("|").split("|")]
    rows: list[list[str]] = []
    for line in lines[2:]:
        if not _TABLE_ROW_RE.match(line.strip()):
            break
        rows.append([cell.strip() for cell in line.strip().strip("|").split("|")])
    if not header:
        return None
    return header, rows


def _add_word_table(document: Document, header: list[str], rows: list[list[str]]) -> None:
    col_count = max(len(header), max((len(r) for r in rows), default=0))
    if col_count == 0:
        return
    table = document.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"
    for col_idx in range(col_count):
        cell_text = header[col_idx] if col_idx < len(header) else ""
        table.rows[0].cells[col_idx].text = redact_internal_identifiers(cell_text)
    for row_idx, row in enumerate(rows, start=1):
        for col_idx in range(col_count):
            cell_text = row[col_idx] if col_idx < len(row) else ""
            table.rows[row_idx].cells[col_idx].text = redact_internal_identifiers(cell_text)


def _render_section_body(
    document: Document,
    text: str,
) -> None:
    sanitized = _strip_internal_tokens(text)
    if not sanitized:
        return

    lines = sanitized.splitlines()
    buffer: list[str] = []
    idx = 0
    while idx < len(lines):
        line = lines[idx]
        if _TABLE_ROW_RE.match(line.strip()):
            table_lines = [line]
            j = idx + 1
            while j < len(lines) and _TABLE_ROW_RE.match(lines[j].strip()):
                table_lines.append(lines[j])
                j += 1
            parsed = _parse_markdown_table_block(table_lines)
            if parsed:
                for para in buffer:
                    if para.strip():
                        document.add_paragraph(para.strip())
                buffer.clear()
                header, rows = parsed
                _add_word_table(document, header, rows)
                idx = j
                continue
        buffer.append(line)
        idx += 1

    for para in buffer:
        stripped = para.strip()
        if stripped:
            document.add_paragraph(stripped)


def render_donor_report_docx(
    *,
    content_json: dict[str, Any],
    template_sections: list[Any],
    format_rules_json: dict[str, Any],
    terminology_map_json: dict[str, Any],
    docx_template_ref: str | None,
    reporting_period_start: str,
    reporting_period_end: str,
    funder_name: str,
    template_name: str,
    ngo_name: str = "",
    generated_at: datetime | None = None,
    knowledge_bank_json: dict[str, Any] | None = None,
    gap_analysis_json: dict[str, Any] | None = None,
) -> tuple[bytes, str]:
    """
    Build export bytes and report which template path was used.

    Returns (docx_bytes, render_mode) where render_mode is 'base_template' or 'from_scratch'.
    """
    base_path = resolve_docx_template_path(docx_template_ref)
    if base_path is not None:
        document = Document(str(base_path))
        render_mode = "base_template"
    else:
        document = Document()
        render_mode = "from_scratch"

    _apply_basic_styles(document)

    when = generated_at or datetime.now(timezone.utc)
    document_title = str(
        format_rules_json.get("document_title") or template_name or "Donor Report"
    )
    add_branded_title_block(
        document,
        org_name=ngo_name or "Organisation",
        document_title=f"Donor Report — {funder_name}",
        subtitle_lines=[
            f"Reporting period: {reporting_period_start} to {reporting_period_end}",
            document_title if document_title != funder_name else "",
        ],
        document_date=when,
    )
    document.add_page_break()

    sections_by_key: dict[str, dict[str, Any]] = {}
    for item in content_json.get("sections") or []:
        if isinstance(item, dict) and item.get("section_key"):
            sections_by_key[str(item["section_key"])] = item

    collected_assumptions: list[str] = []
    table_caveats: list[str] = []
    kb_facts = dict((knowledge_bank_json or {}).get("facts") or {})

    for template_section in template_sections:
        if not isinstance(template_section, dict):
            continue
        section_key = str(template_section.get("section_key") or "")
        heading = strip_markdown_heading_prefix(
            str(template_section.get("label") or section_key)
        )
        # E1: render the funder-authored label verbatim (no canonical_to_funder
        # substitution inside labels, which produced broken mid-sentence headings).
        document.add_heading(redact_internal_identifiers(heading), level=1)

        for table_def in template_section.get("required_tables") or []:
            if not isinstance(table_def, dict):
                continue
            table_label = strip_markdown_heading_prefix(str(table_def.get("label") or ""))
            if table_label:
                document.add_heading(redact_internal_identifiers(table_label), level=2)
            kb_rows = table_rows_for_definition(
                table_def=table_def,
                facts=kb_facts,
                format_rules_json=format_rules_json,
                gap_analysis=gap_analysis_json,
            )
            if kb_rows:
                header = table_headers_for_definition(table_def)
                if header:
                    _add_word_table(document, header, kb_rows)
                    if table_label and is_honest_empty_rows(kb_rows):
                        # D2: state the TRUE reason a declared table could not fill.
                        table_caveats.append(_honest_empty_table_caveat(table_label))

        section = sections_by_key.get(section_key)
        if section is None:
            continue

        status = section.get("generation_status")
        content = section.get("content") or {}
        text = str(content.get("text") or "")
        section_assumptions = content.get("assumptions") or []
        for assumption in section_assumptions:
            if not assumption:
                continue
            # D2: drop the false "schema did not include a table field"
            # misattribution; the engine emits the true table caveat instead.
            if _is_false_table_schema_attribution(str(assumption)):
                continue
            redacted = redact_internal_identifiers(str(assumption))
            if redacted.strip():
                collected_assumptions.append(redacted)

        if status in ("GENERATED", "AWAITING_REVIEW", "ACCEPTED") and text.strip():
            _render_section_body(document, text)
        elif status == "FAILED":
            pass
        # Empty / missing content: heading only — no internal placeholder paragraphs.

    add_assumptions_appendix(document, collected_assumptions + table_caveats)
    add_document_footer(
        document,
        org_name=ngo_name or "Organisation",
        document_label="Donor Report",
    )

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue(), render_mode


def build_export_filename(
    *,
    funder_name: str,
    template_name: str,
    reporting_period_start: str,
    reporting_period_end: str,
) -> str:
    def _slug(value: str) -> str:
        return re.sub(r"[^A-Za-z0-9]+", "_", value).strip("_")[:40] or "report"

    return (
        f"{_slug(funder_name)}_{_slug(template_name)}_"
        f"{reporting_period_start}_{reporting_period_end}.docx"
    )
