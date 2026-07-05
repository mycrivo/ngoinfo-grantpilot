"""Machine-check exported docx against content and template contract (P3-7 S4)."""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from io import BytesIO
from typing import Any

from docx import Document

from app.reports.export.kb_table_renderer import (
    table_headers_for_definition,
    table_rows_for_definition,
)
from app.reports.gap.section_visibility import visible_sections_for_context
from app.reports.services.section_prose import (
    MIN_SECTION_PROSE_CHARS,
    has_non_empty_prose,
    section_meets_minimum_substance,
)


# --- Identifier-leak tripwire (Package 1, E2) -------------------------------
# Independent of the redaction chokepoint: this asserts the *target* (no internal
# identifier reaches NGO-facing output), so the chokepoint must satisfy it rather
# than the assertion being weakened to pass. Patterns require the no-internal-space
# / namespace-prefixed shapes that identifiers have but honest prose does not, so
# times (10:30), ratios (3:1), scripture (John 3:16), decimals (4.2), URLs, and
# domains do not trip it.
#
# Package D WIDENS this tripwire (never narrows it) for the spreadsheet-provenance
# leak class the diagnosis surfaced: (1) A1-notation cell references ("Table2!C12")
# and (2) the em-dash facet suffix carrying a cell ref ("- budget (Table2!C12)") on
# entity labels. Both require the sheet "!" + Column/Row shape, so legitimate prose
# (times, ratios, numbers, em-dashes, "(Q1)", "(2024)") never trips it.
_A1_CELL_REF = r"[A-Za-z][A-Za-z0-9_]*![A-Z]{1,3}[0-9]+(?::[A-Z]{1,3}[0-9]+)?"
_LEAK_PATTERNS: tuple[tuple[str, re.Pattern[str]], ...] = (
    ("colon_item_key", re.compile(r"[A-Za-z][A-Za-z0-9_]*(?::[A-Za-z0-9_]+){2,}")),
    (
        "schema_dotted_path",
        re.compile(
            r"\b(?:financials|indicators?|reporting|objectives|outcomes)(?:\.[A-Za-z0-9_]+){2,}"
        ),
    ),
    ("citation_marker", re.compile(r"\[(?:fact|gap):[^\]]+\]", re.IGNORECASE)),
    ("archetype_token", re.compile(r"\bARCH_[A-Z0-9_]+\b")),
    ("enum_value", re.compile(r"\b(?:cannot_provide|not_applicable)\b")),
    ("generic_placeholder", re.compile(r"the required template items")),
    ("spreadsheet_cell_ref", re.compile(r"\b" + _A1_CELL_REF + r"\b")),
    (
        "entity_facet_provenance",
        re.compile(
            r"[\u2013\u2014]\s*"
            r"(?:budget|actual(?:\s+spend)?|target|milestone|spend|forecast|planned)"
            r"\s*\(" + _A1_CELL_REF + r"\)",
            re.IGNORECASE,
        ),
    ),
)


def scan_identifier_leaks(text: str) -> list[str]:
    """Return violation tags for every internal-identifier pattern found in NGO text."""
    violations: list[str] = []
    for name, pattern in _LEAK_PATTERNS:
        match = pattern.search(text)
        if match:
            violations.append(f"identifier_leak:{name}:{match.group(0)[:60]}")
    return violations


@dataclass
class DocxExportAssertionReport:
    violations: list[str] = field(default_factory=list)

    @property
    def passed(self) -> bool:
        return not self.violations

    def to_summary_dict(self) -> dict[str, Any]:
        return {
            "passed": self.passed,
            "violation_count": len(self.violations),
            "violations": self.violations,
        }


def _docx_table_count(docx_bytes: bytes) -> int:
    document = Document(BytesIO(docx_bytes))
    return len(document.tables)


def _docx_plaintext(docx_bytes: bytes) -> str:
    document = Document(BytesIO(docx_bytes))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def assert_export_docx(
    *,
    docx_bytes: bytes,
    content_json: dict[str, Any],
    template_sections: list[dict[str, Any]],
    format_rules_json: dict[str, Any] | None = None,
    knowledge_bank_json: dict[str, Any] | None = None,
    gap_analysis_json: dict[str, Any] | None = None,
    report_context: dict[str, Any] | None = None,
) -> DocxExportAssertionReport:
    """Assert non-empty NGO section prose, minimum substance, and KB-backed tables."""
    violations: list[str] = []
    ctx = report_context or {"report_type": "annual"}
    visible = visible_sections_for_context(
        template_sections,
        report_context=ctx,
        include_funder_owned=False,
    )
    visible_keys = {str(s.get("section_key") or "") for s in visible}

    sections_by_key: dict[str, dict[str, Any]] = {}
    for section in content_json.get("sections") or []:
        if isinstance(section, dict) and section.get("section_key"):
            sections_by_key[str(section["section_key"])] = section

    for template_section in visible:
        key = str(template_section.get("section_key") or "")
        section = sections_by_key.get(key)
        if section is None:
            violations.append(f"missing_section:{key}")
            continue
        status = section.get("generation_status")
        if status in ("GENERATED", "AWAITING_REVIEW", "ACCEPTED"):
            if not has_non_empty_prose(section):
                violations.append(f"empty_prose:{key}")
            elif not section_meets_minimum_substance(section):
                violations.append(
                    f"insufficient_prose:{key}:min={MIN_SECTION_PROSE_CHARS}"
                )

    kb_facts = dict((knowledge_bank_json or {}).get("facts") or {})
    expected_tables = 0
    for template_section in template_sections:
        if not isinstance(template_section, dict):
            continue
        key = str(template_section.get("section_key") or "")
        if key not in visible_keys:
            continue
        for table_def in template_section.get("required_tables") or []:
            if not isinstance(table_def, dict):
                continue
            rows = table_rows_for_definition(
                table_def=table_def,
                facts=kb_facts,
                format_rules_json=format_rules_json or {},
                gap_analysis=gap_analysis_json,
            )
            if rows:
                expected_tables += 1
                header = table_headers_for_definition(table_def)
                if not header:
                    violations.append(
                        f"table_header_missing:{key}:{table_def.get('table_key')}"
                    )

    actual_tables = _docx_table_count(docx_bytes)
    if expected_tables > 0 and actual_tables < expected_tables:
        violations.append(
            f"table_count_low:expected>={expected_tables}:actual={actual_tables}"
        )

    plaintext = _docx_plaintext(docx_bytes)
    # Identifier-leak tripwire over the entire rendered document (body prose,
    # table cells, and the Assumptions & Caveats appendix). Always runs.
    violations.extend(scan_identifier_leaks(plaintext))

    if violations:
        return DocxExportAssertionReport(violations=violations)

    for template_section in visible:
        key = str(template_section.get("section_key") or "")
        section = sections_by_key.get(key)
        if not section or not has_non_empty_prose(section):
            continue
        snippet = section.get("content", {}).get("text", "")[:60].strip()
        if snippet and snippet not in plaintext:
            violations.append(f"prose_not_in_docx:{key}")

    return DocxExportAssertionReport(violations=violations)
