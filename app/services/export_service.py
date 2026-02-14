from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO
from uuid import UUID

from docx import Document
from docx.shared import Pt
from sqlalchemy.orm import Session

from app.core.errors import DomainError
from app.models.funding_opportunity import FundingOpportunity
from app.models.ngo_profile import NGOProfile
from app.models.usage_ledger import UsageActionType
from app.services.proposal_service import ProposalService
from app.services.quota_service import record_usage


class ExportService:
    def __init__(self, db_session: Session) -> None:
        self.db = db_session
        self.proposals = ProposalService(db_session)

    def export_docx(self, *, user, proposal_id: UUID, export_format: str) -> tuple[bytes, str]:
        normalized_format = (export_format or "").strip().upper()
        if normalized_format != "DOCX":
            raise DomainError(
                error_code="UNSUPPORTED_FORMAT",
                message="Only DOCX export is supported.",
                status_code=422,
            )

        proposal = self.proposals.get_proposal(user=user, proposal_id=proposal_id)
        opportunity = self.db.get(FundingOpportunity, proposal.funding_opportunity_id)
        profile = self.db.query(NGOProfile).filter(NGOProfile.user_id == user.id).one_or_none()

        opportunity_title = opportunity.title if opportunity else "Untitled Opportunity"
        ngo_name = profile.organization_name if profile else user.email
        generated_at = datetime.now(timezone.utc)
        docx_bytes = _build_docx_bytes(
            content_json=proposal.content_json or {},
            opportunity_title=opportunity_title,
            ngo_name=ngo_name,
            generated_at=generated_at,
            proposal_id=str(proposal.id),
            proposal_version=int(proposal.version),
        )

        idempotency_key = (
            f"docx_export:{user.id}:{proposal.id}:v{int(proposal.version)}"
        )
        record_usage(
            self.db,
            user.id,
            UsageActionType.DOCX_EXPORT.value,
            idempotency_key=idempotency_key,
            commit=True,
        )

        filename = f"proposal-{proposal.id}.docx"
        return docx_bytes, filename


def _build_docx_bytes(
    *,
    content_json: dict,
    opportunity_title: str,
    ngo_name: str,
    generated_at: datetime,
    proposal_id: str,
    proposal_version: int,
) -> bytes:
    document = Document()
    _apply_basic_styles(document)

    document.add_heading(opportunity_title, level=0)
    document.add_paragraph(f"NGO: {ngo_name}")
    document.add_paragraph(f"Generated At (UTC): {generated_at.isoformat()}")
    document.add_paragraph(f"Proposal ID: {proposal_id}")
    document.add_paragraph(f"Version: {proposal_version}")
    document.add_page_break()

    sections = content_json.get("sections") or []
    assumptions: list[str] = []
    for section in sections:
        title = section.get("label") or "Untitled Section"
        status = section.get("generation_status")
        content = section.get("content") or {}
        text = content.get("text") or ""
        section_assumptions = content.get("assumptions") or []
        assumptions.extend([a for a in section_assumptions if a])

        document.add_heading(title, level=1)
        if status == "GENERATED":
            document.add_paragraph(text)
        else:
            document.add_paragraph("To be completed manually")

    if assumptions:
        document.add_page_break()
        document.add_heading("Assumptions Appendix", level=1)
        deduped = []
        seen = set()
        for item in assumptions:
            if item not in seen:
                seen.add(item)
                deduped.append(item)
        for item in deduped:
            document.add_paragraph(item, style="List Bullet")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _apply_basic_styles(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(12)
