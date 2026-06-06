"""Proposal content_json → client-facing DOCX bytes (python-docx, from scratch)."""

from __future__ import annotations

from datetime import datetime
from io import BytesIO

from docx import Document

from app.core.docx_presentation import (
    add_assumptions_appendix,
    add_branded_title_block,
    add_document_footer,
    apply_house_styles,
    strip_markdown_heading_prefix,
)


def build_proposal_docx_bytes(
    *,
    content_json: dict,
    opportunity_title: str,
    ngo_name: str,
    generated_at: datetime,
) -> bytes:
    document = Document()
    apply_house_styles(document)

    add_branded_title_block(
        document,
        org_name=ngo_name,
        document_title=f"Grant Proposal — {opportunity_title}",
        document_date=generated_at,
    )
    document.add_page_break()

    sections = content_json.get("sections") or []
    assumptions: list[str] = []
    for section in sections:
        title = strip_markdown_heading_prefix(section.get("label") or "Untitled Section")
        status = section.get("generation_status")
        content = section.get("content") or {}
        text = content.get("text") or ""
        section_assumptions = content.get("assumptions") or []
        assumptions.extend([a for a in section_assumptions if a])

        document.add_heading(title, level=1)
        if status == "GENERATED" and text.strip():
            document.add_paragraph(text)

    add_assumptions_appendix(document, assumptions)
    add_document_footer(
        document,
        org_name=ngo_name,
        document_label="Grant Proposal",
    )

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
