"""Unit tests for app.reports.export.docx_renderer — no API/DB dependencies."""

from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path

from docx import Document

from app.reports.export.docx_renderer import render_donor_report_docx

FCDO_TEMPLATE_PATH = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "artefacts"
    / "me_module"
    / "TEMPLATE_INSTANCE_FCDO.json"
)


def _docx_plaintext(docx_bytes: bytes) -> str:
    document = Document(BytesIO(docx_bytes))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    return "\n".join(parts)


def _load_fcdo_template_fields() -> dict:
    return json.loads(FCDO_TEMPLATE_PATH.read_text(encoding="utf-8"))


def _render_minimal(*, template_sections, content_json, terminology_map_json):
    return render_donor_report_docx(
        content_json=content_json,
        template_sections=template_sections,
        format_rules_json={"document_title": "Test Report"},
        terminology_map_json=terminology_map_json,
        docx_template_ref=None,
        reporting_period_start="2024-10-15",
        reporting_period_end="2025-10-14",
        funder_name="Test Funder",
        template_name="Test Template",
    )


def test_body_prose_unchanged_by_terminology_or_schema_keys():
    fcdo = _load_fcdo_template_fields()
    body = (
        "Risk management remained broadly stable. Year 1 monitoring returns did not report "
        "any major incident trend. actual spend against a budget of GBP 162,000. "
        "A smaller set of indicators fell below milestone. clear limitations that affected "
        "confidence in some results."
    )
    template_sections = [
        {
            "section_key": "risk_and_safeguarding",
            "label": "Risk, Assumptions and Safeguarding",
            "required_tables": [],
        }
    ]
    content_json = {
        "sections": [
            {
                "section_key": "risk_and_safeguarding",
                "generation_status": "ACCEPTED",
                "content": {"text": body},
            }
        ]
    }
    docx_bytes, _ = _render_minimal(
        template_sections=template_sections,
        content_json=content_json,
        terminology_map_json=fcdo["terminology_map_json"],
    )
    text = _docx_plaintext(docx_bytes)
    assert "Risk management remained broadly stable" in text
    assert "did not report any major incident trend" in text
    assert "against a budget of GBP 162,000" in text
    assert "fell below milestone" in text
    assert "clear limitations that affected confidence" in text
    assert "Risk rating / assumptions / controls management" not in text
    assert "did not Annual Review any" not in text
    assert "Budget / forecast and actual costs" not in text
    assert "fell below ." not in text


def test_body_removes_whole_citation_markers_without_orphan_brackets():
    fcdo = _load_fcdo_template_fields()
    body = (
        "Shortfall on OP2.2. [gap:risk_and_safeguarding:indicator:climate_environment_risk] "
        "Climate conditions held. [fact:financials.lines.op1_1.y1_actual] Spend was recorded."
    )
    template_sections = [
        {
            "section_key": "risk_and_safeguarding",
            "label": "Risk update",
            "required_tables": [],
        }
    ]
    content_json = {
        "sections": [
            {
                "section_key": "risk_and_safeguarding",
                "generation_status": "ACCEPTED",
                "content": {"text": body},
            }
        ]
    }
    docx_bytes, _ = _render_minimal(
        template_sections=template_sections,
        content_json=content_json,
        terminology_map_json=fcdo["terminology_map_json"],
    )
    text = _docx_plaintext(docx_bytes)
    assert "Shortfall on OP2.2. Climate conditions held. Spend was recorded." in text
    assert "fact:" not in text
    assert "gap:" not in text
    assert "[ [" not in text
    assert "[ ]" not in text


def test_kb_logframe_table_renders_in_docx():
    fcdo = _load_fcdo_template_fields()
    template_sections = fcdo["report_sections_json"]
    # Real recorded indicator fact shapes (op<n>_<n>.actual/.target), not the
    # deleted favourable distilled logframe_ar1_* skeleton keys.
    facts = {
        "indicators.op1_1_girls_reenrolled.actual": {"value": 684, "semantic_label": "Re-enrolled girls"},
        "indicators.op1_1_girls_reenrolled.target": {"value": 650},
    }
    content_json = {
        "sections": [
            {
                "section_key": "detailed_output_scoring",
                "generation_status": "ACCEPTED",
                "content": {"text": "Output scoring narrative for the review period."},
            }
        ]
    }
    docx_bytes, _ = render_donor_report_docx(
        content_json=content_json,
        template_sections=template_sections,
        format_rules_json=fcdo.get("format_rules_json") or {},
        terminology_map_json=fcdo.get("terminology_map_json") or {},
        docx_template_ref=None,
        reporting_period_start="2024-10-15",
        reporting_period_end="2025-10-14",
        funder_name="FCDO",
        template_name="Annual Review",
        knowledge_bank_json={"facts": facts},
    )
    document = Document(BytesIO(docx_bytes))
    assert len(document.tables) >= 1
    # The real actual (684) and target (650) land in their own columns of a row.
    matching_rows = [
        [c.text for c in row.cells]
        for table in document.tables
        for row in table.rows
        if "684" in [c.text for c in row.cells]
    ]
    assert matching_rows, "expected a logframe row populated from the real actual fact"
    row = matching_rows[0]
    assert "684" in row and "650" in row

    # E1: funder-authored labels render VERBATIM — no canonical_to_funder
    # substitution inside section/table headings.
    terminology_map_json = {
        "canonical_to_funder": {
            "risk": "Risk rating / assumptions / controls",
            "budget": "Budget / forecast and actual costs",
        }
    }
    template_sections = [
        {
            "section_key": "risk_and_safeguarding",
            "label": "Program risk review",
            "required_tables": [
                {
                    "table_key": "risk_register_update",
                    "label": "Annual budget summary",
                    "columns": [],
                }
            ],
        }
    ]
    content_json = {
        "sections": [
            {
                "section_key": "risk_and_safeguarding",
                "generation_status": "ACCEPTED",
                "content": {"text": "Body text only."},
            }
        ]
    }
    docx_bytes, _ = _render_minimal(
        template_sections=template_sections,
        content_json=content_json,
        terminology_map_json=terminology_map_json,
    )
    document = Document(BytesIO(docx_bytes))
    headings = [p.text for p in document.paragraphs if p.style.name.startswith("Heading")]
    assert any("Program risk review" in h for h in headings)
    assert any("Annual budget summary" in h for h in headings)
    assert not any("Risk rating / assumptions / controls" in h for h in headings)
    assert not any("Budget / forecast and actual costs" in h for h in headings)
    assert "Body text only." in _docx_plaintext(docx_bytes)
