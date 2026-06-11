"""Stage H export — docx render, storage persist, download route."""

from __future__ import annotations

import json
import uuid
from datetime import date, datetime, timedelta, timezone
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import func, select

import app.core.security as security
from app.core.config import get_settings
from app.core.config import get_settings as config_get_settings
from app.core.security import create_access_token
from app.db.session import get_db
from app.main import create_app
from app.models.user import User
from app.models.ngo_profile import NGOProfile
from app.core.errors import ForbiddenError
from app.models.usage_ledger import UsageActionType, UsageLedger
from app.services.quota_service import PLAN_IMPACT, get_entitlements, report_create_idempotency_key
from app.reports.export.docx_renderer import render_donor_report_docx
from app.reports.models.donor_report import DonorReport
from app.reports.models.enums import DonorReportStatus
from app.reports.models.funder_report_template import FunderReportTemplate
from app.reports.services.document_storage_service import DocumentStorageService
from app.reports.services.report_export_service import (
    ReportExportServiceError,
    export_and_persist,
)
from tests.worker_validation_seed import (
    create_worker_validation_sessionmaker,
    seed_user_plan,
)

FCDO_TEMPLATE_PATH = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "artefacts"
    / "me_module"
    / "TEMPLATE_INSTANCE_FCDO.json"
)
FCDO_CONTENT_PATH = (
    Path(__file__).resolve().parents[0]
    / "fixtures"
    / "export"
    / "fcdo_recorded_content_json.json"
)

get_settings.cache_clear()


@pytest.fixture(autouse=True)
def _reset_settings_cache():
    get_settings.cache_clear()
    security.get_settings = config_get_settings
    yield
    get_settings.cache_clear()
    security.get_settings = config_get_settings


class MemoryDocumentStorage:
    """In-memory R2 stand-in for dev-loop tests."""

    def __init__(self) -> None:
        self.objects: dict[str, tuple[bytes, str]] = {}

    @staticmethod
    def build_storage_ref(user_id, report_id, filename: str) -> str:
        return DocumentStorageService.build_storage_ref(user_id, report_id, filename)

    def upload_bytes(self, key: str, data: bytes, content_type: str) -> None:
        self.objects[key] = (data, content_type)

    def delete_object(self, key: str) -> None:
        self.objects.pop(key, None)

    def fetch_bytes(self, key: str) -> bytes:
        return self.objects[key][0]


def _docx_plaintext(docx_bytes: bytes) -> str:
    document = Document(BytesIO(docx_bytes))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _load_fcdo_template_fields() -> dict:
    return json.loads(FCDO_TEMPLATE_PATH.read_text(encoding="utf-8"))


def _seed_gate3_ready_report(session) -> tuple[uuid.UUID, uuid.UUID, MemoryDocumentStorage]:
    fcdo = _load_fcdo_template_fields()
    content_json = json.loads(FCDO_CONTENT_PATH.read_text(encoding="utf-8"))
    now = datetime.now(timezone.utc)
    user = User(
        id=uuid.uuid4(),
        email=f"export-test-{uuid.uuid4().hex[:8]}@example.org",
        auth_provider="email",
        created_at=now,
        updated_at=now,
    )
    template = FunderReportTemplate(
        id=uuid.uuid4(),
        funder_name=fcdo["funder_name"],
        template_name=fcdo["template_name"],
        region=fcdo.get("region", "UK"),
        reporting_frequency=fcdo.get("reporting_frequency", "annual"),
        report_sections_json=fcdo["report_sections_json"],
        format_rules_json=fcdo.get("format_rules_json", {}),
        terminology_map_json=fcdo.get("terminology_map_json", {}),
        docx_template_ref=fcdo.get("docx_template_ref"),
        is_active=True,
        version=fcdo.get("version", 1),
        created_at=now,
        updated_at=now,
    )
    kb = {
        "gate2_confirmed_at": now.isoformat(),
        "gate3_confirmed_at": now.isoformat(),
    }
    report = DonorReport(
        id=uuid.uuid4(),
        user_id=user.id,
        funder_report_template_id=template.id,
        reporting_period_start=date(2024, 10, 15),
        reporting_period_end=date(2025, 10, 14),
        status=DonorReportStatus.DRAFT.value,
        knowledge_bank_json=kb,
        gap_analysis_json={},
        indicator_actuals_json={},
        content_json=content_json,
        version=1,
        created_at=now,
        updated_at=now,
    )
    session.add_all([user, template, report])
    session.add(
        NGOProfile(
            id=uuid.uuid4(),
            user_id=user.id,
            organization_name="BridgeLight Education Trust",
            country_of_registration="Uganda",
            mission_statement="Girls' education access",
            focus_sectors=[],
            geographic_areas_of_work=[],
            target_groups=[],
            past_projects=[],
            profile_status="COMPLETE",
            completeness_score=100,
            missing_fields=[],
            funders_worked_with_before=[],
            created_at=now,
            updated_at=now,
        )
    )
    seed_user_plan(session, user.id, plan_name=PLAN_IMPACT)
    session.commit()
    return report.id, user.id, MemoryDocumentStorage()


@pytest.fixture
def export_db():
    return create_worker_validation_sessionmaker()


def test_render_docx_follows_template_order_and_surface_labels(export_db):
    fcdo = _load_fcdo_template_fields()
    content_json = json.loads(FCDO_CONTENT_PATH.read_text(encoding="utf-8"))
    docx_bytes, render_mode = render_donor_report_docx(
        content_json=content_json,
        template_sections=fcdo["report_sections_json"],
        format_rules_json=fcdo["format_rules_json"],
        terminology_map_json=fcdo["terminology_map_json"],
        docx_template_ref=fcdo["docx_template_ref"],
        reporting_period_start="2024-10-15",
        reporting_period_end="2025-10-14",
        funder_name=fcdo["funder_name"],
        template_name=fcdo["template_name"],
    )
    assert render_mode == "from_scratch"
    text = _docx_plaintext(docx_bytes)
    assert "A. Summary and Overview" in text
    assert "C. Detailed Output Scoring" in text
    assert "920,420" in text
    assert "steady progress" in text
    assert "summary_and_overview" not in text
    assert "detailed_output_scoring" not in text
    assert "fact:" not in text
    assert "gap:" not in text
    assert "ARCH_" not in text

    headings = [
        p.text
        for p in Document(BytesIO(docx_bytes)).paragraphs
        if p.style.name.startswith("Heading 1")
    ]
    template_labels = [
        s["label"] for s in fcdo["report_sections_json"] if s.get("label")
    ]
    for label in template_labels[:3]:
        assert label in headings


def test_export_and_persist_happy_path(export_db):
    session = export_db()
    report_id, _user_id, storage = _seed_gate3_ready_report(session)
    session.close()

    result = export_and_persist(export_db(), report_id, storage=storage)

    final = export_db()
    report = final.get(DonorReport, report_id)
    final.close()

    assert report.status == DonorReportStatus.COMPLETE.value
    export_meta = report.content_json["export"]
    assert export_meta["storage_ref"] == result.storage_ref
    assert export_meta["filename"].endswith(".docx")
    assert export_meta["template_version"] == 1
    assert storage.fetch_bytes(export_meta["storage_ref"])

    ledger = final.execute(
        select(func.count())
        .select_from(UsageLedger)
        .where(
            UsageLedger.user_id == _user_id,
            UsageLedger.event_type == UsageActionType.REPORT_CREATE.value,
            UsageLedger.idempotency_key == report_create_idempotency_key(report_id),
        )
    ).scalar_one()
    entitlements = get_entitlements(final, _user_id)
    assert ledger == 1
    assert entitlements["entitlements"]["reports"]["used"] == 1


def test_export_recomplete_does_not_recharge(export_db):
    session = export_db()
    report_id, user_id, storage = _seed_gate3_ready_report(session)
    session.close()

    export_and_persist(export_db(), report_id, storage=storage)
    export_and_persist(export_db(), report_id, storage=storage)

    final = export_db()
    ledger = final.execute(
        select(func.count())
        .select_from(UsageLedger)
        .where(
            UsageLedger.user_id == user_id,
            UsageLedger.event_type == UsageActionType.REPORT_CREATE.value,
            UsageLedger.idempotency_key == report_create_idempotency_key(report_id),
        )
    ).scalar_one()
    entitlements = get_entitlements(final, user_id)
    final.close()
    assert ledger == 1
    assert entitlements["entitlements"]["reports"]["used"] == 1


def test_export_cutover_idempotent_when_create_charge_already_exists(export_db):
    session = export_db()
    report_id, user_id, storage = _seed_gate3_ready_report(session)
    session.add(
        UsageLedger(
            id=uuid.uuid4(),
            user_id=user_id,
            event_type=UsageActionType.REPORT_CREATE.value,
            occurred_at=datetime.now(timezone.utc),
            idempotency_key=report_create_idempotency_key(report_id),
            metadata_json={},
        )
    )
    session.commit()
    session.close()

    export_and_persist(export_db(), report_id, storage=storage)

    final = export_db()
    ledger = final.execute(
        select(func.count())
        .select_from(UsageLedger)
        .where(
            UsageLedger.user_id == user_id,
            UsageLedger.event_type == UsageActionType.REPORT_CREATE.value,
            UsageLedger.idempotency_key == report_create_idempotency_key(report_id),
        )
    ).scalar_one()
    entitlements = get_entitlements(final, user_id)
    final.close()
    assert ledger == 1
    assert entitlements["entitlements"]["reports"]["used"] == 1


def test_export_quota_exceeded_does_not_complete_or_degrade(export_db):
    session = export_db()
    report_id, user_id, storage = _seed_gate3_ready_report(session)
    for index in range(2):
        session.add(
            UsageLedger(
                id=uuid.uuid4(),
                user_id=user_id,
                event_type=UsageActionType.REPORT_CREATE.value,
                occurred_at=datetime.now(timezone.utc) + timedelta(hours=index),
                idempotency_key=f"report:create:seed:{index}",
                metadata_json={},
            )
        )
    session.commit()
    session.close()

    with pytest.raises(ForbiddenError):
        export_and_persist(export_db(), report_id, storage=storage)

    final = export_db()
    report = final.get(DonorReport, report_id)
    ledger = final.execute(
        select(func.count())
        .select_from(UsageLedger)
        .where(
            UsageLedger.user_id == user_id,
            UsageLedger.event_type == UsageActionType.REPORT_CREATE.value,
            UsageLedger.idempotency_key == report_create_idempotency_key(report_id),
        )
    ).scalar_one()
    final.close()
    assert report.status == DonorReportStatus.GENERATING.value
    assert report.status != DonorReportStatus.DEGRADED.value
    assert report.status != DonorReportStatus.COMPLETE.value
    assert ledger == 0
    assert "export" not in (report.content_json or {})


def test_export_failure_sets_degraded(export_db):
    session = export_db()
    report_id, _user_id, storage = _seed_gate3_ready_report(session)
    session.close()

    class FailingStorage(MemoryDocumentStorage):
        def upload_bytes(self, key: str, data: bytes, content_type: str) -> None:
            raise RuntimeError("upload failed")

    with pytest.raises(ReportExportServiceError, match="upload failed"):
        export_and_persist(export_db(), report_id, storage=FailingStorage())

    final = export_db()
    report = final.get(DonorReport, report_id)
    final.close()
    assert report.status == DonorReportStatus.DEGRADED.value
    assert "export" not in (report.content_json or {})


@pytest.fixture
def export_api(monkeypatch):
    session_factory = create_worker_validation_sessionmaker()
    session = session_factory()
    report_id, user_id, storage = _seed_gate3_ready_report(session)
    export_and_persist(session, report_id, storage=storage)
    session.commit()
    session.close()

    monkeypatch.setattr(
        "app.reports.services.report_export_service.DocumentStorageService",
        lambda settings=None: storage,
    )

    app = create_app(
        SimpleNamespace(
            CORS_ALLOWED_ORIGINS="http://localhost:3000",
            ME_MODULE_ENABLED=True,
        )
    )

    def _override_get_db():
        db = session_factory()
        try:
            yield db
        finally:
            db.close()

    app.dependency_overrides[get_db] = _override_get_db
    token, _ = create_access_token(str(user_id), "export@example.org", "free")
    client = TestClient(app)

    yield SimpleNamespace(
        client=client,
        auth_header={"Authorization": f"Bearer {token}"},
        report_id=report_id,
        user_id=user_id,
        storage=storage,
        session_factory=session_factory,
    )


def test_download_route_returns_docx_for_owner(export_api):
    response = export_api.client.get(
        f"/api/reports/{export_api.report_id}/export",
        headers=export_api.auth_header,
    )
    assert response.status_code == 200
    assert (
        response.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert response.content.startswith(b"PK")
    assert "920,420" in _docx_plaintext(response.content)


def test_download_route_404_when_export_missing(export_api):
    session = export_api.session_factory()
    report = session.get(DonorReport, export_api.report_id)
    content = dict(report.content_json)
    content.pop("export", None)
    report.content_json = content
    session.add(report)
    session.commit()
    session.close()

    response = export_api.client.get(
        f"/api/reports/{export_api.report_id}/export",
        headers=export_api.auth_header,
    )
    assert response.status_code == 404
