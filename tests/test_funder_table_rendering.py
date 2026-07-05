"""Package 2 proof: funder tables render from REAL committed/recorded KB shapes.

Gated on the Smoke P0 M&E allowlist. Proves, against the real NLCF c1 re-walk
knowledge bank and the recorded FCDO knowledge bank (NOT the favourable distilled
fixture):

- Declared tables render as real Word tables populated from verified facts.
- No declared table renders as a bare heading.
- The facet->column mapping fails closed (unmappable real fact -> "not provided").
- Variance derives only from two real operands (lone operand -> "not provided").
- Manual / narrative tables render their columns with honest empties, never vanish.
- D2: the false "schema did not include a table field" caveat is suppressed and a
  true engine caveat names tables that genuinely could not fill.
- Package 1 unregressed: no internal identifiers in any rendered cell or heading.
"""

from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from app.reports.eval.docx_export_assertions import scan_identifier_leaks
from app.reports.export.docx_renderer import render_donor_report_docx
from app.reports.export.kb_table_renderer import (
    NOT_PROVIDED,
    table_headers_for_definition,
    table_rows_for_definition,
)

_REPO = Path(__file__).resolve().parents[1]
_NLCF_SNAPSHOT = _REPO / "docs/artefacts/me_module/audits/snapshots/c1_nlcf_rewalk_d8e7518b.json"
# Package D anchors its provenance-leak proof to the post-A walk where the leak is
# REAL: the d8e7518b snapshot above has clean semantic_labels (Table2! lives only in
# provenance.cell_ref there), so it would never exercise PL-a. The 703f0dcf walk
# carries the leaked "... — budget (Table2!C12)" labels in the KB facts. The existing
# d8e7518b proofs are NOT weakened; a real-data leak proof is added alongside.
_NLCF_703_SNAPSHOT = _REPO / "docs/artefacts/me_module/audits/snapshots/pkg2_nlcf_rewalk_703f0dcf.json"
_NLCF_TEMPLATE = _REPO / "docs/artefacts/me_module/TEMPLATE_INSTANCE_NLCF.json"
_FCDO_KB = _REPO / "tests/fixtures/reconciler/recorded/fcdo_bridgelight_recorded_knowledge_bank.json"
_FCDO_TEMPLATE = _REPO / "docs/artefacts/me_module/TEMPLATE_INSTANCE_FCDO.json"


def _load(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def _nlcf_facts() -> dict:
    snap = _load(_NLCF_SNAPSHOT)
    return snap["snapshots"]["after_synthesis"]["report"]["knowledge_bank_json"]["facts"]


def _fcdo_facts() -> dict:
    return _load(_FCDO_KB)["facts"]


def _nlcf_703_facts() -> dict:
    snap = _load(_NLCF_703_SNAPSHOT)
    return snap["snapshots"]["after_reconcile"]["report"]["knowledge_bank_json"]["facts"]


def _all_tables(template: dict) -> list[dict]:
    return [
        t
        for s in template["report_sections_json"]
        for t in (s.get("required_tables") or [])
        if isinstance(t, dict)
    ]


def _find_table(template: dict, table_key: str) -> dict:
    return next(t for t in _all_tables(template) if t.get("table_key") == table_key)


def _render(template: dict, facts: dict, content_json: dict | None = None) -> bytes:
    docx_bytes, _ = render_donor_report_docx(
        content_json=content_json or {"sections": []},
        template_sections=template["report_sections_json"],
        format_rules_json=template.get("format_rules_json") or {},
        terminology_map_json=template.get("terminology_map_json") or {},
        docx_template_ref=None,
        reporting_period_start="2024-01-01",
        reporting_period_end="2024-12-31",
        funder_name="Funder",
        template_name="Template",
        ngo_name="Test NGO",
        knowledge_bank_json={"facts": facts},
    )
    return docx_bytes


def _docx_tables(docx_bytes: bytes) -> list[list[list[str]]]:
    doc = Document(BytesIO(docx_bytes))
    out = []
    for table in doc.tables:
        out.append([[c.text for c in row.cells] for row in table.rows])
    return out


def _docx_plaintext(docx_bytes: bytes) -> str:
    doc = Document(BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


# --------------------------------------------------------------------------- #
# NLCF financials (budget_vs_actual) — real financials.lines.* shapes
# --------------------------------------------------------------------------- #


def test_nlcf_budget_vs_actual_populates_from_real_financials():
    facts = _nlcf_facts()
    table = _find_table(_load(_NLCF_TEMPLATE), "budget_vs_actual")
    rows = table_rows_for_definition(table_def=table, facts=facts)

    flat = [cell for row in rows for cell in row]
    assert "31200" in flat and "29950" in flat  # real budget + actual
    assert "-1250" in flat  # variance computed from two real operands
    # Honest empties exist (lines with a single operand).
    assert NOT_PROVIDED in flat
    # No raw fact-key namespace leaked into any cell.
    assert not any("financials.lines" in cell for cell in flat)


def test_nlcf_outcomes_summary_manual_renders_honest_empty_not_vanished():
    facts = _nlcf_facts()
    table = _find_table(_load(_NLCF_TEMPLATE), "outcomes_summary")
    rows = table_rows_for_definition(table_def=table, facts=facts)
    assert rows, "manual table must render its columns, not vanish"
    assert all(cell == NOT_PROVIDED for row in rows for cell in row)


# --------------------------------------------------------------------------- #
# FCDO indicators (output_score_table) — real indicators.* shapes
# --------------------------------------------------------------------------- #


def test_fcdo_output_score_table_populates_from_real_indicators():
    facts = _fcdo_facts()
    table = _find_table(_load(_FCDO_TEMPLATE), "output_score_table")
    rows = table_rows_for_definition(table_def=table, facts=facts)
    flat = [cell for row in rows for cell in row]

    # Real recorded actuals/targets present (op1_1: actual 985, target 1200).
    assert "985" in flat and "1200" in flat
    # Not the deleted 12-row OP skeleton: many real indicator rows, not placeholders.
    assert len(rows) > 12
    assert not any(cell.strip() in {"OP2.3", "OP4.2"} for cell in flat)
    # Not all honest-empty.
    assert any(c not in ("", NOT_PROVIDED) for c in flat)


@pytest.mark.parametrize("table_key", ["outcome_assessment", "vfm_measures"])
def test_fcdo_narrative_indicator_tables_render_honest_empty(table_key):
    facts = _fcdo_facts()
    table = _find_table(_load(_FCDO_TEMPLATE), table_key)
    rows = table_rows_for_definition(table_def=table, facts=facts)
    flat = [cell for row in rows for cell in row]
    # Indicator actuals must NOT be dumped into a narrative table (wrong table).
    assert "985" not in flat
    assert all(cell == NOT_PROVIDED for cell in flat)


# --------------------------------------------------------------------------- #
# Guards (owner-required)
# --------------------------------------------------------------------------- #


def test_facet_mapping_fails_closed_real_fact_not_forced_into_column():
    table = {
        "table_key": "t",
        "data_source": "financials",
        "columns": [
            {"column_key": "cost_type", "label": "Cost type"},
            {"column_key": "budgeted_amount", "label": "Budget"},
            {"column_key": "actual_spend", "label": "Actual"},
        ],
    }
    facts = {
        "financials.lines.unmapped.budgeted_amount": {"value": 100, "semantic_label": "Unmapped line"},
        "financials.lines.unmapped.forecast_q3": {"value": 4242, "semantic_label": "Unmapped line"},
    }
    row = table_rows_for_definition(table_def=table, facts=facts)[0]
    assert "4242" not in row  # unmappable real fact is never placed in a column
    assert row == ["Unmapped line", "100", NOT_PROVIDED]


def test_variance_guard_lone_operand_never_renders_figure():
    table = {
        "table_key": "t",
        "data_source": "financials",
        "columns": [
            {"column_key": "cost_type", "label": "Cost type"},
            {"column_key": "budgeted_amount", "label": "Budget"},
            {"column_key": "actual_spend", "label": "Actual"},
            {"column_key": "variance", "label": "Variance"},
        ],
    }
    actual_only = table_rows_for_definition(
        table_def=table, facts={"financials.lines.a.actual_spend": {"value": 500, "semantic_label": "A"}}
    )[0]
    assert actual_only[3] == NOT_PROVIDED and actual_only[2] == "500"
    budget_only = table_rows_for_definition(
        table_def=table, facts={"financials.lines.b.budget": {"value": 700, "semantic_label": "B"}}
    )[0]
    assert budget_only[3] == NOT_PROVIDED and budget_only[1] == "700"


# --------------------------------------------------------------------------- #
# Full-render proofs (no bare headings, Package 1 clean, D2 caveat)
# --------------------------------------------------------------------------- #


@pytest.mark.parametrize("template_path,facts_fn", [
    (_NLCF_TEMPLATE, _nlcf_facts),
    (_FCDO_TEMPLATE, _fcdo_facts),
])
def test_no_declared_table_renders_as_bare_heading(template_path, facts_fn):
    template = _load(template_path)
    docx_bytes = _render(template, facts_fn())
    rendered = _docx_tables(docx_bytes)
    declared = [t for t in _all_tables(template) if table_headers_for_definition(t)]
    assert len(rendered) >= len(declared)


@pytest.mark.parametrize("template_path,facts_fn", [
    (_NLCF_TEMPLATE, _nlcf_facts),
    (_FCDO_TEMPLATE, _fcdo_facts),
])
def test_rendered_tables_have_no_identifier_leaks(template_path, facts_fn):
    docx_bytes = _render(_load(template_path), facts_fn())
    assert scan_identifier_leaks(_docx_plaintext(docx_bytes)) == []


# --------------------------------------------------------------------------- #
# Package D — real provenance-leak walk (703f0dcf): clean cells + tripwire
# --------------------------------------------------------------------------- #


def test_pkgd_real_leaked_labels_render_clean_identity_cells():
    facts = _nlcf_703_facts()
    # Precondition: the real walk carries the provenance leak in semantic_label.
    assert any(
        "Table2!" in (f.get("semantic_label") or "") for f in facts.values()
    ), "precondition: 703f0dcf must carry the semantic_label cell-ref leak"

    table = _find_table(_load(_NLCF_TEMPLATE), "budget_vs_actual")
    rows = table_rows_for_definition(table_def=table, facts=facts)
    flat = [cell for row in rows for cell in row]

    # The clean human name reaches the identity cell; provenance does not.
    assert "Sessional youth workers" in flat
    assert not any("Table2!" in cell for cell in flat)
    assert not any(("— budget" in cell or "— actual" in cell) for cell in flat)
    # Real figures still populate from the same facts.
    assert "13600" in flat
    # Whole rendered table is leak-free under the widened tripwire.
    assert scan_identifier_leaks("\n".join(flat)) == []


def test_pkgd_tripwire_fires_on_real_leak_and_silent_after_render():
    facts = _nlcf_703_facts()
    leaked = next(
        f["semantic_label"]
        for f in facts.values()
        if "Table2!" in (f.get("semantic_label") or "")
    )
    assert scan_identifier_leaks(leaked), f"tripwire missed real leak: {leaked}"

    docx_bytes = _render(_load(_NLCF_TEMPLATE), facts)
    assert scan_identifier_leaks(_docx_plaintext(docx_bytes)) == []


def test_d2_false_schema_caveat_suppressed_true_caveat_present():
    template = _load(_FCDO_TEMPLATE)
    first_section_key = template["report_sections_json"][0]["section_key"]
    false_caveat = (
        "A required outcomes table was not populated because the output schema did "
        "not include a table field and so the row-level detail could not be shown."
    )
    content_json = {
        "sections": [
            {
                "section_key": first_section_key,
                "generation_status": "GENERATED",
                "content": {"text": "Narrative body.", "assumptions": [false_caveat]},
            }
        ]
    }
    docx_bytes = _render(template, _fcdo_facts(), content_json=content_json)
    text = _docx_plaintext(docx_bytes)
    assert "schema did not include a table field" not in text
    # A genuinely honest-empty declared table yields the TRUE engine caveat.
    assert "no verified figures were available in the submitted records" in text
