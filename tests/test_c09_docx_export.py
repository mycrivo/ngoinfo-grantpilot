"""
C-09 DOCX Export — End-to-end verification tests.

Tests:
  T1. DOCX binary integrity (ZIP magic bytes PK\\x03\\x04) + file size
  T2. DOCX content (opportunity title, NGO name, section text present)
  T3. Format validation (non-DOCX → 422 UNSUPPORTED_FORMAT)
  T4. Ownership enforcement (user A cannot export user B's proposal → 403)
  T5. 404 for nonexistent proposal
  T6. Idempotency: second export same version → no duplicate ledger row
  T7. Idempotency key format: docx_export:{user}:{proposal}:v{version}
  T8. Version bump → new ledger row (different idem key)
  T9. Route exists at POST /api/proposals/{proposal_id}/export
"""

import uuid
from datetime import datetime, timezone
from io import BytesIO
from types import SimpleNamespace

from sqlalchemy import create_engine, event, text as sa_text
from sqlalchemy.dialects.postgresql import UUID, JSONB
from sqlalchemy.ext.compiler import compiles
from sqlalchemy.orm import sessionmaker

# Compile PostgreSQL types for SQLite
@compiles(UUID, "sqlite")
def _compile_uuid_sqlite(_type, _compiler, **_kw):
    return "CHAR(36)"


@compiles(JSONB, "sqlite")
def _compile_jsonb_sqlite(_type, _compiler, **_kw):
    return "TEXT"


from app.db.base import Base
from app.models.funding_opportunity import FundingOpportunity
from app.models.ngo_profile import NGOProfile
from app.models.proposal import Proposal
from app.models.usage_ledger import UsageLedger, UsageActionType
from app.models.user import User
from app.models.user_plan import UserPlan
# Ensure all models are imported so Base.metadata knows about them
import app.models.fit_scan  # noqa: F401
import app.models.auth_oauth_exchange_code  # noqa: F401
import app.models.auth_refresh_token  # noqa: F401
import app.models.stripe_event  # noqa: F401


def _make_db():
    engine = create_engine("sqlite:///:memory:")

    # Temporarily remove ALL server_defaults and CHECK constraints to avoid Postgres syntax
    _saved_defaults = {}
    _saved_constraints = {}
    for table in Base.metadata.tables.values():
        to_remove = [c for c in table.constraints if hasattr(c, "sqltext")]
        for c in to_remove:
            _saved_constraints.setdefault(table.name, []).append(c)
            table.constraints.discard(c)

        for col in table.columns:
            if col.server_default is not None:
                _saved_defaults[(table.name, col.name)] = col.server_default
                col.server_default = None

    @event.listens_for(engine, "connect")
    def _set_sqlite_pragma(dbapi_conn, _record):
        dbapi_conn.execute("PRAGMA foreign_keys = OFF")

    Base.metadata.create_all(engine)

    # Restore everything
    for (tname, cname), default in _saved_defaults.items():
        Base.metadata.tables[tname].columns[cname].server_default = default
    for tname, constraints in _saved_constraints.items():
        for c in constraints:
            Base.metadata.tables[tname].constraints.add(c)

    return sessionmaker(bind=engine)()


def _seed(db):
    """Seed via ORM objects (avoids SQLite/UUID mismatch issues)."""
    now = datetime.now(timezone.utc)
    user = User(
        id=uuid.uuid4(), email="export-test@example.org", auth_provider="email",
        created_at=now, updated_at=now,
    )
    db.add(user)
    db.flush()

    plan = UserPlan(id=uuid.uuid4(), user_id=user.id, plan_name="GROWTH",
                    plan_activated_at=now, created_at=now, updated_at=now)
    db.add(plan)
    db.flush()

    opp = FundingOpportunity(
        id=uuid.uuid4(), title="Test Grant 2026",
        source_url="https://example.org/grant",
        application_url="https://example.org/apply",
        donor_organization="Test Foundation",
        funding_type="Project Grant",
        applicant_type="NGO",
        location_text="Global",
        focus_areas="Education, Health",
        deadline_type="ROLLING",
        short_summary="A test grant.",
        requirements_json={"variants": [{"submission_items": []}]},
        is_active=True, is_archived=False, status="READY",
        created_at=now, updated_at=now,
    )
    db.add(opp)
    db.flush()

    profile = NGOProfile(
        id=uuid.uuid4(), user_id=user.id,
        organization_name="Test NGO",
        country_of_registration="US",
        mission_statement="Help the world.",
        focus_sectors=["Education"],
        geographic_areas_of_work=["Global"],
        target_groups=["Youth"],
        past_projects=[],
        profile_status="COMPLETE",
        completeness_score=100,
        missing_fields=[],
        funders_worked_with_before=[],
        created_at=now, updated_at=now,
    )
    db.add(profile)
    db.flush()

    content = {
        "sections": [
            {
                "submission_item_id": "s1",
                "label": "Executive Summary",
                "generation_status": "GENERATED",
                "content": {
                    "text": "This is a test proposal section.",
                    "assumptions": ["Assumption A", "Assumption B"],
                },
            },
            {
                "submission_item_id": "s2",
                "label": "Budget Overview",
                "generation_status": "MANUAL_REQUIRED",
                "content": {"text": "", "assumptions": []},
            },
        ],
        "generation_summary": {"total_items": 2, "generated": 1},
    }

    proposal = Proposal(
        id=uuid.uuid4(), user_id=user.id, funding_opportunity_id=opp.id,
        plan_at_creation="GROWTH", prompt_version="v1.0.0",
        version=1, status="READY", regeneration_count=0,
        content_json=content, created_at=now, updated_at=now,
    )
    db.add(proposal)
    db.commit()
    db.refresh(user)
    db.refresh(proposal)
    return user, proposal


def _mock_settings(monkeypatch):
    fake = SimpleNamespace(AUTH_JWT_SIGNING_KEY="x" * 64)
    import app.core.security as security
    monkeypatch.setattr(security, "get_settings", lambda: fake)


def _patch_quota_commits(monkeypatch):
    """Make record_usage and get_or_create_user_plan use flush instead of commit for sqlite.
    Also inject Python-side uuid4 default for UsageLedger.id (Postgres gen_random_uuid
    was stripped for SQLite compatibility)."""
    from app.services import quota_service as qs
    from app.models.usage_ledger import UsageLedger as UL

    _orig_record = qs.record_usage
    def patched_record(db, user_id, event_type, *, idempotency_key=None, commit=True):
        result = _orig_record(db, user_id, event_type, idempotency_key=idempotency_key, commit=False)
        # For new rows, provide Python-side defaults since Postgres functions were stripped
        if result.id is None:
            result.id = uuid.uuid4()
        if result.metadata_json is None:
            result.metadata_json = {}
        return result
    monkeypatch.setattr(qs, "record_usage", patched_record)

    _orig_plan = qs.get_or_create_user_plan
    def patched_plan(db, user_id, *, commit=True):
        return _orig_plan(db, user_id, commit=False)
    monkeypatch.setattr(qs, "get_or_create_user_plan", patched_plan)


# ===== TESTS =====

def test_T1_docx_binary_integrity(monkeypatch):
    """T1: export_docx returns bytes starting with PK\\x03\\x04 (ZIP/DOCX magic)."""
    _mock_settings(monkeypatch)
    _patch_quota_commits(monkeypatch)
    db = _make_db()
    user, proposal = _seed(db)

    from app.services.export_service import ExportService
    service = ExportService(db)
    content, filename = service.export_docx(
        user=user, proposal_id=proposal.id, export_format="DOCX",
    )

    assert isinstance(content, bytes)
    assert len(content) > 100, f"File too small: {len(content)} bytes"
    assert content[:4] == b"PK\x03\x04", f"Not a valid ZIP: first 4 bytes = {content[:4].hex()}"
    assert filename == f"proposal-{proposal.id}.docx"
    print(f"\n  PASS | size={len(content)} bytes | first20hex={content[:20].hex()} | file={filename}")


def test_T2_docx_content(monkeypatch):
    """T2: DOCX contains opportunity title, NGO name, section text."""
    _mock_settings(monkeypatch)
    _patch_quota_commits(monkeypatch)
    db = _make_db()
    user, proposal = _seed(db)

    from app.services.export_service import ExportService
    from docx import Document
    service = ExportService(db)
    content, _ = service.export_docx(user=user, proposal_id=proposal.id, export_format="DOCX")

    doc = Document(BytesIO(content))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    assert "Test Grant 2026" in full_text, "Opportunity title missing"
    assert "Test NGO" in full_text, "NGO name missing"
    assert str(proposal.id) in full_text, "Proposal ID missing"
    assert "This is a test proposal section." in full_text, "Section text missing"
    print(f"\n  PASS | text length={len(full_text)} chars | headings checked")


def test_T3_rejects_non_docx(monkeypatch):
    """T3: non-DOCX format → 422 UNSUPPORTED_FORMAT."""
    _mock_settings(monkeypatch)
    _patch_quota_commits(monkeypatch)
    db = _make_db()
    user, proposal = _seed(db)

    from app.services.export_service import ExportService
    from app.core.errors import DomainError
    service = ExportService(db)
    try:
        service.export_docx(user=user, proposal_id=proposal.id, export_format="PDF")
        assert False, "Should have raised"
    except DomainError as exc:
        assert exc.error_code == "UNSUPPORTED_FORMAT"
        assert exc.status_code == 422
        print(f"\n  PASS | PDF rejected: {exc.error_code} / {exc.status_code}")


def test_T4_ownership_enforced(monkeypatch):
    """T4: User B cannot export User A's proposal → 403 FORBIDDEN."""
    _mock_settings(monkeypatch)
    _patch_quota_commits(monkeypatch)
    db = _make_db()
    user_a, proposal = _seed(db)

    # Create user B via ORM
    user_b = User(
        id=uuid.uuid4(), email="other@example.org", auth_provider="email",
        created_at=datetime.now(timezone.utc), updated_at=datetime.now(timezone.utc),
    )
    db.add(user_b)
    db.commit()
    db.refresh(user_b)

    from app.services.export_service import ExportService
    from app.core.errors import DomainError
    service = ExportService(db)
    try:
        service.export_docx(user=user_b, proposal_id=proposal.id, export_format="DOCX")
        assert False, "Should have raised"
    except DomainError as exc:
        assert exc.error_code == "FORBIDDEN"
        assert exc.status_code == 403
        print(f"\n  PASS | ownership enforced: {exc.error_code}")


def test_T5_nonexistent_proposal(monkeypatch):
    """T5: Exporting nonexistent proposal → 404 PROPOSAL_NOT_FOUND."""
    _mock_settings(monkeypatch)
    _patch_quota_commits(monkeypatch)
    db = _make_db()
    user, _ = _seed(db)

    from app.services.export_service import ExportService
    from app.core.errors import DomainError
    service = ExportService(db)
    try:
        service.export_docx(user=user, proposal_id=uuid.uuid4(), export_format="DOCX")
        assert False, "Should have raised"
    except DomainError as exc:
        assert exc.error_code == "PROPOSAL_NOT_FOUND"
        assert exc.status_code == 404
        print(f"\n  PASS | 404 for missing: {exc.error_code}")


def test_T6_idempotency(monkeypatch):
    """T6: Second export same proposal+version → no duplicate ledger row."""
    _mock_settings(monkeypatch)
    _patch_quota_commits(monkeypatch)
    db = _make_db()
    user, proposal = _seed(db)

    from app.services.export_service import ExportService
    service = ExportService(db)

    # Export #1
    service.export_docx(user=user, proposal_id=proposal.id, export_format="DOCX")
    db.flush()
    count_1 = db.query(UsageLedger).filter(
        UsageLedger.user_id == user.id,
        UsageLedger.event_type == UsageActionType.DOCX_EXPORT.value,
    ).count()

    # Export #2 — same proposal, same version
    service.export_docx(user=user, proposal_id=proposal.id, export_format="DOCX")
    db.flush()
    count_2 = db.query(UsageLedger).filter(
        UsageLedger.user_id == user.id,
        UsageLedger.event_type == UsageActionType.DOCX_EXPORT.value,
    ).count()

    assert count_1 == 1, f"Export #1: expected 1 row, got {count_1}"
    assert count_2 == 1, f"Export #2: expected still 1 row, got {count_2}"
    print(f"\n  PASS | rows after #1: {count_1} | after #2: {count_2} | idempotent")


def test_T7_idempotency_key_format(monkeypatch):
    """T7: Key format = docx_export:{user_id}:{proposal_id}:v{version}."""
    _mock_settings(monkeypatch)
    _patch_quota_commits(monkeypatch)
    db = _make_db()
    user, proposal = _seed(db)

    from app.services.export_service import ExportService
    service = ExportService(db)
    service.export_docx(user=user, proposal_id=proposal.id, export_format="DOCX")
    db.flush()

    row = db.query(UsageLedger).filter(
        UsageLedger.user_id == user.id,
        UsageLedger.event_type == UsageActionType.DOCX_EXPORT.value,
    ).first()

    expected = f"docx_export:{user.id}:{proposal.id}:v{int(proposal.version)}"
    assert row is not None, "No DOCX_EXPORT row"
    assert row.idempotency_key == expected, f"Expected '{expected}', got '{row.idempotency_key}'"
    print(f"\n  PASS | idem_key = {row.idempotency_key}")


def test_T8_new_version_new_row(monkeypatch):
    """T8: Version bump → new ledger row."""
    _mock_settings(monkeypatch)
    _patch_quota_commits(monkeypatch)
    db = _make_db()
    user, proposal = _seed(db)

    from app.services.export_service import ExportService
    service = ExportService(db)

    # Export v1
    service.export_docx(user=user, proposal_id=proposal.id, export_format="DOCX")
    db.flush()

    # Bump version
    proposal.version = 2
    db.flush()

    # Export v2
    service.export_docx(user=user, proposal_id=proposal.id, export_format="DOCX")
    db.flush()

    rows = db.query(UsageLedger).filter(
        UsageLedger.user_id == user.id,
        UsageLedger.event_type == UsageActionType.DOCX_EXPORT.value,
    ).all()

    assert len(rows) == 2, f"Expected 2 rows, got {len(rows)}"
    keys = {r.idempotency_key for r in rows}
    assert f"docx_export:{user.id}:{proposal.id}:v1" in keys
    assert f"docx_export:{user.id}:{proposal.id}:v2" in keys
    print(f"\n  PASS | v1+v2 rows: {len(rows)} | keys: {keys}")


def test_T9_route_exists():
    """T9: POST /api/proposals/{proposal_id}/export route is registered."""
    from app.api.routes.proposals import router

    export_routes = [r for r in router.routes if hasattr(r, "path") and "/export" in r.path]
    assert len(export_routes) == 1
    route = export_routes[0]
    assert "POST" in route.methods
    assert route.path == "/api/proposals/{proposal_id}/export"
    print(f"\n  PASS | POST {route.path}")
