"""A-JSON — synthesis JSON-parse resilience (Smoke P0 M&E allowlist).

Anchored to the REAL P3-8 failure class: a large synthesis JSON object truncated
mid-string ("Unterminated string starting at ... char 8092"). Proves:

  1. A truncated-mid-string payload is NEVER bound as a complete section — it surfaces
     as the honest synthesis_parse_failure terminal state (completeness gate).
  2. The parse ladder recovers a genuinely COMPLETE object (strict / fenced).
  3. One bounded identical retry recovers a transient failure.
  4. synthesis_parse_failure is distinct from insufficient_data (state + prose).
  5. The resume retry is bounded — it settles into the terminal state.
  6. Gate-3 treats synthesis_parse_failure as export-ready (report completes).
  7. Raw payload is trace-only: captured in the diagnostics sink, NEVER on the section,
     the DOCX, or the caveats.
"""

from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path

from docx import Document

from app.reports.eval.docx_export_assertions import scan_identifier_leaks
from app.reports.export.docx_renderer import render_donor_report_docx
from app.reports.parsing.json_from_text import (
    extract_complete_json_object,
    extract_json_object_from_text,
)
from app.reports.schemas.content_json_v1 import (
    MAX_SYNTHESIS_PARSE_FAILURE_CYCLES,
    section_needs_synthesis,
)
from app.reports.services.gate3_confirmation_service import _sections_not_export_ready
from app.reports.services.report_synthesis_service import _generate_one_section
from app.reports.services.section_prose import (
    STRUCTURED_BIND_STATUS_INSUFFICIENT_DATA,
    STRUCTURED_BIND_STATUS_SYNTHESIS_PARSE_FAILURE,
    build_insufficient_data_section,
    build_parse_failure_statement,
    build_synthesis_parse_failure_section,
)
from app.reports.services.synthesis_parse import parse_synthesis_response

ROOT = Path(__file__).resolve().parents[1]
FCDO_TEMPLATE = ROOT / "docs" / "artefacts" / "me_module" / "TEMPLATE_INSTANCE_FCDO.json"

RAW_SENTINEL = "RAW_PAYLOAD_SENTINEL_do_not_leak_9f2c"


# ---------------------------------------------------------------------------
# Fixtures: real P3-8 truncation class + genuinely-complete counterparts
# ---------------------------------------------------------------------------
def _section_from_template(section_key: str) -> dict:
    payload = json.loads(FCDO_TEMPLATE.read_text(encoding="utf-8"))
    for section in payload["report_sections_json"]:
        if section.get("section_key") == section_key:
            return section
    raise KeyError(section_key)


def _truncated_midstring_content(sentinel: str = RAW_SENTINEL) -> str:
    """A large synthesis object with COMPLETE inner claim objects, then cut off mid-string.

    Mirrors the real char-8092 unterminated-string failure: everything before the cut is
    well-formed (including whole inner claim objects a naive salvager would grab), but the
    outer object never closes.
    """
    head = (
        '{"section_key": "evidence_and_evaluation", '
        '"generation_status": "GENERATED", '
        '"archetype": "ARCH_EVIDENCE_AND_EVALUATION_REVIEW", '
        '"generated_content": {"claims": ['
    )
    claims = ", ".join(
        '{"text": "Delivery evidence point %d recorded against confirmed indicator '
        'records for the reporting period.", "source_refs": '
        '["fact:indicators.OP1.%d.actual"], "value_tokens": []}' % (i, i)
        for i in range(30)
    )
    # Unterminated final string carrying the sentinel — no closing quote or braces.
    tail = (
        '], "text": "The programme delivered against its indicators. ' + sentinel
        + " and then the completion was cut off mid-sentence before the string could"
    )
    return head + claims + tail


def _complete_synthesis_object(section_key: str) -> dict:
    return {
        "section_key": section_key,
        "generation_status": "GENERATED",
        "archetype": None,
        "generated_content": {
            "claims": [
                {
                    "text": "Programme delivery continued against confirmed indicator records.",
                    "source_refs": ["fact:indicators.OP1.1.ar1_actual"],
                    "value_tokens": [],
                }
            ],
            "text": "Programme delivery continued against confirmed indicator records.",
            "assumptions": [],
        },
        "constraints_applied": {"word_limit": 900, "word_limit_respected": True},
        "warnings": [],
    }


def _raw_response(content: str, *, finish_reason: str = "length") -> dict:
    return {
        "choices": [{"finish_reason": finish_reason, "message": {"content": content}}],
        "usage": {"prompt_tokens": 1200, "completion_tokens": 2250},
    }


def _fcdo_kb() -> dict:
    from tests.test_gap_compliance_agent import _load_distilled_fcdo_kb

    return _load_distilled_fcdo_kb()


def _run_section(*, raw_contents: list[str], prior_cycles: int = 0):
    """Drive _generate_one_section with a scripted sequence of raw responses."""
    section = _section_from_template("performance_and_conclusions")
    kb = _fcdo_kb()
    calls = {"n": 0}
    diagnostics: list[dict] = []

    def _raw_fn(_sk, _sys, _user):
        idx = min(calls["n"], len(raw_contents) - 1)
        calls["n"] += 1
        finish = "stop" if raw_contents[idx].rstrip().endswith("}") else "length"
        return _raw_response(raw_contents[idx], finish_reason=finish), 10, 20

    result, in_tok, out_tok = _generate_one_section(
        section=section,
        report_inputs={"knowledge_bank": kb},
        knowledge_bank_json=kb,
        report_context={"report_type": "annual"},
        query_fn_synthesis=None,
        user_id=None,
        raw_response_fn=_raw_fn,
        prior_parse_failure_cycles=prior_cycles,
        diagnostics_out=diagnostics,
    )
    return result, diagnostics, calls["n"]


def _docx_text(docx_bytes: bytes) -> str:
    document = Document(BytesIO(docx_bytes))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


# ===========================================================================
# 1. Completeness gate — truncated NEVER binds (adjustment #1, moat-critical)
# ===========================================================================
def test_truncated_midstring_is_not_parseable_as_complete():
    attempt = parse_synthesis_response(_raw_response(_truncated_midstring_content()))
    assert attempt.ok is False
    assert attempt.payload is None
    assert attempt.parse_strategy == "none"
    assert attempt.finish_reason == "length"
    assert "Unterminated string" in (attempt.parse_error or "")


def test_completeness_gate_refuses_the_salvageable_inner_fragment():
    """The old extractor WOULD salvage an inner claim object; the completeness gate must not."""
    truncated = _truncated_midstring_content()
    salvaged = extract_json_object_from_text(truncated)
    assert isinstance(salvaged, dict)  # naive extractor grabs a fragment...
    assert salvaged.get("generation_status") is None  # ...which is NOT the section envelope
    # The completeness-preserving path refuses it outright.
    assert extract_complete_json_object(truncated) is None


def test_truncated_midstring_section_surfaces_parse_failure_not_bound():
    truncated = _truncated_midstring_content()
    result, _diags, calls = _run_section(raw_contents=[truncated])
    # One bounded identical retry => two attempts before settling.
    assert calls == 2
    assert result["generation_status"] == "GENERATED"
    assert (
        result["content"]["structured_bind_status"]
        == STRUCTURED_BIND_STATUS_SYNTHESIS_PARSE_FAILURE
    )
    # Never silently converted to a data gap, and never an empty/complete-looking section.
    assert result["content"]["structured_bind_status"] != STRUCTURED_BIND_STATUS_INSUFFICIENT_DATA
    assert len(result["content"]["text"].strip()) >= 40
    assert result["content"]["claims"] == []


# ===========================================================================
# 2. Ladder recovers a genuinely COMPLETE object
# ===========================================================================
def test_strict_parse_of_complete_object():
    content = json.dumps(_complete_synthesis_object("performance_and_conclusions"))
    attempt = parse_synthesis_response(_raw_response(content, finish_reason="stop"))
    assert attempt.ok is True
    assert attempt.parse_strategy == "json.loads"
    assert attempt.payload["generation_status"] == "GENERATED"


def test_fenced_complete_object_recovered_by_balanced_ladder():
    inner = json.dumps(_complete_synthesis_object("performance_and_conclusions"))
    fenced = "```json\n" + inner + "\n```"
    attempt = parse_synthesis_response(_raw_response(fenced, finish_reason="stop"))
    assert attempt.ok is True
    assert attempt.parse_strategy == "balanced_object"
    assert attempt.payload["generation_status"] == "GENERATED"


def test_preamble_prose_then_complete_object_recovered():
    inner = json.dumps(_complete_synthesis_object("performance_and_conclusions"))
    content = "Here is the section:\n" + inner
    attempt = parse_synthesis_response(_raw_response(content, finish_reason="stop"))
    assert attempt.ok is True
    assert attempt.parse_strategy == "balanced_object"
    assert attempt.payload["generation_status"] == "GENERATED"


def test_complete_object_with_trailing_garbage_refused():
    # Conservative completeness gate: content after the closed object is ambiguous
    # (possible dropped content), so it is refused rather than silently salvaged.
    content = json.dumps(_complete_synthesis_object("x")) + " <<trailing noise>>"
    assert extract_complete_json_object(content) is None


# ===========================================================================
# 3. Bounded retry recovers a transient failure
# ===========================================================================
def test_one_bounded_retry_recovers_on_second_attempt():
    truncated = _truncated_midstring_content()
    complete = json.dumps(_complete_synthesis_object("performance_and_conclusions"))
    result, diagnostics, calls = _run_section(raw_contents=[truncated, complete])
    assert calls == 2
    assert result["generation_status"] == "GENERATED"
    bind_status = result["content"].get("structured_bind_status")
    assert bind_status in ("bound", "honest_empty")
    assert bind_status != STRUCTURED_BIND_STATUS_SYNTHESIS_PARSE_FAILURE
    # The first (failed) attempt is still recorded for diagnosis.
    assert len(diagnostics) == 1
    assert diagnostics[0]["attempt"] == 1


# ===========================================================================
# 4. Distinct from insufficient_data (state + prose)
# ===========================================================================
def test_parse_failure_distinct_from_insufficient_data():
    section = _section_from_template("performance_and_conclusions")
    parse_fail = build_synthesis_parse_failure_section(section=section)
    insufficient = build_insufficient_data_section(section=section)

    assert (
        parse_fail["content"]["structured_bind_status"]
        == STRUCTURED_BIND_STATUS_SYNTHESIS_PARSE_FAILURE
    )
    assert (
        insufficient["content"]["structured_bind_status"]
        == STRUCTURED_BIND_STATUS_INSUFFICIENT_DATA
    )
    pf_text = parse_fail["content"]["text"].lower()
    ins_text = insufficient["content"]["text"].lower()
    assert pf_text != ins_text
    # Parse failure blames the drafting system, NOT missing evidence.
    assert "drafting system" in pf_text
    assert "left this section blank" not in pf_text
    # Insufficiency is the opposite framing.
    assert "left this section blank" in ins_text


def test_parse_failure_prose_carries_no_fabricated_claims_or_raw_payload():
    section = _section_from_template("performance_and_conclusions")
    text = build_parse_failure_statement(section=section)
    assert scan_identifier_leaks(text) == [], text
    assert not any(ch.isdigit() for ch in text)
    assert RAW_SENTINEL not in text


# ===========================================================================
# 5. Bounded resume — settles into terminal state (adjustment #2)
# ===========================================================================
def test_resume_retry_is_bounded_and_settles():
    section = _section_from_template("performance_and_conclusions")
    below = build_synthesis_parse_failure_section(section=section, parse_failure_cycles=1)
    assert below["content"]["parse_failure_cycles"] == 1
    assert section_needs_synthesis(below) is True  # 1 < ceiling => retried on resume

    settled = build_synthesis_parse_failure_section(
        section=section, parse_failure_cycles=MAX_SYNTHESIS_PARSE_FAILURE_CYCLES
    )
    assert section_needs_synthesis(settled) is False  # ceiling reached => settles


def test_parse_failure_cycle_counter_increments_from_prior():
    truncated = _truncated_midstring_content()
    result, _diags, _calls = _run_section(raw_contents=[truncated], prior_cycles=1)
    assert result["content"]["parse_failure_cycles"] == 2


# ===========================================================================
# 6. Gate-3 export-ready — a single unreadable section does not freeze the report
# ===========================================================================
def test_gate3_treats_parse_failure_as_export_ready():
    content_json = {
        "sections": [
            {
                "section_key": "s_parse",
                "generation_status": "ACCEPTED",
                "content": {
                    "text": "This section could not be finalised (engine prose).",
                    "citation_mode": "structured",
                    "structured_bind_status": STRUCTURED_BIND_STATUS_SYNTHESIS_PARSE_FAILURE,
                },
            },
            {
                "section_key": "s_insufficient",
                "generation_status": "ACCEPTED",
                "content": {
                    "text": "No citable source supplied material for this section.",
                    "citation_mode": "structured",
                    "structured_bind_status": STRUCTURED_BIND_STATUS_INSUFFICIENT_DATA,
                },
            },
        ]
    }
    assert _sections_not_export_ready(content_json) == []


def test_gate3_still_blocks_unknown_bind_status():
    content_json = {
        "sections": [
            {
                "section_key": "s_bad",
                "generation_status": "ACCEPTED",
                "content": {
                    "text": "some prose",
                    "citation_mode": "structured",
                    "structured_bind_status": "BIND_FAILED",
                },
            }
        ]
    }
    assert _sections_not_export_ready(content_json) == ["s_bad"]


# ===========================================================================
# 7. Raw payload is trace-only (adjustment #3) — both directions
# ===========================================================================
def test_raw_payload_captured_in_trace_but_never_on_section_or_export():
    truncated = _truncated_midstring_content()
    result, diagnostics, _calls = _run_section(raw_contents=[truncated])

    # Direction A: the raw payload IS captured in the diagnostics sink (trace-only).
    assert diagnostics, "expected parse-failure diagnostics for the trace"
    trace_blob = json.dumps(diagnostics)
    assert RAW_SENTINEL in trace_blob
    assert any(d.get("finish_reason") == "length" for d in diagnostics)
    assert all("response_head" in d or "response_tail" in d for d in diagnostics)

    # Direction B: the raw payload NEVER rides on the persisted section.
    section_blob = json.dumps(result)
    assert RAW_SENTINEL not in section_blob
    assert "Unterminated string" not in section_blob

    # Direction B: nor does it reach the DOCX body or caveats/assumptions appendix.
    template_sections = [
        {
            "section_key": "performance_and_conclusions",
            "label": "Performance and conclusions",
            "required_tables": [],
        }
    ]
    render_section = dict(result)
    render_section["generation_status"] = "ACCEPTED"
    docx_bytes, _ = render_donor_report_docx(
        content_json={"sections": [render_section]},
        template_sections=template_sections,
        format_rules_json={"document_title": "Test Report"},
        terminology_map_json={},
        docx_template_ref=None,
        reporting_period_start="2024-10-15",
        reporting_period_end="2025-10-14",
        funder_name="Test Funder",
        template_name="Test Template",
    )
    out = _docx_text(docx_bytes)
    assert RAW_SENTINEL not in out
    assert "Unterminated string" not in out
    assert scan_identifier_leaks(out) == [], out
    # The honest engine prose IS shown.
    assert "drafting system" in out.lower()
