"""A-04 structural hardening acceptance tests for proposal and M&E DOCX exports."""

from __future__ import annotations

import json
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from app.core.docx_presentation import ASSUMPTIONS_APPENDIX_TITLE
from app.reports.export.docx_renderer import render_donor_report_docx
from app.services.proposal_docx_renderer import build_proposal_docx_bytes

FCDO_TEMPLATE_PATH = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "artefacts"
    / "me_module"
    / "TEMPLATE_INSTANCE_FCDO.json"
)
FCDO_CONTENT_PATH = (
    Path(__file__).resolve().parents[0]
    / "fixtures"
    / "export"
    / "fcdo_recorded_content_json.json"
)

INTERNAL_MARKERS = (
    "[Section not generated]",
    "[Not generated",
    "Proposal ID:",
    "Generated At (UTC):",
    "Version:",
)

KNOWN_ME_BODY_SNIPPET = (
    "The programme made steady progress during the Annual Review period"
)
KNOWN_PROPOSAL_BODY = (
    "BridgeLight will re-enrol out-of-school girls across three districts."
)


def _load_fcdo_template_fields() -> dict:
    return json.loads(FCDO_TEMPLATE_PATH.read_text(encoding="utf-8"))


def _docx_all_text(docx_bytes: bytes) -> str:
    document = Document(BytesIO(docx_bytes))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    for section in document.sections:
        footer = section.footer
        for paragraph in footer.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)
    return "\n".join(parts)


def _heading_paragraphs(docx_bytes: bytes) -> list:
    document = Document(BytesIO(docx_bytes))
    return [
        p
        for p in document.paragraphs
        if p.style.name.startswith("Heading")
    ]


def _footer_has_page_fields(docx_bytes: bytes) -> bool:
    document = Document(BytesIO(docx_bytes))
    for section in document.sections:
        for paragraph in section.footer.paragraphs:
            xml = paragraph._element.xml
            if "PAGE" in xml and "NUMPAGES" in xml:
                return True
    return False


def _sample_proposal_content_json() -> dict:
    return {
        "sections": [
            {
                "label": "Executive Summary",
                "generation_status": "GENERATED",
                "content": {
                    "text": KNOWN_PROPOSAL_BODY,
                    "assumptions": [
                        "Exchange rate held at planning assumptions.",
                        "Partner schools remain accessible during term.",
                    ],
                },
            },
            {
                "label": "Project Design",
                "generation_status": "MANUAL_REQUIRED",
                "content": {"text": "", "assumptions": []},
            },
            {
                "label": "Budget Narrative",
                "generation_status": "FAILED",
                "content": {"text": "", "assumptions": ["Budget line items subject to audit."]},
            },
        ]
    }


def _render_proposal_sample() -> bytes:
    return build_proposal_docx_bytes(
        content_json=_sample_proposal_content_json(),
        opportunity_title="FCDO Girls' Education Fund",
        ngo_name="BridgeLight Education Trust",
        generated_at=datetime(2026, 6, 6, 12, 0, tzinfo=timezone.utc),
    )


def _render_me_sample(*, include_missing_section: bool = False) -> bytes:
    fcdo = _load_fcdo_template_fields()
    content_json = json.loads(FCDO_CONTENT_PATH.read_text(encoding="utf-8"))
    if include_missing_section:
        risk_section = next(
            s for s in content_json["sections"] if s.get("section_key") == "risk_and_safeguarding"
        )
        content_json["sections"] = [
            s
            for s in content_json["sections"]
            if s.get("section_key") != "recommendations_and_actions"
        ]
        content_json["sections"].append(
            {
                "section_key": "risk_and_safeguarding",
                "label": "Risk, Assumptions and Safeguarding",
                "generation_status": "ACCEPTED",
                "content": {
                    "text": risk_section["content"]["text"],
                    "assumptions": [
                        "Teacher focal points remain in post for the review period.",
                        "Safeguarding referrals follow agreed escalation routes.",
                    ],
                },
            }
        )
    return render_donor_report_docx(
        content_json=content_json,
        template_sections=fcdo["report_sections_json"],
        format_rules_json=fcdo["format_rules_json"],
        terminology_map_json=fcdo["terminology_map_json"],
        docx_template_ref=None,
        reporting_period_start="2024-10-15",
        reporting_period_end="2025-10-14",
        funder_name=fcdo["funder_name"],
        template_name=fcdo["template_name"],
        ngo_name="BridgeLight Education Trust",
        generated_at=datetime(2026, 6, 6, 12, 0, tzinfo=timezone.utc),
    )[0]


@pytest.mark.parametrize("render_fn", [_render_proposal_sample, lambda: _render_me_sample()])
def test_no_internal_artifacts(render_fn):
    text = _docx_all_text(render_fn())
    for marker in INTERNAL_MARKERS:
        assert marker not in text


def test_proposal_branded_title_block():
    text = _docx_all_text(_render_proposal_sample())
    assert "BridgeLight Education Trust" in text
    assert "Grant Proposal — FCDO Girls' Education Fund" in text
    assert "6 June 2026" in text


def test_me_branded_title_block():
    text = _docx_all_text(_render_me_sample())
    assert "BridgeLight Education Trust" in text
    assert "Donor Report —" in text
    assert "Reporting period: 2024-10-15 to 2025-10-14" in text
    assert "6 June 2026" in text


@pytest.mark.parametrize(
    "render_fn",
    [_render_proposal_sample, lambda: _render_me_sample(include_missing_section=True)],
)
def test_section_headings_use_word_styles(render_fn):
    headings = _heading_paragraphs(render_fn())
    assert headings
    for heading in headings:
        assert heading.style.name.startswith("Heading")
        assert "#" not in heading.text


def test_known_body_text_preserved_verbatim():
    proposal_text = _docx_all_text(_render_proposal_sample())
    assert KNOWN_PROPOSAL_BODY in proposal_text

    me_text = _docx_all_text(_render_me_sample())
    assert KNOWN_ME_BODY_SNIPPET in me_text


def test_assumptions_consolidated_in_single_appendix():
    proposal_text = _docx_all_text(_render_proposal_sample())
    assert ASSUMPTIONS_APPENDIX_TITLE in proposal_text
    assert "Exchange rate held at planning assumptions." in proposal_text
    assert "Budget line items subject to audit." in proposal_text
    assert proposal_text.count(ASSUMPTIONS_APPENDIX_TITLE) == 1

    me_bytes = _render_me_sample(include_missing_section=True)
    me_doc = Document(BytesIO(me_bytes))
    appendix_headings = [
        p.text
        for p in me_doc.paragraphs
        if p.style.name == "Heading 1" and p.text == ASSUMPTIONS_APPENDIX_TITLE
    ]
    assert len(appendix_headings) == 1
    me_text = _docx_all_text(me_bytes)
    assert "Teacher focal points remain in post for the review period." in me_text
    inline_assumption_headings = [
        p.text for p in me_doc.paragraphs if p.style.name == "Heading 3" and p.text == "Assumptions"
    ]
    assert inline_assumption_headings == []


def test_footer_present_with_page_number_fields():
    for render_fn in (_render_proposal_sample, lambda: _render_me_sample()):
        docx_bytes = render_fn()
        text = _docx_all_text(docx_bytes)
        assert "BridgeLight Education Trust" in text
        assert "Page" in text
        assert _footer_has_page_fields(docx_bytes)


def test_missing_me_section_renders_heading_only():
    docx_bytes = _render_me_sample(include_missing_section=True)
    text = _docx_all_text(docx_bytes)
    assert "Recommendations and Action Points" in text
    assert "[Section not generated]" not in text


def test_docxtpl_not_installed():
    with pytest.raises(ImportError):
        __import__("docxtpl")
