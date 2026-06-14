"""Package 1 - NGO-facing identifier-leak proving tests.

Proves the single redaction chokepoint clears every NGO-facing surface and that
the export tripwire fails the build on any identifier leak while staying silent on
legitimate NGO prose. On the Smoke Test P0 M&E allowlist (.github/workflows).
"""

from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document

from app.reports.eval.docx_export_assertions import (
    assert_export_docx,
    scan_identifier_leaks,
)
from app.reports.export.docx_renderer import _add_word_table, render_donor_report_docx
from app.reports.services.ngo_text_redaction import redact_internal_identifiers
from app.reports.services.section_prose import build_insufficiency_statement

# --- The five identifier shapes (one per leak class) ------------------------
LEAK_SHAPES = {
    "colon_item_key": "spend_summary:table:budget_vs_actual",
    "schema_dotted_path": "financials.lines.part_time_coordinator.budget",
    "citation_marker": "[fact:indicator.op_volunteers_recruited.actual]",
    "archetype_token": "ARCH_OUTCOMES_WITH_STORIES_AND_NUMBERS",
    "enum_value": "cannot_provide",
}

# --- Verbatim leaked strings from c1_nlcf_rewalk_export_d8e7518b.docx --------
COMMITTED_WALK_ASSUMPTIONS = [
    "No budget versus actual spend summary was available from gap answer "
    "spend_summary:table:budget_vs_actual.",
    "A full budget versus actual comparison was not available because "
    "spend_summary:table:budget_vs_actual was marked cannot_provide, so no variance "
    "analysis has been stated.",
    "No beneficiary numbers were available from gap answer "
    "difference_made:indicator:beneficiary_numbers.",
    "No gap answer content was provided for changes_and_next_steps:indicator:changes_made.",
    "No gap answer content was provided for changes_and_next_steps:indicator:planned_changes.",
    "No gap answer content was provided for changes_and_next_steps:indicator:support_needed.",
]

# --- Legitimate NGO prose that MUST NOT trip the tripwire -------------------
LEGITIMATE_NGO_SENTENCES = [
    "The Saturday sessions ran from 10:30 to 12:30 each week.",
    "We achieved a 3:1 ratio of volunteers to paid staff.",
    "Full details are on our website at https://example.org/impact and www.charity.org.",
    "Our closing reflection drew on John 3:16 during the session.",
    "Total spend was 29,950 against a budget of 31,200, a variance of 4.2 per cent.",
    "We ran sessions on Mon, Wed and Fri; attendance grew steadily.",
    "The team agreed three priorities: outreach, training and follow-up.",
    "Note: the budget for the year was tight but manageable.",
]


def _docx_text(docx_bytes: bytes) -> str:
    document = Document(BytesIO(docx_bytes))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _render(*, template_sections, content_json):
    docx_bytes, _ = render_donor_report_docx(
        content_json=content_json,
        template_sections=template_sections,
        format_rules_json={"document_title": "Test Report"},
        terminology_map_json={},
        docx_template_ref=None,
        reporting_period_start="2024-10-15",
        reporting_period_end="2025-10-14",
        funder_name="Test Funder",
        template_name="Test Template",
    )
    return docx_bytes


# ===========================================================================
# Tripwire - BOTH directions
# ===========================================================================
@pytest.mark.parametrize("shape_name,shape", LEAK_SHAPES.items())
def test_tripwire_fires_on_each_leak_shape(shape_name, shape):
    text = f"The funder asked for evidence on {shape} during the period."
    violations = scan_identifier_leaks(text)
    assert violations, f"tripwire missed leak shape {shape_name}: {shape}"


def test_tripwire_fires_on_generic_placeholder():
    assert scan_identifier_leaks("The template requires the required template items here.")


@pytest.mark.parametrize("sentence", LEGITIMATE_NGO_SENTENCES)
def test_tripwire_silent_on_legitimate_ngo_prose(sentence):
    assert scan_identifier_leaks(sentence) == [], f"false positive on: {sentence}"


def test_tripwire_in_assert_export_docx_fails_build_on_leak():
    document = Document()
    document.add_paragraph(
        "No beneficiary numbers were available from gap answer "
        "difference_made:indicator:beneficiary_numbers."
    )
    buffer = BytesIO()
    document.save(buffer)
    report = assert_export_docx(
        docx_bytes=buffer.getvalue(),
        content_json={"sections": []},
        template_sections=[],
    )
    assert not report.passed
    assert any(v.startswith("identifier_leak:") for v in report.violations)


def test_tripwire_in_assert_export_docx_passes_on_clean_doc():
    document = Document()
    document.add_paragraph("No beneficiary numbers were available in the submitted records.")
    buffer = BytesIO()
    document.save(buffer)
    report = assert_export_docx(
        docx_bytes=buffer.getvalue(),
        content_json={"sections": []},
        template_sections=[],
    )
    assert report.passed, report.violations


# ===========================================================================
# Chokepoint - meaning preservation on the committed-walk strings
# ===========================================================================
def test_committed_walk_assumptions_clean_and_meaning_preserved():
    for raw in COMMITTED_WALK_ASSUMPTIONS:
        assert scan_identifier_leaks(raw), f"precondition: raw should leak: {raw}"
        clean = redact_internal_identifiers(raw)
        assert scan_identifier_leaks(clean) == [], f"leak survived: {clean}"

    by_text = {raw: redact_internal_identifiers(raw) for raw in COMMITTED_WALK_ASSUMPTIONS}

    # Honest head clauses preserved verbatim; identifiers gone.
    assert by_text[COMMITTED_WALK_ASSUMPTIONS[0]] == (
        "No budget versus actual spend summary was available."
    )
    assert "could not be provided" in by_text[COMMITTED_WALK_ASSUMPTIONS[1]]
    assert "no variance analysis has been stated" in by_text[COMMITTED_WALK_ASSUMPTIONS[1]]
    assert by_text[COMMITTED_WALK_ASSUMPTIONS[2]] == (
        "No beneficiary numbers were available."
    )
    assert "support needed" in by_text[COMMITTED_WALK_ASSUMPTIONS[5]]
    assert "was provided" in by_text[COMMITTED_WALK_ASSUMPTIONS[5]]


# ===========================================================================
# Surface 1 - section BODY prose
# ===========================================================================
def test_surface_body_prose_clean():
    body = (
        "We delivered weekly sessions through the year. Evidence on "
        "difference_made:indicator:beneficiary_numbers and "
        "financials.lines.part_time_coordinator.budget was reviewed by the team."
    )
    assert scan_identifier_leaks(body)
    template_sections = [
        {"section_key": "story", "label": "The story of your project", "required_tables": []}
    ]
    content_json = {
        "sections": [
            {
                "section_key": "story",
                "generation_status": "ACCEPTED",
                "content": {"text": body},
            }
        ]
    }
    out = _docx_text(_render(template_sections=template_sections, content_json=content_json))
    assert scan_identifier_leaks(out) == [], out
    assert "We delivered weekly sessions through the year." in out


# ===========================================================================
# Surface 2 - assumptions / caveats appendix
# ===========================================================================
def test_surface_assumptions_appendix_clean():
    template_sections = [
        {"section_key": "spend", "label": "What you spent", "required_tables": []}
    ]
    content_json = {
        "sections": [
            {
                "section_key": "spend",
                "generation_status": "ACCEPTED",
                "content": {
                    "text": "We spent the grant on staffing and sessions across the year.",
                    "assumptions": COMMITTED_WALK_ASSUMPTIONS,
                },
            }
        ]
    }
    out = _docx_text(_render(template_sections=template_sections, content_json=content_json))
    assert scan_identifier_leaks(out) == [], out
    assert "No beneficiary numbers were available." in out


# ===========================================================================
# Surface 3 - insufficiency prose (colon-path-ref AND empty-refs)
# ===========================================================================
def test_surface_insufficiency_colon_path_ref_translated():
    section = {
        "section_key": "changes_and_next_steps",
        "label": "How you are changing what you do",
        "required_indicators": [
            "changes_and_next_steps:indicator:changes_made",
            "changes_and_next_steps:indicator:support_needed",
        ],
    }
    statement = build_insufficiency_statement(
        section=section,
        unsatisfied_refs=section["required_indicators"],
    )
    assert scan_identifier_leaks(statement) == [], statement
    assert "changes made" in statement
    assert "support needed" in statement
    assert "this section could not be drafted" in statement.lower()


def test_surface_insufficiency_empty_refs_kills_placeholder():
    section = {
        "section_key": "project_story",
        "label": "The story of your project this year",
        "required_indicators": [],
    }
    statement = build_insufficiency_statement(section=section)
    assert "the required template items" not in statement
    assert scan_identifier_leaks(statement) == [], statement
    assert "The story of your project this year" in statement


# ===========================================================================
# Surface 4 - table CELL content (no row-building; Package 2 fence honoured)
# ===========================================================================
def test_surface_table_cell_content_clean():
    document = Document()
    _add_word_table(
        document,
        header=["Cost type", "spend_summary:table:budget_vs_actual"],
        rows=[["financials.lines.part_time_coordinator.budget", "cannot_provide"]],
    )
    buffer = BytesIO()
    document.save(buffer)
    out = _docx_text(buffer.getvalue())
    assert scan_identifier_leaks(out) == [], out
    assert "Cost type" in out


# ===========================================================================
# Surface 5 - headings (section + table labels)
# ===========================================================================
def test_surface_headings_clean():
    template_sections = [
        {
            "section_key": "difference_made",
            "label": "Outcomes for difference_made:indicator:beneficiary_numbers",
            "required_tables": [
                {"table_key": "t", "label": "Spend on financials.lines.coordinator.budget"}
            ],
        }
    ]
    out = _docx_text(
        _render(template_sections=template_sections, content_json={"sections": []})
    )
    assert scan_identifier_leaks(out) == [], out
