from __future__ import annotations

import json
import sys
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

from app.reports.extraction.spreadsheet_input import (
    compute_spreadsheet_hash,
    get_cell_at_ref,
    list_data_row_ids,
    parse_csv_file,
    parse_docx_tables,
    parse_spreadsheet_from_path,
    spreadsheet_to_json_text,
)

FIXTURES = Path(__file__).resolve().parent / "fixtures" / "indicator_extractor"
XLSX = FIXTURES / "fcdo_bridgelight_indicator_data.xlsx"
ANSWER_KEY = FIXTURES / "keys" / "fcdo_bridgelight_indicator_data_answer_key.json"


def _load_key() -> dict:
    return json.loads(ANSWER_KEY.read_text(encoding="utf-8"))


@pytest.fixture(scope="module")
def parsed_workbook() -> dict:
    assert XLSX.is_file(), f"Run scripts/build_fcdo_indicator_fixture.py — missing {XLSX}"
    return parse_spreadsheet_from_path(XLSX)


def test_parse_xlsx_row_ids_match_answer_key(parsed_workbook: dict) -> None:
    key = _load_key()
    row_ids = list_data_row_ids(parsed_workbook, sheet_name=key["sheet_name"])
    assert row_ids == key["source_row_ids"]
    assert len(row_ids) == key["data_row_count"]


def test_cell_state_zero_blank_na(parsed_workbook: dict) -> None:
    key = _load_key()
    demo = key["cell_state_demo"]["expected"]
    target = get_cell_at_ref(
        parsed_workbook, "Indicators", key["cell_state_demo"]["target_cell_ref"]
    )
    actual = get_cell_at_ref(
        parsed_workbook, "Indicators", key["cell_state_demo"]["actual_cell_ref"]
    )
    unit = get_cell_at_ref(
        parsed_workbook, "Indicators", key["cell_state_demo"]["unit_cell_ref"]
    )
    assert target is not None
    assert target["cell_state"] == demo["target_cell_state"]
    assert target["raw"] == demo["target_raw"]
    assert actual is not None
    assert actual["cell_state"] == demo["actual_cell_state"]
    assert unit is not None
    assert unit["cell_state"] == demo["unit_cell_state"]
    assert unit["raw"] == demo["unit_raw"]


def test_blank_actual_cell_planted_b(parsed_workbook: dict) -> None:
    key = _load_key()
    cell = get_cell_at_ref(
        parsed_workbook, "Indicators", key["structural_cells"]["blank_actual_ref"]
    )
    assert cell is not None
    assert cell["cell_state"] == "blank"
    assert cell["raw"] is None


def test_hidden_row_present_in_grid(parsed_workbook: dict) -> None:
    key = _load_key()
    row_ids = list_data_row_ids(parsed_workbook)
    planted_id = key["planted_conflicts"]["planted_c_row_integrity"]["row_id"]
    assert planted_id in row_ids


def test_spreadsheet_hash_stable(parsed_workbook: dict) -> None:
    h1 = compute_spreadsheet_hash(parsed_workbook)
    h2 = compute_spreadsheet_hash(parse_spreadsheet_from_path(XLSX))
    assert h1 == h2


def test_spreadsheet_to_json_roundtrip(parsed_workbook: dict) -> None:
    text, truncated = spreadsheet_to_json_text(parsed_workbook)
    assert len(text) > 100
    assert truncated is False
    loaded = json.loads(text)
    assert loaded["format"] == "xlsx"


def _write_messy_monitoring_csv(path: Path) -> None:
    path.write_text(
        "Indicator,Target,Actual,Notes\n"
        "BEN-01,100,N/A,Quarterly headcount\n"
        "BEN-02,50,,Pending verification\n"
        "BEN-03,0,0,Zero baseline\n",
        encoding="utf-8",
    )


def test_parse_csv_messy_monitoring_fixture(tmp_path: Path) -> None:
    csv_path = tmp_path / "messy_actuals.csv"
    _write_messy_monitoring_csv(csv_path)
    parsed = parse_csv_file(csv_path)
    assert parsed["format"] == "csv"
    na_cell = get_cell_at_ref(parsed, csv_path.stem, "C2")
    blank_cell = get_cell_at_ref(parsed, csv_path.stem, "C3")
    zero_cell = get_cell_at_ref(parsed, csv_path.stem, "B4")
    assert na_cell is not None and na_cell["cell_state"] == "not_applicable"
    assert blank_cell is not None and blank_cell["cell_state"] == "blank"
    assert zero_cell is not None and zero_cell["raw"] == "0"


def _write_logframe_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("NLCF monitoring logframe — Q4 actuals")
    table = doc.add_table(rows=4, cols=4)
    headers = ("Indicator ID", "Target", "Actual", "Notes")
    for col_idx, label in enumerate(headers):
        table.rows[0].cells[col_idx].text = label
    table.rows[1].cells[0].text = "OUT-01"
    table.rows[1].cells[1].text = "120"
    table.rows[1].cells[2].text = "118"
    table.rows[1].cells[3].text = "Verified"
    table.rows[2].cells[0].text = "OUT-02"
    table.rows[2].cells[1].text = "80"
    table.rows[2].cells[2].text = "N/A"
    table.rows[2].cells[3].text = "Missing disaggregation"
    table.rows[3].cells[0].text = "OUT-03"
    table.rows[3].cells[1].text = "40"
    table.rows[3].cells[2].text = ""
    table.rows[3].cells[3].text = "  pending  "
    doc.save(path)


def test_parse_docx_tables_via_docling_mock(tmp_path: Path) -> None:
    docx_path = tmp_path / "nlcf_logframe.docx"
    _write_logframe_docx(docx_path)

    import pandas as pd

    mock_df = pd.DataFrame(
        {
            "Indicator ID": ["OUT-01", "OUT-02", "OUT-03"],
            "Target": ["120", "80", "40"],
            "Actual": ["118", "N/A", ""],
            "Notes": ["Verified", "Missing disaggregation", "  pending  "],
        }
    )
    mock_table = MagicMock()
    mock_table.export_to_dataframe.return_value = mock_df
    mock_document = MagicMock()
    mock_document.tables = [mock_table]
    mock_result = MagicMock()
    mock_result.document = mock_document
    mock_converter = MagicMock()
    mock_converter.convert.return_value = mock_result
    mock_converter_class = MagicMock(return_value=mock_converter)
    fake_docling = MagicMock()
    fake_docling.document_converter.DocumentConverter = mock_converter_class

    with patch.dict(
        sys.modules,
        {
            "docling": MagicMock(),
            "docling.document_converter": fake_docling.document_converter,
        },
    ):
        parsed = parse_docx_tables(docx_path)

    assert parsed["format"] == "docx"
    assert parsed["sheets"][0]["name"] == "Table1"
    actual_cell = get_cell_at_ref(parsed, "Table1", "C3")
    blank_cell = get_cell_at_ref(parsed, "Table1", "C4")
    assert actual_cell is not None and actual_cell["cell_state"] == "not_applicable"
    assert blank_cell is not None and blank_cell["cell_state"] == "blank"
    assert get_cell_at_ref(parsed, "Table1", "D4")["raw"] == "pending"


def test_parse_spreadsheet_from_path_routes_docx(tmp_path: Path) -> None:
    docx_path = tmp_path / "monitoring.docx"
    _write_logframe_docx(docx_path)

    import pandas as pd

    mock_df = pd.DataFrame({"A": ["1"], "B": ["2"]})
    mock_table = MagicMock()
    mock_table.export_to_dataframe.return_value = mock_df
    mock_document = MagicMock()
    mock_document.tables = [mock_table]
    mock_result = MagicMock()
    mock_result.document = mock_document
    mock_converter = MagicMock()
    mock_converter.convert.return_value = mock_result
    mock_converter_class = MagicMock(return_value=mock_converter)
    fake_docling = MagicMock()
    fake_docling.document_converter.DocumentConverter = mock_converter_class

    with patch.dict(
        sys.modules,
        {
            "docling": MagicMock(),
            "docling.document_converter": fake_docling.document_converter,
        },
    ):
        parsed = parse_spreadsheet_from_path(docx_path)

    assert parsed["format"] == "docx"


def test_parse_docx_without_tables_raises_for_degrade_path(tmp_path: Path) -> None:
    docx_path = tmp_path / "empty.docx"
    Document().save(docx_path)

    mock_document = MagicMock()
    mock_document.tables = []
    mock_result = MagicMock()
    mock_result.document = mock_document
    mock_converter = MagicMock()
    mock_converter.convert.return_value = mock_result
    mock_converter_class = MagicMock(return_value=mock_converter)
    fake_docling = MagicMock()
    fake_docling.document_converter.DocumentConverter = mock_converter_class

    with patch.dict(
        sys.modules,
        {
            "docling": MagicMock(),
            "docling.document_converter": fake_docling.document_converter,
        },
    ):
        with pytest.raises(ValueError, match="No tables found"):
            parse_spreadsheet_from_path(docx_path)
