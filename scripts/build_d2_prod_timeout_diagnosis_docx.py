#!/usr/bin/env python3
"""Generate Word doc for D2 proposal-extractor prod timeout diagnosis."""

from __future__ import annotations

import io
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

REPO_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = REPO_ROOT / "D2_PROPOSAL_EXTRACTOR_PROD_TIMEOUT_DIAGNOSIS_2026-05-31.docx"


def _load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/segoeui.ttf",
        "C:/Windows/Fonts/arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    if bold:
        candidates = [
            "C:/Windows/Fonts/segoeuib.ttf",
            "C:/Windows/Fonts/arialbd.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        ] + candidates
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def _draw_diagram() -> bytes:
    w, h = 920, 520
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    title_font = _load_font(18, bold=True)
    box_font = _load_font(14)
    note_font = _load_font(12)
    small_font = _load_font(11)

    draw.text((w // 2, 24), "D2 extract stage — wall-clock boundaries", fill="#1a1a1a", font=title_font, anchor="mm")

    boxes = [
        (120, 70, 320, 130, "Extract stage\n(orchestrator)", "#E8F0FE"),
        (380, 70, 580, 130, "load_document_text\n(Docling re-parse)", "#FFF4E5"),
        (640, 70, 840, 130, "extract_proposal_text\nasyncio.wait_for 90s", "#FCE8E6"),
        (640, 190, 840, 250, "claude_agent_sdk.query\n(subprocess + model)", "#FCE8E6"),
        (640, 310, 840, 370, "Stream open — no\nResultMessage yet", "#FCE8E6"),
        (640, 430, 840, 490, "TimeoutError at 90s\nSTOP_TIMEOUT", "#EA4335"),
    ]

    def rounded_box(x1, y1, x2, y2, text, fill, outline="#5F6368"):
        draw.rounded_rectangle([x1, y1, x2, y2], radius=12, fill=fill, outline=outline, width=2)
        lines = text.split("\n")
        line_h = 18
        total_h = len(lines) * line_h
        start_y = y1 + ((y2 - y1) - total_h) // 2 + 4
        for i, line in enumerate(lines):
            draw.text(((x1 + x2) // 2, start_y + i * line_h), line, fill="#202124", font=box_font, anchor="mm")

    for x1, y1, x2, y2, text, fill in boxes:
        rounded_box(x1, y1, x2, y2, text, fill)

    def arrow(x1, y1, x2, y2, color="#5F6368"):
        draw.line([x1, y1, x2, y2], fill=color, width=2)
        if y2 > y1:
            draw.polygon([(x2, y2), (x2 - 6, y2 - 10), (x2 + 6, y2 - 10)], fill=color)
        elif x2 > x1:
            draw.polygon([(x2, y2), (x2 - 10, y2 - 6), (x2 - 10, y2 + 6)], fill=color)

    arrow(320, 100, 380, 100)
    arrow(580, 100, 640, 100)
    arrow(740, 130, 740, 190)
    arrow(740, 250, 740, 310)
    arrow(740, 370, 740, 430)

    draw.text((400, 145), "Outside 90s timer", fill="#B06000", font=note_font)
    draw.line([395, 132, 470, 132], fill="#B06000", width=1)
    draw.text((700, 395), "Inside 90s timer", fill="#C5221F", font=note_font)

    draw.text(
        (40, 200),
        "Classify stage (earlier):\nalso Docling-parses all 4 uploads\n— also outside D2 90s",
        fill="#5F6368",
        font=small_font,
    )

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_para(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def _add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = val
    doc.add_paragraph()


def build() -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("D2 Proposal Extractor — Production Timeout Diagnosis", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f"Read-only diagnosis · {date.today().isoformat()}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.font.color.rgb = RGBColor(0x5F, 0x63, 0x68)
        r.font.size = Pt(10)

    _add_para(
        doc,
        "Verdict: The failure is a clean stage failure at the D2 wall-clock backstop "
        "(asyncio.wait_for, 90s), fired while the Agent SDK stream had not yet yielded a "
        "ResultMessage. The best-supported explanation is genuine model/SDK-work latency on "
        "this document's slow tail (same input that once completed in ~72s now sometimes needs "
        ">90s). Docling adds job time but is not inside that 90s bucket. SDK subprocess hang "
        "remains possible but is not separable from slow completion with current prod logs. "
        "D3/D4 did not run on these walks; D2 is isolated and uniquely hard-fails on timeout.",
    )

    _add_heading(doc, "What failed (prod evidence)", 2)
    _add_table(
        doc,
        ["Attempt", "Report ID", "Job ID", "Error"],
        [
            [
                "1",
                "1f35153d-5277-46e1-9ca6-2d2a2c37eafa",
                "3e765004-fff3-40ca-b538-ed2764558883",
                "extract: Proposal extractor exceeded 90s timeout",
            ],
            [
                "2",
                "39805e82-d47e-4255-bba0-2b691449ea2e",
                "e4e6ee07-ec96-4974-8d50-ed136fbc2ad3",
                "Same",
            ],
        ],
    )
    _add_bullet(doc, "Timeout raised in extract_proposal_text → ProposalExtractorError(STOP_TIMEOUT)")
    _add_bullet(doc, "Normalized to StageFailure → mark_job_failed(..., event=pipeline_exception)")
    _add_bullet(doc, "Stack frame at proposal_extractor.py:303 — inside async for message in query_fn(...), before any ResultMessage")
    _add_bullet(doc, "Post-hardening clean failure — not a stuck worker")
    _add_para(doc, "Smoke-walk polling (POLL_SECONDS=12):", bold=True)
    _add_bullet(doc, "Attempt 1: ~2 polls in classify, ~8 in extract → ~96s+ visible in extract")
    _add_bullet(doc, "Attempt 2: ~1 poll in classify, ~7 in extract → ~84s+ visible in extract")

    _add_heading(doc, "Wall-clock breakdown", 2)
    _add_para(doc, "Sequence diagram (phase boundaries):", bold=True)
    doc.add_picture(io.BytesIO(_draw_diagram()), width=Inches(6.5))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_table(
        doc,
        ["Phase", "In 90s timer?", "Evidence"],
        [
            ["S3 fetch + temp write", "No", "document_intake.load_document_text before extract_and_persist_proposal"],
            ["Docling parse (proposal)", "No", "Same — runs in pipeline before D2"],
            ["SDK spawn + model + structured JSON", "Yes", "asyncio.wait_for(_run_extractor_query, timeout=TIMEOUT_SECONDS)"],
            ["Persist / dispatch", "No", "Only after D2 returns or raises"],
        ],
    )
    _add_para(
        doc,
        "Classify stage also Docling-parses all four uploads (3× .docx + 1× .xlsx) before extract; "
        "that time is outside the D2 90s but adds total job latency.",
    )
    _add_para(
        doc,
        "Gap: Prod DB agent_trace_json and proposal_extractor start log lines were not available "
        "from this environment, so the 90s cannot be split into spawn vs TTFT vs generation on prod.",
        bold=True,
    )

    _add_heading(doc, "Document characteristics", 2)
    _add_table(
        doc,
        ["Metric", "Value"],
        [
            ["File size", "41,503 bytes"],
            ["Docling text", "12,993 chars (~1,545 words)"],
            ["Truncated?", "No (MAX_INPUT_CHARS=120,000)"],
            ["Content hash", "422b2e10… — matches recorded gate fixture"],
            ["Structure", "Markdown tables, logframe-style indicators, VfM section"],
            ["Model task", "Large structured extraction: 2 objectives + activities + 16 indicators via json_schema"],
        ],
    )
    _add_para(doc, "Recorded successful gate run on this exact hash:", bold=True)
    _add_bullet(doc, "model_used: haiku")
    _add_bullet(doc, "latency_ms: 72,363 (~72.4s SDK-reported duration)")
    _add_bullet(doc, "output_tokens: 12,904 (~13k structured JSON)")
    _add_bullet(doc, "Headroom at 90s ceiling: ~17.6s")
    _add_para(
        doc,
        "Decision log D-033 documented bimodal behavior: successes at 72–74s, failures >80s at a 75s ceiling; "
        "90s was chosen as margin. Prod now hits the slow tail above even 90s.",
    )
    _add_para(
        doc,
        "Local Docling (same file): first parse in cold process ~74s (model load); warm re-parse ~0.3–9s. "
        "Affects classify/extract preamble, not the error string.",
    )

    _add_heading(doc, "Which timeout fired?", 2)
    _add_table(
        doc,
        ["Layer", "Value", "Applies to D2?"],
        [
            ["asyncio.wait_for in extract_proposal_text", "ME_CLASSIFIER_TIMEOUT_SECONDS default 90", "Yes — this fired"],
            ["API_TIMEOUT_MS in SDK subprocess env", "90,000 ms", "Parallel inner ceiling"],
            ["dispatch_stage(per_call_timeout_seconds=…)", "Not passed for D2", "No extra dispatch timeout"],
            ["Worker ME_WORKER_JOB_TIMEOUT_SECONDS", "3600", "No — job failed at ~2 min"],
        ],
    )
    _add_para(doc, "Prod env: ME_CLASSIFIER_MODEL=haiku; no ME_CLASSIFIER_TIMEOUT_SECONDS override.")

    _add_heading(doc, "Candidate causes — separated", 2)

    _add_heading(doc, "1. Genuine model latency — PRIMARY (strongest evidence)", 3)
    for b in [
        "Same document hash as a 72.3s successful extraction with ~13k output tokens.",
        "D-033: instrumented bimodal latency; 90s was a bet on covering the slow tail — prod shows it is not always sufficient.",
        "Failures align with exact 90s backstop, reproducible across two runs.",
    ]:
        _add_bullet(doc, b)

    _add_heading(doc, "2. SDK subprocess overhead / hang — possible, not proven", 3)
    for b in [
        "D2 uses claude_agent_sdk.query (CLI subprocess); timeout fires while still awaiting stream messages.",
        "Same signature as hang or slow model; no prod evidence of subprocess exit code or partial stream events.",
        "D-033 noted timeouts were wall-time long poles, not structured-output retry subtypes — slightly favors slow completion.",
    ]:
        _add_bullet(doc, b)

    _add_heading(doc, "3. Document-handling cost — contributing, not the 90s trigger", 3)
    for b in [
        "Docling is explicitly outside asyncio.wait_for in the code path.",
        "Re-parsing at extract adds seconds (warm) to minutes (cold worker) but does not consume the 90s D2 budget.",
        "Cost driver is LLM structured output volume, not parse failure.",
    ]:
        _add_bullet(doc, b)

    _add_heading(doc, "4. Environment difference — contributing amplifier", 3)
    for b in [
        "Same model (haiku) and timeout defaults as gate/dev.",
        "Smoke walk uploads 4 docs vs earlier 3-doc prod success — extra classify Docling load before D2.",
        "Attempt 2 reached extract faster than attempt 1 — consistent with warm Docling / worker state.",
    ]:
        _add_bullet(doc, b)

    _add_heading(doc, "D3/D4 on the same run?", 2)
    _add_para(doc, "No. Proposal is first in upload order; D2 hard-fails the stage. D3/D4 never executed.")
    _add_table(
        doc,
        ["Agent", "On 90s timeout"],
        [
            ["D2 proposal", "Hard fail — ProposalExtractorError → stage failure"],
            ["D3 grant-terms", "2×90s retry → degraded (D-035)"],
            ["D4 indicator", "2×90s retry → degraded (D-036)"],
        ],
    )

    _add_heading(doc, "What would settle remaining ambiguity (read-only)", 2)
    for b in [
        "agent_trace_json on failed jobs (stages.classify.completed_at vs failure.at) — needs DB access inside Railway network.",
        "Worker logs with proposal_extractor start filename=… chars=12993 correlated with failure ~90s later.",
        "SDK-level observation: time of first stream message vs ResultMessage.duration_ms on read-only replay.",
    ]:
        _add_bullet(doc, b)
    _add_para(doc, "Without those, cause 1 and 2 are observationally identical at the 90s boundary.")

    _add_heading(doc, "Recommended fix direction (diagnosis only — no changes made)", 2)
    _add_table(
        doc,
        ["Priority", "Fix class", "Rationale"],
        [
            ["1", "Resilience parity with D3/D4", "D-035-style bounded retry + degraded continuation; orchestrator already continues on degraded extraction."],
            ["2", "D2-specific timeout budget", "Separate env (e.g. ME_PROPOSAL_TIMEOUT_SECONDS) above 90s; 72s typical / >90s tail."],
            ["3", "Output/latency reduction", "Shrink structured payload to cut ~13k output tokens."],
            ["4", "Runtime path (longer-term)", "Messages API migration for D2 (as D1/E1/E3); removes subprocess uncertainty."],
            ["Lower", "Docling caching classify→extract", "Improves total job time; won't fix pure >90s model tail alone."],
        ],
    )

    _add_heading(doc, "Bottom line", 2)
    _add_para(
        doc,
        "Why D2 exceeds 90s on prod for this document: The Agent SDK extraction for this FCDO proposal "
        "sometimes requires more than 90 seconds of wall-clock work to finish streaming structured output. "
        "That is expected slow-tail behavior for this agent (documented in D-033, observed at 72.3s on success), "
        "now manifesting above the current ceiling. Docling and classify overhead add job latency but are not "
        "what triggers the error. Subprocess hang cannot be excluded but is less likely than slow completion "
        "given prior instrumentation and identical input hash. D2 alone hard-fails where sibling extractors would "
        "degrade and continue.",
    )

    _add_para(doc, "Scope: Read-only diagnosis. No code, config, env, timeout, or deploy changes.", bold=True)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path}")
