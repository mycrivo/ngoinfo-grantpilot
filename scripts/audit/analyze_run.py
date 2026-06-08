#!/usr/bin/env python3
"""Analyze a completed walk artifact: sections, critic flags, gap-wall, docx-vs-KB leak.

Usage: python -m scripts.audit.analyze_run <walk_artifact.json> [export.docx]
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from scripts.audit._common import ARTIFACT_DIR


def _norm(s: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", str(s).lower())


# Gap-wall category heuristics over required_item_ref.
FUNDER_SIDE = {
    "output_scores", "output_score_table", "impact_weightings", "risk_ratings",
    "economy", "efficiency", "effectiveness", "equity", "vfm_measures",
    "commercial_improvement_where_relevant", "FCDO_management_actions",
}
DATA_BACKED_HINTS = {
    "actual_results": ["ar1_actual"],
    "output_indicators": ["indicators.", "ar1_milestone_target"],
    "outcome_indicators": ["indicators.", "proposal_target"],
    "logframe_milestones": ["ar1_milestone_target"],
    "progress_against_expected_results": ["ar1_actual", "ar1_milestone_target"],
    "forecast_vs_actual_costs": ["financials.lines"],
    "forecast_vs_actual_spend": ["financials.lines"],
    "financial_delivery": ["financials.lines"],
    "cost_drivers": ["financials.lines"],
}
NARRATIVE = {
    "overall_progress", "main_results_achieved", "main_issues", "key_recommendations",
    "major_deviations", "new_evidence", "evaluation_progress", "evidence_base_strength",
    "data_quality_limitations", "new_risks", "realised_assumptions",
    "funds_not_used_as_intended_risk", "climate_environment_risk",
    "safeguarding_risk_where_relevant", "partner_performance",
    "supplier_or_consultant_performance", "commercial_or_procurement_issues",
    "recommendations_from_current_review", "updates_on_previous_recommendations",
    "priorities_for_next_period", "recommendations_action_plan", "review_summary_sheet",
    "outcome_assessment", "gender_age_or_vulnerable_group_disaggregation_where_relevant",
}


def categorize_gap(ref: str, fact_keys: list[str]) -> tuple[str, str]:
    if ref in FUNDER_SIDE:
        return "funder_side_assessment", "reviewer/funder-authored field; engine asks NGO to author it"
    if ref.startswith("logframe_row:"):
        rid = ref.split(":", 1)[1].replace("_", ".").upper()  # op2_3 -> OP2.3
        token = _norm(rid)
        has_actual = any(_norm(k).startswith(_norm("indicators" + rid) ) for k in fact_keys) or \
            any(token in _norm(k) and "ar1actual" in _norm(k) for k in fact_keys)
        if has_actual:
            return "data_backed_but_asked", f"actual present for {rid} but flagged"
        return "genuine_data_gap", f"no AR1 actual fact for {rid} in KB"
    if ref in DATA_BACKED_HINTS:
        hints = DATA_BACKED_HINTS[ref]
        present = [h for h in hints if any(h in k for k in fact_keys)]
        if present:
            return "data_backed_but_asked", f"KB holds matching data ({', '.join(present)}) but matcher missed it"
        return "genuine_data_gap", "no matching data facts in KB"
    if ref in NARRATIVE:
        return "narrative_or_judgment", "narrative content; should be synthesized or human-authored"
    return "other", "uncategorized"


def analyze(artifact_path: Path) -> dict:
    d = json.loads(artifact_path.read_text(encoding="utf-8"))
    snaps = d.get("snapshots") or {}
    final = snaps.get("after_export") or snaps.get("after_critique") or snaps.get("after_synthesis") or {}
    report = final.get("report") or {}
    kb = report.get("knowledge_bank_json") or {}
    facts = kb.get("facts") or {}
    fact_keys = list(facts.keys())
    content = report.get("content_json") or {}
    sections = content.get("sections") or []

    # Gap-wall from after_gap snapshot.
    gap_report = (snaps.get("after_gap") or {}).get("report") or {}
    gaps = (gap_report.get("gap_analysis_json") or {}).get("gaps") or []
    gap_cats: dict[str, int] = {}
    gap_detail = []
    for g in gaps:
        cat, why = categorize_gap(g.get("required_item_ref", ""), fact_keys)
        gap_cats[cat] = gap_cats.get(cat, 0) + 1
        gap_detail.append({"ref": g.get("required_item_ref"), "section": g.get("section_key"),
                           "type": g.get("required_item_type"), "category": cat, "why": why})

    # Section + critic-flag analysis.
    sec_rows = []
    all_block_flags = []
    for s in sections:
        block = s.get("content") or {}
        flags = s.get("critic_flags") or []
        blocks = [f for f in flags if isinstance(f, dict) and f.get("severity") == "BLOCK"]
        warns = [f for f in flags if isinstance(f, dict) and f.get("severity") == "WARN"]
        for f in blocks:
            all_block_flags.append({"section": s.get("section_key"), **{k: f.get(k) for k in
                                    ("claim", "reason", "severity", "specific", "status", "detail", "message")}})
        sec_rows.append({
            "section_key": s.get("section_key"),
            "generation_status": s.get("generation_status"),
            "text_len": len(block.get("text") or ""),
            "evidence_used": block.get("evidence_used") or [],
            "dropped_citations": block.get("dropped_citations") or [],
            "critic_flags_total": len(flags),
            "critic_blocks": len(blocks),
            "critic_warns": len(warns),
            "failure_reason": s.get("failure_reason"),
        })

    return {
        "report_id": d.get("report_id"),
        "verdict": d.get("verdict"),
        "facts_total": len(facts),
        "gaps_total": len(gaps),
        "gap_categories": gap_cats,
        "gap_detail": gap_detail,
        "sections": sec_rows,
        "block_flags": all_block_flags,
        "report_status": report.get("status"),
        "created_at": str(report.get("created_at")),
        "updated_at": str(report.get("updated_at")),
        "gate_timestamps": {k: kb.get(k) for k in
                            ("gate1_confirmed_at", "gate2_confirmed_at", "gate3_confirmed_at")},
    }


def docx_vs_kb(docx_path: Path, artifact_path: Path) -> dict:
    from docx import Document

    d = json.loads(artifact_path.read_text(encoding="utf-8"))
    snaps = d.get("snapshots") or {}
    final = snaps.get("after_export") or {}
    kb = (final.get("report") or {}).get("knowledge_bank_json") or {}
    facts = kb.get("facts") or {}
    gap_answers = kb.get("gap_answers") or {}

    doc = Document(str(docx_path))
    text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            text += "\n" + " | ".join(c.text for c in row.cells)

    # All numbers in the rendered doc.
    doc_numbers = set(re.findall(r"\b\d[\d,]*(?:\.\d+)?\b", text))
    # Numbers backed by KB facts or gap answers.
    kb_text = json.dumps(facts) + json.dumps(gap_answers)
    kb_numbers = set(re.findall(r"\b\d[\d,]*(?:\.\d+)?\b", kb_text))

    unbacked = sorted([n for n in doc_numbers if n not in kb_numbers and len(n.replace(",", "")) >= 2],
                      key=lambda x: -len(x))
    return {
        "docx_paragraphs": len(doc.paragraphs),
        "docx_tables": len(doc.tables),
        "docx_chars": len(text),
        "doc_numbers_count": len(doc_numbers),
        "kb_numbers_count": len(kb_numbers),
        "numbers_in_doc_not_in_kb": unbacked[:60],
        "placeholder_markers": text.count("[Section not generated]") + text.count("not generated"),
    }


def main() -> int:
    if len(sys.argv) < 2:
        cands = sorted(ARTIFACT_DIR.glob("walk_*.json"))
        if not cands:
            print("no artifact found"); return 1
        art = cands[-1]
    else:
        art = Path(sys.argv[1])
    result = analyze(art)
    if len(sys.argv) >= 3:
        result["docx_vs_kb"] = docx_vs_kb(Path(sys.argv[2]), art)
    out = ARTIFACT_DIR / f"analysis_{result['report_id'][:8]}.json"
    out.write_text(json.dumps(result, indent=2, default=str), encoding="utf-8")
    print(json.dumps({k: v for k, v in result.items()
                      if k in ("report_id", "verdict", "facts_total", "gaps_total",
                               "gap_categories", "report_status", "created_at",
                               "updated_at", "gate_timestamps")}, indent=2, default=str))
    print("SECTIONS:")
    for s in result["sections"]:
        print(f"  {s['section_key']}: status={s['generation_status']} text_len={s['text_len']} "
              f"evidence={len(s['evidence_used'])} blocks={s['critic_blocks']} warns={s['critic_warns']} "
              f"fail={s['failure_reason']}")
    print(f"BLOCK_FLAGS ({len(result['block_flags'])}):")
    for f in result["block_flags"]:
        print(f"  [{f['section']}] {json.dumps({k:v for k,v in f.items() if k!='section' and v})[:300]}")
    if "docx_vs_kb" in result:
        print("DOCX_VS_KB:")
        print(json.dumps(result["docx_vs_kb"], indent=2)[:2000])
    print(f"ANALYSIS={out}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
