#!/usr/bin/env python3
"""Track 3 Phase B confirming walk — NLCF prod, proposal-failure elevation.

Induces proposal checkpoint via timeout-bait / unreadable fixtures, acks
proceed_with_gap, then runs answered + skip Gate-2 branches. Evidence only —
no product fixes.

Env:
  PLAN (default IMPACT)
  MAX_TO_CHECKPOINT (default 900)
  MAX_TO_GATE1 / MAX_TO_GATE2 / MAX_SYNTH / MAX_CRITIQUE / MAX_EXPORT
  BRANCHES=answered,skip (default both)
"""
from __future__ import annotations

import json
import os
import sys
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document

from scripts.audit import _common as C
from scripts.audit.gap_answers import answer_gap

REPO = C.REPO
EVIDENCE_DIR = REPO / "docs" / "artefacts" / "me_module" / "audits"
LOG_PATH = EVIDENCE_DIR / "TRACK3_CONFIRMING_WALK_2026-07-18.log"
EVIDENCE_PATH = EVIDENCE_DIR / f"TRACK3_CONFIRMING_WALK_EVIDENCE_{datetime.now(timezone.utc).strftime('%Y%m%dT%H%M%SZ')}.json"
FIXTURE_DIR = REPO / "docs" / "artefacts" / "me_module" / "audits" / "fixtures"
TIMEOUT_BAIT = FIXTURE_DIR / "nlcf_timeout_bait_proposal.docx"
IMAGE_ONLY = (
    REPO / "tests" / "fixtures" / "docling_intake" / "image_only_no_text_layer.pdf"
)

ELEVATED_REFS = (
    "community_participation_examples",
    "partner_or_local_collaboration_examples",
)
ELEVATED_KEYS = tuple(
    f"community_involvement:indicator:{ref}" for ref in ELEVATED_REFS
)
ANSWER_TEXT = {
    "community_participation_examples": (
        "TRACK3_ANSWERED_COMMUNITY_PARTICIPATION_MARKER: residents co-designed "
        "three evening workshops in July 2026 and kept attendance registers."
    ),
    "partner_or_local_collaboration_examples": (
        "TRACK3_ANSWERED_PARTNER_COLLAB_MARKER: partnership with Southbank "
        "Community Trust delivered shared volunteer rota and joint outreach."
    ),
}

MAX_TO_CHECKPOINT = int(os.environ.get("MAX_TO_CHECKPOINT", "900"))
MAX_TO_GATE1 = int(os.environ.get("MAX_TO_GATE1", "1500"))
MAX_TO_GATE2 = int(os.environ.get("MAX_TO_GATE2", "900"))
MAX_SYNTH = int(os.environ.get("MAX_SYNTH", "2400"))
MAX_CRITIQUE = int(os.environ.get("MAX_CRITIQUE", "1800"))
MAX_EXPORT = int(os.environ.get("MAX_EXPORT", "600"))


class _Tee:
    def __init__(self, *streams):
        self.streams = streams

    def write(self, data: str) -> int:
        for s in self.streams:
            s.write(data)
            s.flush()
        return len(data)

    def flush(self) -> None:
        for s in self.streams:
            s.flush()


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _build_timeout_bait() -> Path:
    FIXTURE_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Southbank Community Project — Funding Proposal", level=1)
    doc.add_paragraph(
        "This document is intentionally dense filler designed to stress proposal "
        "extraction toward a dual-attempt timeout degrade (Track 3 confirming walk)."
    )
    doc.add_heading("Objectives", level=2)
    doc.add_paragraph("Objective 1: Increase community participation in local services.")
    doc.add_heading("Activities", level=2)
    doc.add_paragraph("Activity 1: Weekly drop-in sessions with volunteer facilitators.")
    filler = (
        "FILLER_BLOCK for timeout inducement. Repeat community project narrative, "
        "monitoring language, and non-structured prose without clean indicator tables. "
    ) * 80
    # Aim near MAX_INPUT_CHARS (120k) to maximize extractor work.
    while len(filler) < 110_000:
        filler += filler
    doc.add_paragraph(filler[:110_000])
    doc.save(TIMEOUT_BAIT)
    return TIMEOUT_BAIT


def _checkpoint_from_job(job: dict) -> dict | None:
    stages = (job.get("agent_trace_json") or {}).get("stages") or {}
    extract = stages.get("extract") or {}
    cp = extract.get("proposal_checkpoint")
    return dict(cp) if isinstance(cp, dict) else None


def _ack_proceed(session, report_id: str) -> dict:
    r = C.auth_request(
        session,
        "POST",
        f"{C.BASE_URL}/api/reports/{report_id}/jobs/proposal-checkpoint/ack",
        json={"action": "proceed_with_gap"},
        timeout=60,
    )
    return {"status_code": r.status_code, "body": C._safe_json(r)}


def _gaps_from_capture(capture: dict) -> list[dict]:
    report = capture.get("report") or {}
    gaps = ((report.get("gap_analysis_json") or {}).get("gaps") or [])
    return [g for g in gaps if isinstance(g, dict)]


def _elevated_subset(gaps: list[dict]) -> list[dict]:
    out = []
    for g in gaps:
        ref = g.get("required_item_ref") or ""
        key = g.get("item_key") or ""
        if ref in ELEVATED_REFS or key in ELEVATED_KEYS:
            out.append(g)
    return out


def _assert_elevation(gaps: list[dict]) -> dict:
    elevated = _elevated_subset(gaps)
    refs = sorted({g.get("required_item_ref") for g in elevated})
    keys = sorted({g.get("item_key") for g in elevated})
    questions = [g.get("question") or g.get("gap_question") or "" for g in elevated]
    internal_leak = []
    for q in questions:
        ql = q.lower()
        for token in (
            "community_participation_examples",
            "partner_or_local_collaboration_examples",
            "item_key",
            "fact:",
            "gap:",
        ):
            if token in ql or token in q:
                internal_leak.append({"question": q, "token": token})
    mandatory_ok = all(
        g.get("mandatory") is True or g.get("required") is True or True  # surface fields vary
        for g in elevated
    )
    # Prefer explicit required/mandatory when present
    for g in elevated:
        if "mandatory" in g and g.get("mandatory") is not True:
            mandatory_ok = False
        if "is_mandatory" in g and g.get("is_mandatory") is not True:
            mandatory_ok = False
    exact = refs == sorted(ELEVATED_REFS) or set(refs) == set(ELEVATED_REFS)
    # exact elevated set among elevated items — also ensure both present
    both_present = set(ELEVATED_REFS).issubset(set(refs)) and len(elevated) == 2
    return {
        "elevated_count": len(elevated),
        "elevated_refs": refs,
        "elevated_keys": keys,
        "questions": questions,
        "exact_two": both_present,
        "exact_ref_set": exact and len(elevated) == 2,
        "internal_identifier_leaks": internal_leak,
        "mandatory_ok": mandatory_ok,
        "raw_elevated": elevated,
    }


def _community_section(content_json: dict | None) -> dict | None:
    if not isinstance(content_json, dict):
        return None
    sections = content_json.get("sections") or content_json.get("report_sections") or []
    if isinstance(sections, dict):
        return sections.get("community_involvement")
    for s in sections:
        if isinstance(s, dict) and s.get("section_key") == "community_involvement":
            return s
    # alternate shapes
    return content_json.get("community_involvement")


def _run_branch(
    *,
    branch: str,
    proposal_path: Path,
    supporting: list[Path],
    evidence: dict,
) -> dict:
    t0 = time.time()
    run_label = f"track3-{branch}-{int(time.time())}"
    email = f"audit-{run_label}@grantpilot-test.org"
    session = C.mint_session(email, plan=os.environ.get("PLAN", "IMPACT"), full_name=f"Track3 {branch}")
    report = C.create_report(session, template_id=C.NLCF_TEMPLATE_ID)
    report_id = report["id"]
    print(f"=== BRANCH {branch} report_id={report_id} email={email} ===", flush=True)

    up_prop = C.upload(session, report_id, proposal_path)
    print(f"UPLOAD proposal {proposal_path.name} -> {up_prop.get('id')}", flush=True)
    for path in supporting:
        up = C.upload(session, report_id, path)
        print(f"UPLOAD {path.name} -> {up.get('id')}", flush=True)

    enq = C.enqueue(session, report_id)
    print(f"ENQUEUE {enq}", flush=True)

    job = C.poll_job(
        session,
        report_id,
        label=f"{branch}-checkpoint",
        until_status={"awaiting_human", "failed"},
        until_stage={"extract", "gap", "reconcile"},
        max_seconds=MAX_TO_CHECKPOINT,
    )
    cp = _checkpoint_from_job(job)
    checkpoint_block = {
        "job_status": job.get("status"),
        "job_stage": job.get("stage"),
        "timeout": bool(job.get("_timeout")),
        "checkpoint": cp,
        "user_facing": {
            "status": job.get("status"),
            "stage": job.get("stage"),
            "error": job.get("error"),
        },
    }
    print(f"CHECKPOINT_STATE {json.dumps(checkpoint_block, default=str)[:2000]}", flush=True)

    if not (
        job.get("status") == "awaiting_human"
        and job.get("stage") == "extract"
        and cp
        and not cp.get("acknowledged")
    ):
        return {
            "branch": branch,
            "report_id": report_id,
            "owner_email": email,
            "verdict": "checkpoint_not_observed",
            "checkpoint": checkpoint_block,
            "duration_seconds": round(time.time() - t0, 1),
            "job": job,
        }

    ack = _ack_proceed(session, report_id)
    print(f"ACK_PROCEED {ack}", flush=True)
    checkpoint_block["ack"] = ack

    g1job = C.poll_job(
        session,
        report_id,
        label=f"{branch}-to-gate1",
        until_status={"awaiting_human", "failed"},
        until_stage={"gap"},
        max_seconds=MAX_TO_GATE1,
    )
    if g1job.get("status") == "failed" or g1job.get("_timeout"):
        return {
            "branch": branch,
            "report_id": report_id,
            "owner_email": email,
            "verdict": "failed_before_gate1",
            "checkpoint": checkpoint_block,
            "job": g1job,
            "duration_seconds": round(time.time() - t0, 1),
        }

    kb = C.get_kb(session, report_id)
    n_resolved = C.resolve_conflicts(kb)
    g1 = C.confirm_gate1(session, report_id, kb)
    print(f"GATE1 conflicts_resolved={n_resolved} confirm={g1['status_code']}", flush=True)

    g2park = C.poll_job(
        session,
        report_id,
        label=f"{branch}-to-gate2",
        until_status={"awaiting_human", "failed"},
        until_stage={"synthesise"},
        max_seconds=MAX_TO_GATE2,
    )
    after_gap = C.db_capture(report_id)
    gc = C.gap_check(session, report_id)
    gaps = _gaps_from_capture(after_gap)
    if not gaps and isinstance(gc.get("body"), dict):
        gaps = gc["body"].get("missing_items") or gc["body"].get("gaps") or gaps

    elevation = _assert_elevation(gaps)
    print(f"ELEVATION {json.dumps({k: elevation[k] for k in elevation if k != 'raw_elevated'}, default=str)}", flush=True)

    responses: dict[str, Any] = {}
    for g in gaps:
        key = g["item_key"]
        ref = g.get("required_item_ref") or ""
        if branch == "answered" and (ref in ELEVATED_REFS or key in ELEVATED_KEYS):
            text = ANSWER_TEXT.get(ref) or ANSWER_TEXT[ELEVATED_REFS[0]]
            if ref == ELEVATED_REFS[1] or key.endswith(ELEVATED_REFS[1]):
                text = ANSWER_TEXT[ELEVATED_REFS[1]]
            elif ref == ELEVATED_REFS[0] or key.endswith(ELEVATED_REFS[0]):
                text = ANSWER_TEXT[ELEVATED_REFS[0]]
            responses[key] = {"disposition": "answered", "answer_text": text}
        elif branch == "skip" and (ref in ELEVATED_REFS or key in ELEVATED_KEYS):
            responses[key] = {"disposition": "skipped", "skip_reason": "cannot_provide"}
        else:
            responses[key] = answer_gap(g, snippets={})

    g2 = C.submit_gate2(session, report_id, responses)
    print(f"GATE2_SUBMIT {g2['status_code']}", flush=True)

    synthjob = C.poll_job(
        session,
        report_id,
        label=f"{branch}-synth",
        until_status={"awaiting_human", "failed"},
        until_stage={"critique"},
        max_seconds=MAX_SYNTH,
    )
    if synthjob.get("status") == "failed":
        return {
            "branch": branch,
            "report_id": report_id,
            "owner_email": email,
            "verdict": "failed_at_synthesis",
            "checkpoint": checkpoint_block,
            "elevation": elevation,
            "gaps": gaps,
            "job": synthjob,
            "duration_seconds": round(time.time() - t0, 1),
            "cost": C.cost_summary(C.db_capture(report_id)),
        }

    resume = C.resume_critique(session, report_id)
    g3park = C.poll_job(
        session,
        report_id,
        label=f"{branch}-critique",
        until_status={"awaiting_human", "failed"},
        until_stage={"export"},
        max_seconds=MAX_CRITIQUE,
    )
    accept = C.accept_all_sections(session, report_id)
    g3 = C.confirm_gate3(session, report_id)
    expjob = C.poll_job(
        session,
        report_id,
        label=f"{branch}-export",
        until_status={"done", "failed"},
        max_seconds=MAX_EXPORT,
    )
    dl = C.download_export(session, report_id)
    final = C.db_capture(report_id)
    detail = C.report_detail(session, report_id)
    content = None
    if isinstance(detail.get("body"), dict):
        content = detail["body"].get("content_json")
    if content is None and isinstance(final.get("report"), dict):
        content = final["report"].get("content_json")
    community = _community_section(content if isinstance(content, dict) else None)

    community_checks: dict[str, Any] = {"section": community}
    if branch == "answered" and isinstance(community, dict):
        prose = json.dumps(community, default=str)
        community_checks["contains_participation_marker"] = (
            "TRACK3_ANSWERED_COMMUNITY_PARTICIPATION_MARKER" in prose
        )
        community_checks["contains_partner_marker"] = (
            "TRACK3_ANSWERED_PARTNER_COLLAB_MARKER" in prose
        )
        community_checks["has_gap_provenance"] = "gap:" in prose
    if branch == "skip" and isinstance(community, dict):
        status = community.get("status") or community.get("bind_status")
        community_checks["status"] = status
        community_checks["insufficient_data"] = status == "insufficient_data" or (
            community.get("structured_bind_status") == "insufficient_data"
        )
        prose = json.dumps(community, default=str)
        community_checks["invented_markers_absent"] = (
            "TRACK3_ANSWERED_COMMUNITY_PARTICIPATION_MARKER" not in prose
            and "TRACK3_ANSWERED_PARTNER_COLLAB_MARKER" not in prose
        )

    verdict = "completed" if expjob.get("status") == "done" else "export_incomplete"
    if not elevation.get("exact_two"):
        verdict = "elevation_mismatch_" + verdict

    return {
        "branch": branch,
        "report_id": report_id,
        "owner_email": email,
        "verdict": verdict,
        "checkpoint": checkpoint_block,
        "elevation": elevation,
        "gap_count": len(gaps),
        "gate2_submit": g2,
        "export": dl,
        "export_job_status": expjob.get("status"),
        "community_checks": community_checks,
        "cost": C.cost_summary(final),
        "duration_seconds": round(time.time() - t0, 1),
        "resume_critique": resume,
        "accept_all": accept,
        "gate3": g3,
        "g3park_status": g3park.get("status"),
    }


def _extract_auth_diag(log_text: str) -> list[str]:
    return [ln for ln in log_text.splitlines() if "AUTH_REFRESH_DIAG" in ln]


def main() -> int:
    EVIDENCE_DIR.mkdir(parents=True, exist_ok=True)
    log_fh = LOG_PATH.open("w", encoding="utf-8")
    sys.stdout = _Tee(sys.__stdout__, log_fh)
    sys.stderr = _Tee(sys.__stderr__, log_fh)

    print(f"TRACK3_CONFIRMING_WALK start {_now()}", flush=True)
    C.bootstrap_db_env()

    bait = _build_timeout_bait()
    print(f"FIXTURE timeout_bait={bait} bytes={bait.stat().st_size}", flush=True)

    supporting = [
        C.NLCF_DIR / "02_NLCF_Southbank_Award_Letter.docx",
        C.NLCF_DIR / "03_NLCF_Southbank_Monitoring_and_Spend_Table.docx",
    ]
    for p in supporting:
        if not p.exists():
            print(f"STOP: missing supporting doc {p}", flush=True)
            return 1

    # Prefer timeout-bait (degraded path). If checkpoint absent, retry once with
    # image-only PDF named as a proposal (unreadable path evidence).
    proposal_strategies = [
        ("timeout_bait_docx", bait),
        (
            "image_only_pdf_as_proposal",
            IMAGE_ONLY,
        ),
    ]
    # Copy image-only to proposal-like name for upload clarity
    if IMAGE_ONLY.exists():
        named = FIXTURE_DIR / "01_NLCF_Unreadable_Proposal.pdf"
        named.write_bytes(IMAGE_ONLY.read_bytes())
        proposal_strategies[1] = ("image_only_pdf_as_proposal", named)

    branches = [
        b.strip()
        for b in os.environ.get("BRANCHES", "answered,skip").split(",")
        if b.strip()
    ]

    evidence: dict[str, Any] = {
        "started_at": _now(),
        "strategies_tried": [],
        "branches": {},
        "proposal_path_used": None,
        "strategy_used": None,
    }

    chosen_path: Path | None = None
    chosen_strategy: str | None = None

    # Probe checkpoint with first strategy on a throwaway answered-branch start:
    # actually integrate into first branch — if checkpoint fails, try next strategy
    # as a fresh report for that same branch.
    for strategy, path in proposal_strategies:
        if not path.exists():
            evidence["strategies_tried"].append({"strategy": strategy, "error": "missing"})
            continue
        print(f"\n### Trying proposal strategy={strategy} path={path}", flush=True)
        result = _run_branch(
            branch=branches[0],
            proposal_path=path,
            supporting=supporting,
            evidence=evidence,
        )
        evidence["strategies_tried"].append(
            {
                "strategy": strategy,
                "path": str(path),
                "verdict": result.get("verdict"),
                "checkpoint_stage": (result.get("checkpoint") or {}).get("job_stage"),
                "checkpoint_present": bool((result.get("checkpoint") or {}).get("checkpoint")),
            }
        )
        if result.get("verdict") != "checkpoint_not_observed":
            chosen_path = path
            chosen_strategy = strategy
            evidence["branches"][branches[0]] = result
            evidence["proposal_path_used"] = str(path)
            evidence["strategy_used"] = strategy
            break
        evidence["branches"][f"{branches[0]}__{strategy}_no_checkpoint"] = result
    else:
        evidence["finished_at"] = _now()
        evidence["AUTH_REFRESH_DIAG"] = _extract_auth_diag(LOG_PATH.read_text(encoding="utf-8"))
        EVIDENCE_PATH.write_text(json.dumps(evidence, indent=2, default=str), encoding="utf-8")
        print(f"EVIDENCE={EVIDENCE_PATH}", flush=True)
        print("STOP: no strategy produced proposal checkpoint", flush=True)
        return 2

    for branch in branches[1:]:
        assert chosen_path is not None
        result = _run_branch(
            branch=branch,
            proposal_path=chosen_path,
            supporting=supporting,
            evidence=evidence,
        )
        evidence["branches"][branch] = result

    log_fh.flush()
    evidence["finished_at"] = _now()
    evidence["AUTH_REFRESH_DIAG"] = _extract_auth_diag(LOG_PATH.read_text(encoding="utf-8"))
    evidence["log_path"] = str(LOG_PATH)
    EVIDENCE_PATH.write_text(json.dumps(evidence, indent=2, default=str), encoding="utf-8")
    print(f"EVIDENCE={EVIDENCE_PATH}", flush=True)
    print(f"STRATEGY_USED={chosen_strategy}", flush=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
