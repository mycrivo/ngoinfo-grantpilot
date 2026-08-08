#!/usr/bin/env python3
"""Owner-triggered: read persisted report → ScoreableBundle JSON (local file).

Read-only. Does not call the engine or any model. Does not commit the bundle.
By default does not fetch DOCX bytes; pass --fetch-export-text to pull plaintext
from the already-persisted export object in S3 (still read-only).

Usage:
  python scripts/audit/bundle_export_run.py --railway --out /tmp/bundle.json
  python scripts/audit/bundle_export_run.py --railway --fetch-export-text --out /tmp/bundle.json
"""

from __future__ import annotations

import argparse
import json
import os
import shutil
import subprocess
import sys
from pathlib import Path

REPO = Path(__file__).resolve().parents[2]
DEFAULT_REPORT_ID = "dfd17248-9b46-48d9-8bc6-5348eab44a1c"


def _bootstrap_railway_db_url() -> str:
    railway = shutil.which("railway.cmd") or shutil.which("railway")
    if not railway:
        raise SystemExit("railway CLI not found; set DATABASE_URL or pass --railway after install")
    pg = json.loads(
        subprocess.check_output(
            [railway, "variables", "--json", "--service", "Postgres"],
            cwd=str(REPO),
            text=True,
        )
    )
    url = pg.get("DATABASE_PUBLIC_URL") or pg.get("DATABASE_URL")
    if not url:
        raise SystemExit("Postgres service has no DATABASE_PUBLIC_URL / DATABASE_URL")
    return str(url)


def _load_record(report_id: str):
    import psycopg2
    from psycopg2.extras import RealDictCursor

    from app.reports.eval.bundle_export import PersistedReportRecord

    url = os.environ.get("DATABASE_PUBLIC_URL") or os.environ.get("DATABASE_URL")
    if not url:
        raise SystemExit("No DATABASE_URL / DATABASE_PUBLIC_URL")

    conn = psycopg2.connect(url)
    conn.set_session(readonly=True, autocommit=True)
    try:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute(
            """
            SELECT id, status, version,
                   reporting_period_start, reporting_period_end,
                   knowledge_bank_json, gap_analysis_json,
                   indicator_actuals_json, content_json
            FROM donor_reports
            WHERE id = %s::uuid
            """,
            (report_id,),
        )
        row = cur.fetchone()
        if not row:
            raise SystemExit(f"report not found: {report_id}")
        cur.execute(
            """
            SELECT agent_trace_json
            FROM report_jobs
            WHERE donor_report_id = %s::uuid
            ORDER BY started_at DESC NULLS LAST, id DESC
            LIMIT 1
            """,
            (report_id,),
        )
        job = cur.fetchone()
        cur.close()
    finally:
        conn.close()

    return PersistedReportRecord(
        report_id=str(row["id"]),
        status=row.get("status"),
        version=row.get("version"),
        reporting_period_start=str(row.get("reporting_period_start")),
        reporting_period_end=str(row.get("reporting_period_end")),
        knowledge_bank_json=row.get("knowledge_bank_json"),
        gap_analysis_json=row.get("gap_analysis_json"),
        content_json=row.get("content_json"),
        indicator_actuals_json=row.get("indicator_actuals_json"),
        agent_trace_json=(job or {}).get("agent_trace_json"),
    )


def _fetch_export_plaintext(content_json: dict) -> str:
    """Read plaintext from the already-persisted export DOCX (no re-render)."""
    from io import BytesIO

    from docx import Document

    from app.reports.services.document_storage_service import DocumentStorageService

    export_meta = (content_json or {}).get("export") or {}
    storage_ref = export_meta.get("storage_ref")
    if not storage_ref:
        raise SystemExit("content_json.export.storage_ref absent — cannot fetch export text")
    storage = DocumentStorageService()
    data = storage.fetch_bytes(str(storage_ref))
    document = Document(BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--report-id", default=DEFAULT_REPORT_ID)
    parser.add_argument("--railway", action="store_true")
    parser.add_argument("--out", type=Path, required=True, help="Local path (do not commit)")
    parser.add_argument(
        "--fetch-export-text",
        action="store_true",
        help="Fetch plaintext from persisted export DOCX via storage_ref",
    )
    args = parser.parse_args()

    if args.railway:
        os.environ["DATABASE_URL"] = _bootstrap_railway_db_url()

    sys.path.insert(0, str(REPO))
    from app.reports.eval.bundle_export import (
        bundle_to_export_dict,
        export_scoreable_bundle,
    )

    record = _load_record(args.report_id)
    if args.fetch_export_text:
        record.export_plaintext = _fetch_export_plaintext(record.content_json or {})

    result = export_scoreable_bundle(record)
    payload = bundle_to_export_dict(result)
    args.out.parent.mkdir(parents=True, exist_ok=True)
    args.out.write_text(json.dumps(payload, indent=2, ensure_ascii=False, default=str) + "\n", encoding="utf-8")
    print(f"wrote {args.out}")
    print(f"stages_present={result.bundle.stages_present}")
    print(f"observations={len(result.observations)}")
    print("NOTE: do not commit this bundle file — it may contain identifiable organisation content.")


if __name__ == "__main__":
    main()
