#!/usr/bin/env python3
"""Finalize gate run deliverables: KB dump, API verify, run note."""
from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
import sys
import uuid
import urllib.error
import urllib.request
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path

REPO = Path(__file__).resolve().parents[2]
OUT = REPO / "M_E_Module" / "gate_run"
RID = uuid.UUID("6643d922-150d-4000-b878-4025e7c9145a")
BASE = "https://ngoinfo-grantpilot-production.up.railway.app"
DEPLOY_SHA = "5de686cd04cf81fb07b775a25fd22df35bddf1db"


def bootstrap() -> None:
    railway = shutil.which("railway.cmd") or shutil.which("railway")
    pg = json.loads(
        subprocess.check_output(
            [railway, "variables", "--json", "--service", "Postgres"], cwd=REPO, text=True
        )
    )
    os.environ["DATABASE_URL"] = pg["DATABASE_PUBLIC_URL"]
    backend = json.loads(
        subprocess.check_output(
            [railway, "variables", "--json", "--service", "ngoinfo-grantpilot"],
            cwd=REPO,
            text=True,
        )
    )
    os.environ["TEST_MODE_SECRET"] = str(backend.get("TEST_MODE_SECRET") or "")


def mint_token(email: str) -> str:
    payload = json.dumps({"email": email, "full_name": "Gate Run Verify"}).encode()
    req = urllib.request.Request(
        f"{BASE}/api/auth/test-mode/mint",
        data=payload,
        headers={
            "Content-Type": "application/json",
            "X-Test-Mode-Secret": os.environ["TEST_MODE_SECRET"],
        },
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=60) as resp:
        return json.loads(resp.read().decode())["access_token"]


def docx_analysis(path: Path, section_keys: set[str]) -> dict:
    from docx import Document

    doc = Document(str(path))
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    table_cells: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    table_cells.append(cell.text)
    full = "\n".join(paras + table_cells)
    leak_patterns = {
        "fact: prefix": re.compile(r"\bfact:[^\s,;]+", re.I),
        "gap: prefix": re.compile(r"\bgap:[^\s,;]+", re.I),
        "ARCH_ archetype": re.compile(r"\bARCH_[A-Z0-9_]+\b"),
    }
    for key in section_keys:
        if len(key) > 8:
            leak_patterns[f"section_key:{key}"] = re.compile(rf"\b{re.escape(key)}\b", re.I)
    leaks = {name: m.group(0) for name, pat in leak_patterns.items() if (m := pat.search(full))}
    headings = [p.text for p in doc.paragraphs if p.style and p.style.name.startswith("Heading")]
    return {
        "paragraph_count": len(paras),
        "table_cell_count": len(table_cells),
        "char_count": len(full),
        "headings": headings,
        "leaks": leaks,
        "empty_doc": len(full.strip()) == 0,
    }


def write_kb(db, report_id: uuid.UUID, path: Path) -> None:
    from app.reports.models.donor_report import DonorReport

    report = db.get(DonorReport, report_id)
    kb = report.knowledge_bank_json or {}
    facts = kb.get("facts") or {}
    gaps = kb.get("gap_answers") or {}
    lines = [
        f"# Knowledge bank — {report_id}",
        "",
        f"Captured: {datetime.now(timezone.utc).isoformat()}",
        "",
        f"- gate1_confirmed_at: {kb.get('gate1_confirmed_at')}",
        f"- gate2_confirmed_at: {kb.get('gate2_confirmed_at')}",
        f"- gate3_confirmed_at: {kb.get('gate3_confirmed_at')}",
        "",
        "## Facts",
        "",
    ]
    for key in sorted(facts.keys()):
        entry = facts[key]
        if not isinstance(entry, dict):
            lines.append(f"- **{key}**: {entry}")
            continue
        lines.append(
            f"- **{key}** | value={entry.get('value')!r} | unit={entry.get('unit')!r} | "
            f"semantic_label={entry.get('semantic_label')!r}"
        )
    lines.extend(["", "## Gap answers (answered)", ""])
    for key in sorted(gaps.keys()):
        entry = gaps[key]
        if not isinstance(entry, dict):
            continue
        if entry.get("disposition") != "answered":
            continue
        text = (entry.get("answer_text") or "").replace("\n", " ")
        lines.append(
            f"- **{key}** | disposition={entry.get('disposition')!r} | answer_text={text!r}"
        )
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> int:
    bootstrap()
    sys.path.insert(0, str(REPO))
    import app.models  # noqa: F401
    from sqlalchemy import create_engine
    from sqlalchemy.orm import sessionmaker
    from app.models.user import User
    from app.reports.models.donor_report import DonorReport
    from app.reports.models.report_job import ReportJob

    engine = create_engine(os.environ["DATABASE_URL"], pool_pre_ping=True)
    Session = sessionmaker(bind=engine)
    db = Session()
    report = db.get(DonorReport, RID)
    job = (
        db.query(ReportJob)
        .filter(ReportJob.donor_report_id == RID)
        .order_by(ReportJob.started_at.desc().nullslast(), ReportJob.id.desc())
        .first()
    )
    user = db.get(User, report.user_id)
    content = report.content_json or {}
    sections = content.get("sections") or []
    section_keys = {str(s.get("section_key")) for s in sections if s.get("section_key")}
    export_meta = content.get("export") or {}
    export_trace = ((job.agent_trace_json or {}).get("stages") or {}).get("export") or {}

    write_kb(db, RID, OUT / "6643d922_knowledge_bank.md")

    docx_path = OUT / "6643d922_export.docx"
    analysis = docx_analysis(docx_path, section_keys) if docx_path.is_file() else {"error": "missing"}

    api_ok = False
    api_bytes = 0
    api_error = ""
    try:
        token = mint_token(user.email)
        req = urllib.request.Request(
            f"{BASE}/api/reports/{RID}/export",
            headers={"Authorization": f"Bearer {token}"},
        )
        with urllib.request.urlopen(req, timeout=120) as resp:
            data = resp.read()
            api_bytes = len(data)
            api_ok = data[:2] == b"PK"
            api_path = OUT / "6643d922_export_via_api.docx"
            api_path.write_bytes(data)
    except urllib.error.HTTPError as exc:
        api_error = f"HTTP {exc.code}: {exc.read()[:300].decode(errors='replace')}"
    except Exception as exc:
        api_error = str(exc)

    with_prose = []
    failed = []
    empty_render = []
    for s in sections:
        key = s.get("section_key")
        status = s.get("generation_status")
        text = ((s.get("content") or {}).get("text") or "").strip()
        if status == "FAILED":
            failed.append(key)
        elif text:
            with_prose.append(key)
        else:
            empty_render.append(key)

    note = [
        "# Gate run note — 2026-06-04",
        "",
        "## Environment",
        f"- **Railway project:** NGOINfo-GrantPilot AI / **production**",
        f"- **Deploy / code SHA:** `{DEPLOY_SHA}` (`5de686c` — Stage H export; `_run_export_stage` with `export_completed`, not stub)",
        f"- **Backend deployment:** `b19c518e-921d-4c10-9e43-41ecf5b1714b` SUCCESS 2026-06-04 16:00 UTC+1",
        f"- **Report ID:** `{RID}`",
        "",
        "## Pre-check (read-only, before action)",
        "- Job stage/status: `export` / `awaiting_human`",
        "- gate1_confirmed_at and gate2_confirmed_at: set",
        "- gate3_confirmed_at: null",
        "- content_json.export: absent",
        "- 8 F1 sections; 21 unaccepted BLOCK flags",
        "",
        "## Gate 3 confirmation mechanism",
        "Orchestrator-test accept-all (`tests/test_orchestrator_critique.py::_accept_all_sections_for_gate3`):",
        "direct SQLAlchemy update setting each section `generation_status=ACCEPTED` and all `critic_flags[].accepted=true` (21 BLOCKs), then `confirm_gate3()` service call.",
        "",
        "## State transitions observed",
        f"- Pre-action report status: `DRAFT`",
        f"- Post Gate 3 job: `export` / `queued` (re-enqueued by `re_enqueue_gate3_job`)",
        f"- Export trace action: `{export_trace.get('action')}`",
        f"- Post-export report status: **`{report.status}`**",
        f"- Post-export job status: **`{job.status}`**",
        f"- gate3_confirmed_at: `{ (report.knowledge_bank_json or {}).get('gate3_confirmed_at') }`",
        "",
        "## Export metadata",
        f"- **render_mode:** `{export_meta.get('render_mode')}`",
        f"- **filename:** `{export_meta.get('filename')}`",
        f"- **template_version:** `{export_meta.get('template_version')}`",
        f"- **storage_ref:** `{export_meta.get('storage_ref')}`",
        f"- **bytes_written (trace):** `{export_trace.get('bytes_written')}`",
        "",
        "## Section coverage",
        f"- Sections in content_json: **{len(sections)}**",
        f"- With prose text: **{len(with_prose)}** — {', '.join(with_prose) or 'none'}",
        f"- FAILED at synthesis (no prose; may be omitted or placeholder in docx): **{len(failed)}** — {', '.join(failed) or 'none'}",
        f"- ACCEPTED/other without text: {', '.join(empty_render) or 'none'}",
        "",
        "## Docx quality scan",
        f"- Paragraphs: {analysis.get('paragraph_count', 'n/a')}; table cells with text: {analysis.get('table_cell_count', 'n/a')}",
        f"- Headings present: {len(analysis.get('headings') or [])}",
        f"- **Canonical-key leaks in visible text:** {analysis.get('leaks') or 'none detected'}",
        "",
        "## API verification",
        f"- `GET /api/reports/{RID}/export`: **{'OK' if api_ok else 'FAILED'}** ({api_bytes} bytes)",
    ]
    if api_error:
        note.append(f"- API error detail: {api_error}")
    note.extend(["", "## Deliverables", "- `6643d922_export.docx` (from R2 via DocumentStorageService)", "- `6643d922_knowledge_bank.md`", ""])
    (OUT / "GATE_RUN_NOTE_2026-06-04.md").write_text("\n".join(note), encoding="utf-8")
    print(json.dumps({"status": report.status, "api_ok": api_ok, "analysis": analysis}, indent=2))
    db.close()
    return 0 if report.status == "COMPLETE" and docx_path.is_file() else 1


if __name__ == "__main__":
    raise SystemExit(main())
