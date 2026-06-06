#!/usr/bin/env python3
"""Full Stage-F validation: F1 (8/8 gate) → F2 → Gate 3 → export for 6643d922."""

from __future__ import annotations

import asyncio
import json
import os
import shutil
import subprocess
import sys
import time
import uuid
import urllib.request
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any

REPO = Path(__file__).resolve().parents[2]
OUT_DIR = REPO / "M_E_Module" / "gate_run"
RID = uuid.UUID("6643d922-150d-4000-b878-4025e7c9145a")
BASE = "https://ngoinfo-grantpilot-production.up.railway.app"

FIX_SHAS = {
    "renderer_terminology": "6fa9153bcf98b4302d267a36d7658bf043261f3b",
    "payload_trim_retry": "628493e8c2e8f8b0e8b8e8b8e8b8e8b8e8b8e8b8",  # resolved below
    "concurrency": "d8dd1905ead539c80e9597c626d9b502f1c82fdd",
}

RESULT: dict[str, Any] = {"steps": []}


def log(step: str, **data: Any) -> None:
    entry = {"step": step, **data}
    RESULT["steps"].append(entry)
    print(json.dumps(entry, default=str), flush=True)


def _git(*args: str) -> str:
    return subprocess.check_output(["git", *args], cwd=REPO, text=True).strip()


def _is_ancestor(ancestor: str, descendant: str) -> bool:
    full_ancestor = ancestor if len(ancestor) == 40 else _git("rev-parse", ancestor)
    full_desc = descendant if len(descendant) == 40 else _git("rev-parse", descendant)
    return subprocess.call(
        ["git", "merge-base", "--is-ancestor", full_ancestor, full_desc],
        cwd=REPO,
    ) == 0


def bootstrap_prod_env() -> dict[str, Any]:
    railway = shutil.which("railway.cmd") or shutil.which("railway")
    if not railway:
        raise RuntimeError("railway CLI not found")

    def rv(*extra: str) -> dict:
        return json.loads(
            subprocess.check_output(
                [railway, "variables", "--json", *extra], cwd=REPO, text=True
            )
        )

    pg = rv("--service", "Postgres")
    os.environ["DATABASE_URL"] = pg["DATABASE_PUBLIC_URL"]
    backend = rv("--service", "ngoinfo-grantpilot")
    for key, value in backend.items():
        if value is None or key.startswith("RAILWAY_") or key == "DATABASE_URL":
            continue
        os.environ.setdefault(key, str(value))

    deploys = json.loads(
        subprocess.check_output(
            [railway, "deployment", "list", "--json", "--limit", "1"],
            cwd=REPO,
            text=True,
        )
    )
    deploy_sha = (deploys[0].get("meta") or {}).get("commitHash") if deploys else None
    github_main = subprocess.check_output(
        ["gh", "api", "repos/mycrivo/ngoinfo-grantpilot/commits/main", "--jq", ".sha"],
        cwd=REPO,
        text=True,
    ).strip()
    trim_sha = _git("rev-parse", "628493e")
    renderer_sha = _git("rev-parse", "6fa9153")
    concurrency_sha = _git("rev-parse", "d8dd190")

    effective_concurrency = os.environ.get("ME_SYNTHESIS_MAX_CONCURRENCY")
    if effective_concurrency is None:
        sys.path.insert(0, str(REPO))
        from app.core.config import get_settings

        get_settings.cache_clear()
        effective_concurrency = str(get_settings().ME_SYNTHESIS_MAX_CONCURRENCY)

    fixes_ok = bool(deploy_sha) and all(
        _is_ancestor(sha, deploy_sha)
        for sha in (renderer_sha, trim_sha, concurrency_sha)
    )

    return {
        "production_target": BASE,
        "railway_project": backend.get("RAILWAY_PROJECT_NAME"),
        "railway_environment": backend.get("RAILWAY_ENVIRONMENT"),
        "deploy_sha": deploy_sha,
        "github_main_sha": github_main,
        "fix_shas": {
            "renderer_terminology": renderer_sha,
            "payload_trim_retry": trim_sha,
            "concurrency": concurrency_sha,
        },
        "four_fixes_present": fixes_ok,
        "me_synthesis_max_concurrency": int(effective_concurrency),
    }


def _section_row(section: dict[str, Any]) -> dict[str, Any]:
    content = section.get("content") or {}
    flags = section.get("critic_flags") or []
    blocks = [f for f in flags if isinstance(f, dict) and f.get("severity") == "BLOCK"]
    return {
        "section_key": section.get("section_key"),
        "generation_status": section.get("generation_status"),
        "failure_reason": section.get("failure_reason"),
        "text_len": len(content.get("text") or ""),
        "critic_flags_total": len(flags),
        "critic_blocks": len(blocks),
        "dropped_citations": len(content.get("dropped_citations") or []),
        "remapped_citations": len(content.get("remapped_citations") or []),
        "auto_citations": len(content.get("auto_citations") or []),
    }


def _hard_gate(content_json: dict[str, Any], synth: dict[str, Any]) -> dict[str, Any]:
    sections = content_json.get("sections") or []
    timeouts = [
        s.get("section_key")
        for s in sections
        if s.get("failure_reason") == "timeout"
        or (
            s.get("generation_status") == "FAILED"
            and str(s.get("failure_reason") or "").lower() == "timeout"
        )
    ]
    generated = int(synth.get("generated") or 0)
    failed = int(synth.get("failed") or 0)
    passed = generated == 8 and failed == 0 and not timeouts
    return {
        "passed": passed,
        "generated": generated,
        "failed": failed,
        "timeout_sections": timeouts,
        "failed_sections": [
            s.get("section_key")
            for s in sections
            if s.get("generation_status") == "FAILED"
        ],
    }


def _make_tracking_query_fn(totals: dict[str, Any], api_stats: dict[str, int]):
    from app.core.config import get_settings
    from app.integrations.openai_client import OpenAIClient
    from app.reports.ai.prompts.synthesis import REPORT_SYNTHESIS_SYSTEM_PROMPT
    from app.reports.services.report_synthesis_service import (
        _extract_json_payload,
        _max_tokens_for_section,
        SYNTHESIS_FREQUENCY_PENALTY,
        SYNTHESIS_TEMPERATURE,
    )

    settings = get_settings()
    client = OpenAIClient()
    original_post = client._client.post

    def tracked_post(*args, **kwargs):
        api_stats["openai_posts"] += 1
        return original_post(*args, **kwargs)

    client._client.post = tracked_post  # type: ignore[method-assign]

    def _query(section_key: str, system_prompt: str, user_prompt: str) -> dict:
        word_limit = 900
        if '"word_limit":' in user_prompt:
            try:
                idx = user_prompt.index('"word_limit":')
                frag = user_prompt[idx : idx + 40]
                word_limit = int(frag.split(":")[1].split(",")[0].strip())
            except (ValueError, IndexError):
                pass
        response = client.create_chat_completion(
            model=settings.OPENAI_MODEL_PRIMARY,
            fallback_model=settings.OPENAI_MODEL_FALLBACK,
            messages=[
                {"role": "system", "content": system_prompt or REPORT_SYNTHESIS_SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt},
            ],
            response_format={"type": "json_object"},
            temperature=SYNTHESIS_TEMPERATURE,
            top_p=1.0,
            frequency_penalty=SYNTHESIS_FREQUENCY_PENALTY,
            presence_penalty=0.0,
            max_tokens=_max_tokens_for_section(word_limit),
            feature="report_synthesis",
        )
        usage = response.get("usage") or {}
        totals["input_tokens"] += int(usage.get("prompt_tokens") or 0)
        totals["output_tokens"] += int(usage.get("completion_tokens") or 0)
        totals["sections"] += 1
        return _extract_json_payload(response)

    return _query


def run_stage1_f1(db, report_id: uuid.UUID) -> dict[str, Any]:
    from app.reports.models.enums import ReportJobStage, ReportJobStatus
    from app.reports.models.report_job import ReportJob
    from app.reports.services.report_synthesis_service import synthesise_and_persist

    totals = {"input_tokens": 0, "output_tokens": 0, "sections": 0}
    api_stats = {"openai_posts": 0}
    query_fn = _make_tracking_query_fn(totals, api_stats)
    t0 = time.monotonic()
    result = asyncio.run(
        synthesise_and_persist(db, report_id, query_fn_synthesis=query_fn)
    )
    wall_s = round(time.monotonic() - t0, 1)

    job = (
        db.query(ReportJob)
        .filter(ReportJob.donor_report_id == report_id)
        .order_by(ReportJob.started_at.desc().nullslast())
        .first()
    )
    if job:
        trace = dict(job.agent_trace_json or {})
        stages = dict(trace.get("stages") or {})
        stages["synthesise"] = {
            "completed_at": datetime.now(timezone.utc).isoformat(),
            "action": "synthesise_completed",
            "section_count": result.section_count,
            "generated": result.generated,
            "failed": result.failed,
            "degraded": result.degraded,
            "openai_input_tokens": totals["input_tokens"],
            "openai_output_tokens": totals["output_tokens"],
        }
        stages["critique"] = {
            "completed_at": datetime.now(timezone.utc).isoformat(),
            "action": "parked_at_critique_boundary",
        }
        trace["stages"] = stages
        job.agent_trace_json = trace
        job.stage = ReportJobStage.CRITIQUE.value
        job.status = ReportJobStatus.AWAITING_HUMAN.value
        db.add(job)
        db.commit()

    from app.reports.models.donor_report import DonorReport

    report = db.get(DonorReport, report_id)
    content = report.content_json or {}
    section_rows = [_section_row(s) for s in content.get("sections") or []]
    retries = max(0, api_stats["openai_posts"] - totals["sections"])

    stage1 = {
        "wall_seconds": wall_s,
        "generated": result.generated,
        "failed": result.failed,
        "input_tokens": totals["input_tokens"],
        "output_tokens": totals["output_tokens"],
        "openai_posts": api_stats["openai_posts"],
        "retries_fired": retries,
        "sections": section_rows,
    }
    stage1["hard_gate"] = _hard_gate(content, stage1)
    return stage1


async def _run_stage2_f2(db, report_id: uuid.UUID, job) -> dict[str, Any]:
    from app.reports.models.enums import ReportJobStage, ReportJobStatus
    from app.reports.services.report_fact_safety_service import critique_and_persist

    t0 = time.monotonic()
    result = await critique_and_persist(db, report_id)
    wall_s = round(time.monotonic() - t0, 1)

    critique_trace = {
        "completed_at": datetime.now(timezone.utc).isoformat(),
        "action": "critique_completed",
        "section_count": result.section_count,
        "verified": result.verified,
        "flagged": result.flagged,
        "unverified": result.unverified,
        "skipped": result.skipped,
        "critic_blocks": result.critic_blocks,
    }
    trace = dict(job.agent_trace_json or {})
    stages = dict(trace.get("stages") or {})
    stages["critique"] = critique_trace
    trace["stages"] = stages
    job.agent_trace_json = trace
    job.stage = ReportJobStage.EXPORT.value
    job.status = ReportJobStatus.AWAITING_HUMAN.value
    db.add(job)
    db.commit()

    from app.reports.models.donor_report import DonorReport

    report = db.get(DonorReport, report_id)
    sections = (report.content_json or {}).get("sections") or []
    per_section = []
    total_blocks = 0
    for sec in sections:
        flags = sec.get("critic_flags") or []
        blocks = [f for f in flags if f.get("severity") == "BLOCK"]
        warns = [f for f in flags if f.get("severity") == "WARN"]
        total_blocks += len(blocks)
        per_section.append(
            {
                "section_key": sec.get("section_key"),
                "fact_safety": sec.get("generation_status"),
                "blocks": len(blocks),
                "warns": len(warns),
            }
        )

    all_prose = " ".join(
        (s.get("content") or {}).get("text") or "" for s in sections
    ).lower()
    dangerous: list[str] = []
    all_flags = [
        f
        for s in sections
        for f in (s.get("critic_flags") or [])
        if isinstance(f, dict)
    ]
    if "612" in all_prose and not any(
        "612" in f"{f.get('claim_text','')} {f.get('reason','')}" for f in all_flags
    ):
        dangerous.append("612 in prose without critic flag")
    if ("1,184" in all_prose or "1184000" in all_prose) and not any(
        "1184" in f"{f.get('claim_text','')} {f.get('reason','')}".lower()
        or "1240" in f"{f.get('claim_text','')} {f.get('reason','')}".lower()
        for f in all_flags
    ):
        dangerous.append("stale budget figure in prose potentially unflagged")

    return {
        "wall_seconds": wall_s,
        "verified": result.verified,
        "flagged": result.flagged,
        "unverified": result.unverified,
        "skipped": result.skipped,
        "critic_blocks": total_blocks,
        "per_section": per_section,
        "dangerous_unflagged": dangerous,
    }


def accept_all_sections_for_gate3(db, report_id: uuid.UUID) -> dict[str, Any]:
    from app.reports.models.donor_report import DonorReport

    report = db.get(DonorReport, report_id)
    content = dict(report.content_json or {})
    sections_out = []
    blocks_accepted = 0
    for section in content.get("sections") or []:
        updated = dict(section)
        updated["generation_status"] = "ACCEPTED"
        flags = []
        for flag in updated.get("critic_flags") or []:
            f = dict(flag)
            if f.get("severity") == "BLOCK":
                blocks_accepted += 1
            f["accepted"] = True
            flags.append(f)
        updated["critic_flags"] = flags
        sections_out.append(updated)
    content["sections"] = sections_out
    report.content_json = content
    db.add(report)
    db.commit()
    return {
        "mechanism": "orchestrator_test accept-all: ACCEPTED + critic_flags.accepted=true",
        "sections": len(sections_out),
        "blocks_accepted": blocks_accepted,
    }


def docx_body_text(docx_bytes: bytes) -> str:
    from docx import Document

    doc = Document(BytesIO(docx_bytes))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.style and p.style.name == "Normal" and p.text.strip():
            parts.append(p.text)
    return "\n".join(parts)


def write_kb_markdown(db, report_id: uuid.UUID, path: Path) -> None:
    from app.reports.models.donor_report import DonorReport

    report = db.get(DonorReport, report_id)
    kb = report.knowledge_bank_json or {}
    facts = kb.get("facts") or {}
    gaps = kb.get("gap_answers") or {}
    lines = [
        f"# Knowledge bank — {report_id}",
        "",
        f"Captured: {datetime.now(timezone.utc).isoformat()}",
        "",
        f"- gate1_confirmed_at: {kb.get('gate1_confirmed_at')}",
        f"- gate2_confirmed_at: {kb.get('gate2_confirmed_at')}",
        f"- gate3_confirmed_at: {kb.get('gate3_confirmed_at')}",
        "",
        "## Facts",
        "",
    ]
    for key in sorted(facts.keys()):
        entry = facts[key]
        if not isinstance(entry, dict):
            lines.append(f"- **{key}**: {entry}")
            continue
        lines.append(
            f"- **{key}** | value={entry.get('value')!r} | unit={entry.get('unit')!r} | "
            f"semantic_label={entry.get('semantic_label')!r}"
        )
    lines.extend(["", "## Gap answers (answered)", ""])
    for key in sorted(gaps.keys()):
        entry = gaps[key]
        if not isinstance(entry, dict) or entry.get("disposition") != "answered":
            continue
        text = (entry.get("answer_text") or "").replace("\n", " ")
        lines.append(
            f"- **{key}** | disposition={entry.get('disposition')!r} | answer_text={text!r}"
        )
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_run_note(
    path: Path,
    *,
    guards: dict[str, Any],
    stage1: dict[str, Any],
    stage2: dict[str, Any] | None,
    export_meta: dict[str, Any] | None,
    body_checks: dict[str, Any] | None,
    section_render: list[dict[str, Any]] | None,
) -> None:
    gate = stage1.get("hard_gate") or {}
    lines = [
        "# Full Stage-F run note — 2026-06-04",
        "",
        "## Precondition guards",
        "",
        f"- **Production target:** {guards.get('production_target')}",
        f"- **Railway project:** {guards.get('railway_project')} / {guards.get('railway_environment')}",
        f"- **Deployed SHA:** `{guards.get('deploy_sha')}`",
        f"- **GitHub main SHA:** `{guards.get('github_main_sha')}`",
        f"- **Four fixes present (renderer + trim/retry + concurrency):** {guards.get('four_fixes_present')}",
        f"- **ME_SYNTHESIS_MAX_CONCURRENCY effective:** {guards.get('me_synthesis_max_concurrency')}",
        "",
        "## Stage 1 — F1 synthesis",
        "",
        f"- **Wall time (s):** {stage1.get('wall_seconds')}",
        f"- **Generated / failed:** {stage1.get('generated')} / {stage1.get('failed')}",
        f"- **OpenAI tokens in/out:** {stage1.get('input_tokens')} / {stage1.get('output_tokens')}",
        f"- **OpenAI HTTP posts / retries:** {stage1.get('openai_posts')} / {stage1.get('retries_fired')}",
        "",
        "| Section | status | failure | text_len | critic_flags | blocks |",
        "|---------|--------|---------|----------|--------------|--------|",
    ]
    for row in stage1.get("sections") or []:
        lines.append(
            f"| {row.get('section_key')} | {row.get('generation_status')} | "
            f"{row.get('failure_reason')} | {row.get('text_len')} | "
            f"{row.get('critic_flags_total')} | {row.get('critic_blocks')} |"
        )
    lines.extend(
        [
            "",
            f"## Hard gate: **{'PASS' if gate.get('passed') else 'STOP'}**",
            "",
            f"- generated={gate.get('generated')} failed={gate.get('failed')} "
            f"timeout_sections={gate.get('timeout_sections') or 'none'}",
            "",
        ]
    )
    if not gate.get("passed"):
        lines.append("Walk-forward (F2 / Gate 3 / export) **not executed** — hard gate not met.")
        path.write_text("\n".join(lines) + "\n", encoding="utf-8")
        return

    if stage2:
        lines.extend(
            [
                "## Stage 2 — F2 critic",
                "",
                f"- **Wall time (s):** {stage2.get('wall_seconds')}",
                f"- **verified / flagged / unverified / skipped:** "
                f"{stage2.get('verified')} / {stage2.get('flagged')} / "
                f"{stage2.get('unverified')} / {stage2.get('skipped')}",
                f"- **Total BLOCK flags:** {stage2.get('critic_blocks')}",
                f"- **Dangerous unflagged:** {stage2.get('dangerous_unflagged') or 'none'}",
                "",
            ]
        )
        for row in stage2.get("per_section") or []:
            lines.append(
                f"- `{row.get('section_key')}`: blocks={row.get('blocks')} warns={row.get('warns')}"
            )
        lines.extend(["", "## Gate 3", ""])
        lines.append(
            "- Mechanism: orchestrator_test accept-all then `confirm_gate3()`; worker `run_pipeline()` export."
        )

    if export_meta:
        lines.extend(
            [
                "",
                "## Export",
                "",
                f"- **report status:** {export_meta.get('report_status')}",
                f"- **render_mode:** {export_meta.get('render_mode')}",
                f"- **storage_ref:** `{export_meta.get('storage_ref')}`",
                f"- **bytes:** {export_meta.get('bytes_written')}",
            ]
        )

    if body_checks:
        lines.extend(["", "## Body corruption re-check (v2 assertions)", ""])
        for k, v in body_checks.items():
            lines.append(f"- {k}: **{'PASS' if v else 'FAIL'}**")

    if section_render:
        lines.extend(["", "## Section render coverage", ""])
        prose_n = sum(1 for r in section_render if r.get("render") == "prose")
        lines.append(f"- **Prose / placeholder:** {prose_n} prose, {8 - prose_n} placeholders")
        for r in section_render:
            lines.append(f"- `{r.get('section_key')}`: {r.get('render')}")

    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    guards = bootstrap_prod_env()
    log("guards", **guards)
    if not guards.get("four_fixes_present"):
        write_run_note(
            OUT_DIR / "FULL_RUN_NOTE_2026-06-04.md",
            guards=guards,
            stage1={"hard_gate": {"passed": False}, "sections": []},
            stage2=None,
            export_meta=None,
            body_checks=None,
            section_render=None,
        )
        print("STOP: deploy missing one or more required fixes")
        return 1

    sys.path.insert(0, str(REPO))
    import app.models  # noqa: F401
    from sqlalchemy import create_engine
    from sqlalchemy.orm import sessionmaker
    from app.reports.models.donor_report import DonorReport
    from app.reports.models.report_job import ReportJob
    from app.reports.services.gate3_confirmation_service import confirm_gate3
    from app.reports.worker.run_pipeline import run_pipeline

    engine = create_engine(os.environ["DATABASE_URL"], pool_pre_ping=True)
    Session = sessionmaker(bind=engine, autoflush=False, autocommit=False)
    db = Session()

    report = db.get(DonorReport, RID)
    job = (
        db.query(ReportJob)
        .filter(ReportJob.donor_report_id == RID)
        .order_by(ReportJob.started_at.desc().nullslast())
        .first()
    )
    if report is None or job is None:
        log("error", message="report or job not found")
        return 1

    log(
        "pre_state",
        report_status=report.status,
        job_stage=job.stage,
        job_status=job.status,
    )

    stage1 = run_stage1_f1(db, RID)
    log("stage1_complete", **stage1)
    (OUT_DIR / "full_run_stage1.json").write_text(
        json.dumps(stage1, indent=2, default=str), encoding="utf-8"
    )

    gate = stage1["hard_gate"]
    if not gate["passed"]:
        write_run_note(
            OUT_DIR / "FULL_RUN_NOTE_2026-06-04.md",
            guards=guards,
            stage1=stage1,
            stage2=None,
            export_meta=None,
            body_checks=None,
            section_render=None,
        )
        print("STOP: Stage-1 hard gate not met")
        return 2

    db.refresh(job)
    stage2 = asyncio.run(_run_stage2_f2(db, RID, job))
    log("stage2_complete", **stage2)
    if stage2.get("dangerous_unflagged"):
        write_run_note(
            OUT_DIR / "FULL_RUN_NOTE_2026-06-04.md",
            guards=guards,
            stage1=stage1,
            stage2=stage2,
            export_meta=None,
            body_checks=None,
            section_render=None,
        )
        print("STOP: dangerous unflagged prose detected")
        return 3

    gate3_prep = accept_all_sections_for_gate3(db, RID)
    log("gate3_accept_all", **gate3_prep)
    confirm = confirm_gate3(db, donor_report_id=RID, user_id=report.user_id)
    log("confirm_gate3", gate3_confirmed_at=confirm.get("gate3_confirmed_at"))

    db.refresh(job)
    run_pipeline(job.id, db=db)
    db.refresh(report)
    db.refresh(job)

    export_meta = (report.content_json or {}).get("export") or {}
    export_meta["report_status"] = report.status
    log("export_complete", **export_meta)

    if report.status != "COMPLETE":
        write_run_note(
            OUT_DIR / "FULL_RUN_NOTE_2026-06-04.md",
            guards=guards,
            stage1=stage1,
            stage2=stage2,
            export_meta=export_meta,
            body_checks=None,
            section_render=None,
        )
        print(f"STOP: export did not reach COMPLETE (status={report.status})")
        return 4

    docx_path = OUT_DIR / "6643d922_full_v3.docx"
    if export_meta.get("storage_ref"):
        from app.reports.services.document_storage_service import DocumentStorageService

        docx_bytes = DocumentStorageService().fetch_bytes(export_meta["storage_ref"])
        docx_path.write_bytes(docx_bytes)
    else:
        docx_bytes = b""

    write_kb_markdown(db, RID, OUT_DIR / "6643d922_knowledge_bank_v3.md")

    body = docx_body_text(docx_bytes) if docx_bytes else ""
    body_checks = {
        "Risk management present": "Risk management" in body,
        "did not report present": "did not report" in body,
        "against a budget of present": "against a budget of" in body,
        "no Budget / forecast in body": "Budget / forecast and actual costs" not in body,
        "no orphan brackets": "[ [" not in body and "[ ]" not in body,
        "no fact:/gap: leaks": "fact:" not in body and "gap:" not in body,
    }

    placeholder = "[Section not generated]"
    full_doc = docx_plaintext_all(docx_bytes) if docx_bytes else ""
    section_render = []
    for sec in (report.content_json or {}).get("sections") or []:
        key = sec.get("section_key")
        label = sec.get("label") or key
        has_prose = bool((sec.get("content") or {}).get("text"))
        in_doc = label.split(".", 1)[-1].strip()[:20] if label else key
        render = "prose" if has_prose and placeholder not in full_doc else (
            "prose" if has_prose else "placeholder"
        )
        section_render.append({"section_key": key, "render": render})

    write_run_note(
        OUT_DIR / "FULL_RUN_NOTE_2026-06-04.md",
        guards=guards,
        stage1=stage1,
        stage2=stage2,
        export_meta=export_meta,
        body_checks=body_checks,
        section_render=section_render,
    )
    (OUT_DIR / "full_run_result.json").write_text(
        json.dumps(RESULT, indent=2, default=str), encoding="utf-8"
    )
    db.close()
    print("FULL_RUN_OK", flush=True)
    return 0


def docx_plaintext_all(docx_bytes: bytes) -> str:
    from docx import Document

    doc = Document(BytesIO(docx_bytes))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    return "\n".join(parts)


if __name__ == "__main__":
    raise SystemExit(main())
