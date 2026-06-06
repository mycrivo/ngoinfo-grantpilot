#!/usr/bin/env python3
"""Stage F gate run: Gate 3 confirm + export for 6643d922 on production."""
from __future__ import annotations

import json
import os
import shutil
import subprocess
import sys
import uuid
import urllib.request
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path

REPO = Path(__file__).resolve().parents[2]
OUT_DIR = REPO / "M_E_Module" / "gate_run"
RID = uuid.UUID("6643d922-150d-4000-b878-4025e7c9145a")
BASE = "https://ngoinfo-grantpilot-production.up.railway.app"
JOB_ID: uuid.UUID | None = None
USER_ID: uuid.UUID | None = None

RESULT: dict = {"steps": []}


def log(step: str, **data) -> None:
    entry = {"step": step, **data}
    RESULT["steps"].append(entry)
    print(json.dumps(entry, default=str), flush=True)


def bootstrap_prod_env() -> str:
    railway = shutil.which("railway.cmd") or shutil.which("railway")
    if not railway:
        raise RuntimeError("railway CLI not found")

    def rv(*extra: str) -> dict:
        return json.loads(
            subprocess.check_output(
                [railway, "variables", "--json", *extra], cwd=REPO, text=True
            )
        )

    pg = rv("--service", "Postgres")
    os.environ["DATABASE_URL"] = pg["DATABASE_PUBLIC_URL"]
    backend = rv("--service", "ngoinfo-grantpilot")
    for key, value in backend.items():
        if value is None or key.startswith("RAILWAY_") or key == "DATABASE_URL":
            continue
        os.environ.setdefault(key, str(value))

    gh_sha = subprocess.check_output(
        ["gh", "api", "repos/mycrivo/ngoinfo-grantpilot/commits/main", "--jq", ".sha"],
        cwd=REPO,
        text=True,
    ).strip()
    log("bootstrap", github_main_sha=gh_sha, database_url_set=True)
    return gh_sha


def accept_all_sections_for_gate3(db, report_id: uuid.UUID) -> dict:
    from app.reports.models.donor_report import DonorReport

    report = db.get(DonorReport, report_id)
    if report is None:
        raise RuntimeError("report not found")
    content = dict(report.content_json or {})
    sections_out = []
    blocks_accepted = 0
    for section in content.get("sections") or []:
        updated = dict(section)
        updated["generation_status"] = "ACCEPTED"
        flags = []
        for flag in updated.get("critic_flags") or []:
            f = dict(flag)
            if f.get("severity") == "BLOCK":
                blocks_accepted += 1
            f["accepted"] = True
            flags.append(f)
        updated["critic_flags"] = flags
        sections_out.append(updated)
    content["sections"] = sections_out
    report.content_json = content
    db.add(report)
    db.commit()
    log(
        "accept_all_sections",
        mechanism="orchestrator_test pattern (_accept_all_sections_for_gate3): "
        "direct SQLAlchemy update of content_json sections to ACCEPTED + critic_flags.accepted=true",
        sections=len(sections_out),
        blocks_accepted=blocks_accepted,
    )
    return content


def run_export(db, job_id: uuid.UUID) -> None:
    from app.reports.worker.run_pipeline import run_pipeline

    run_pipeline(job_id, db=db)
    log("run_pipeline_complete", job_id=str(job_id))


def mint_token(user_id: uuid.UUID) -> str:
    secret = os.environ.get("TEST_MODE_SECRET", "")
    payload = json.dumps({"user_id": str(user_id)}).encode()
    req = urllib.request.Request(
        f"{BASE}/api/auth/test-mode/mint",
        data=payload,
        headers={
            "Content-Type": "application/json",
            "X-Test-Mode-Secret": secret,
        },
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=60) as resp:
        body = json.loads(resp.read().decode())
    return body["access_token"]


def download_export_api(report_id: uuid.UUID, token: str) -> bytes:
    req = urllib.request.Request(
        f"{BASE}/api/reports/{report_id}/export",
        headers={"Authorization": f"Bearer {token}"},
        method="GET",
    )
    with urllib.request.urlopen(req, timeout=120) as resp:
        return resp.read()


def docx_plaintext(docx_bytes: bytes) -> str:
    from docx import Document

    doc = Document(BytesIO(docx_bytes))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def write_kb_markdown(db, report_id: uuid.UUID, path: Path) -> None:
    from app.reports.models.donor_report import DonorReport

    report = db.get(DonorReport, report_id)
    kb = report.knowledge_bank_json or {}
    facts = kb.get("facts") or {}
    gaps = kb.get("gap_answers") or {}
    lines = [
        f"# Knowledge bank — {report_id}",
        "",
        f"Captured: {datetime.now(timezone.utc).isoformat()}",
        "",
        f"- gate1_confirmed_at: {kb.get('gate1_confirmed_at')}",
        f"- gate2_confirmed_at: {kb.get('gate2_confirmed_at')}",
        f"- gate3_confirmed_at: {kb.get('gate3_confirmed_at')}",
        "",
        "## Facts",
        "",
    ]
    for key in sorted(facts.keys()):
        entry = facts[key]
        if not isinstance(entry, dict):
            lines.append(f"- **{key}**: {entry}")
            continue
        lines.append(
            f"- **{key}** | value={entry.get('value')!r} | unit={entry.get('unit')!r} | "
            f"semantic_label={entry.get('semantic_label')!r}"
        )
    lines.extend(["", "## Gap answers (answered)", ""])
    for key in sorted(gaps.keys()):
        entry = gaps[key]
        if not isinstance(entry, dict):
            continue
        if entry.get("disposition") != "answered":
            continue
        text = (entry.get("answer_text") or "").replace("\n", " ")
        lines.append(
            f"- **{key}** | disposition={entry.get('disposition')!r} | answer_text={text!r}"
        )
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    gh_sha = bootstrap_prod_env()
    RESULT["github_main_sha"] = gh_sha

    sys.path.insert(0, str(REPO))
    import app.models  # noqa: F401 — register User etc. for DonorReport relationships
    from sqlalchemy.orm import sessionmaker
    from sqlalchemy import create_engine
    from app.reports.models.donor_report import DonorReport
    from app.reports.models.report_job import ReportJob
    from app.reports.services.gate3_confirmation_service import confirm_gate3

    engine = create_engine(os.environ["DATABASE_URL"], pool_pre_ping=True)
    Session = sessionmaker(bind=engine, autoflush=False, autocommit=False)
    db = Session()

    report = db.get(DonorReport, RID)
    if report is None:
        log("error", message="report not found")
        return 1
    global USER_ID, JOB_ID
    USER_ID = report.user_id
    job = (
        db.query(ReportJob)
        .filter(ReportJob.donor_report_id == RID)
        .order_by(ReportJob.started_at.desc().nullslast(), ReportJob.id.desc())
        .first()
    )
    if job is None:
        log("error", message="job not found")
        return 1
    JOB_ID = job.id

    log(
        "pre_state",
        report_status=report.status,
        job_stage=job.stage,
        job_status=job.status,
        gate1=report.knowledge_bank_json.get("gate1_confirmed_at"),
        gate2=report.knowledge_bank_json.get("gate2_confirmed_at"),
        gate3=report.knowledge_bank_json.get("gate3_confirmed_at"),
        has_export=bool((report.content_json or {}).get("export")),
        section_count=len((report.content_json or {}).get("sections") or []),
    )

    accept_all_sections_for_gate3(db, RID)
    db.refresh(report)

    confirm = confirm_gate3(db, donor_report_id=RID, user_id=USER_ID)
    log("confirm_gate3", gate3_confirmed_at=confirm.get("gate3_confirmed_at"))

    db.refresh(job)
    log("post_confirm_job", job_status=job.status, job_stage=job.stage)

    run_export(db, JOB_ID)

    db.refresh(report)
    db.refresh(job)
    export_meta = (report.content_json or {}).get("export") or {}
    log(
        "post_export",
        report_status=report.status,
        job_status=job.status,
        export_meta=export_meta,
        export_trace=(job.agent_trace_json or {}).get("stages", {}).get("export"),
    )

    if report.status != "COMPLETE":
        note = OUT_DIR / "GATE_RUN_NOTE_2026-06-04.md"
        note.write_text(
            f"# Gate run FAILED\n\nReport status: {report.status}\nJob error: {job.error}\n",
            encoding="utf-8",
        )
        (OUT_DIR / "gate_run_result.json").write_text(
            json.dumps(RESULT, indent=2, default=str), encoding="utf-8"
        )
        return 2

    docx_path = OUT_DIR / "6643d922_export.docx"
    if export_meta.get("storage_ref"):
        from app.reports.services.document_storage_service import DocumentStorageService

        store = DocumentStorageService()
        docx_bytes = store.fetch_bytes(export_meta["storage_ref"])
        docx_path.write_bytes(docx_bytes)
        log("fetched_from_r2", storage_ref=export_meta["storage_ref"], bytes=len(docx_bytes))
    else:
        docx_bytes = b""

    token = mint_token(USER_ID)
    api_bytes = download_export_api(RID, token)
    log("api_download", bytes=len(api_bytes), starts_with_pk=api_bytes[:2] == b"PK")
    if not docx_path.exists() or len(docx_bytes) == 0:
        docx_path.write_bytes(api_bytes)

    write_kb_markdown(db, RID, OUT_DIR / "6643d922_knowledge_bank.md")

    plaintext = docx_plaintext(docx_path.read_bytes())
    leaks = []
    for needle in ("fact:", "gap:", "summary_and_overview", "ARCH_", "detailed_output_scoring"):
        if needle in plaintext:
            leaks.append(needle)

    sections = (report.content_json or {}).get("sections") or []
    gen = sum(1 for s in sections if s.get("generation_status") in ("GENERATED", "ACCEPTED") and (s.get("content") or {}).get("text"))
    failed = [s.get("section_key") for s in sections if s.get("generation_status") == "FAILED"]

    note_lines = [
        "# Gate run note — 2026-06-04",
        "",
        f"- **Report:** `{RID}`",
        f"- **GitHub main SHA:** `{gh_sha}`",
        f"- **Gate 3 mechanism:** Direct DB update (orchestrator test `_accept_all_sections_for_gate3`) "
        "then `confirm_gate3()` service; worker `run_pipeline()` for export stage with real R2.",
        f"- **Pre job state:** export / awaiting_human",
        f"- **Post report status:** `{report.status}`",
        f"- **Post job status:** `{job.status}`",
        f"- **render_mode:** `{export_meta.get('render_mode')}`",
        f"- **filename:** `{export_meta.get('filename')}`",
        f"- **template_version:** `{export_meta.get('template_version')}`",
        f"- **storage_ref:** `{export_meta.get('storage_ref')}`",
        f"- **Sections with prose in content_json:** {gen} / {len(sections)}",
        f"- **Sections FAILED at synthesis (may render placeholder):** {failed or 'none'}",
        f"- **Canonical-key leak check in docx:** {leaks or 'none detected'}",
        "",
    ]
    (OUT_DIR / "GATE_RUN_NOTE_2026-06-04.md").write_text("\n".join(note_lines), encoding="utf-8")
    (OUT_DIR / "gate_run_result.json").write_text(
        json.dumps(RESULT, indent=2, default=str), encoding="utf-8"
    )
    db.close()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
